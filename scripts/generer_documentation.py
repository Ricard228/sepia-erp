"""Génère la documentation Word de la plateforme SEPIA.

Usage :  python scripts/generer_documentation.py [chemin_de_sortie.docx]
Le document produit décrit l'ensemble des fonctionnalités, la méthodologie de
suivi-évaluation intégrée, l'architecture technique, le guide d'utilisation et
la procédure de déploiement.
"""
import os
import sys
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

RACINE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, RACINE)

from app.config import APP_LONG_NAME, APP_NAME, APP_VERSION  # noqa: E402

BLEU = RGBColor(0x1F, 0x4E, 0x79)
BLEU_MOYEN = RGBColor(0x2E, 0x75, 0xB6)
GRIS = RGBColor(0x55, 0x55, 0x55)
MOIS_FR = ["janvier", "février", "mars", "avril", "mai", "juin", "juillet", "août",
           "septembre", "octobre", "novembre", "décembre"]


# ---------------------------------------------------------------------------
# Utilitaires de mise en forme
# ---------------------------------------------------------------------------
def ombrer(cellule, couleur_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), couleur_hex)
    cellule._tc.get_or_add_tcPr().append(shd)


def texte_cellule(cellule, texte, gras=False, taille=9, blanc=False, centre=False):
    cellule.text = ""
    paragraphe = cellule.paragraphs[0]
    paragraphe.paragraph_format.space_after = Pt(2)
    if centre:
        paragraphe.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraphe.add_run("" if texte is None else str(texte))
    run.bold = gras
    run.font.size = Pt(taille)
    if blanc:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def tableau(document, entetes, lignes, largeurs=None, taille=9):
    table = document.add_table(rows=1, cols=len(entetes))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, titre in enumerate(entetes):
        cellule = table.rows[0].cells[index]
        ombrer(cellule, "1F4E79")
        texte_cellule(cellule, titre, gras=True, blanc=True, centre=True, taille=taille)
        if largeurs:
            cellule.width = Cm(largeurs[index])
    for ligne in lignes:
        cellules = table.add_row().cells
        for index, valeur in enumerate(ligne):
            texte_cellule(cellules[index], valeur, taille=taille,
                          gras=(index == 0 and len(entetes) > 2))
            if largeurs:
                cellules[index].width = Cm(largeurs[index])
    document.add_paragraph()
    return table


def titre1(document, texte):
    h = document.add_heading(texte, level=1)
    h.runs[0].font.color.rgb = BLEU
    h.runs[0].font.size = Pt(16)
    return h


def titre2(document, texte):
    h = document.add_heading(texte, level=2)
    h.runs[0].font.color.rgb = BLEU_MOYEN
    h.runs[0].font.size = Pt(13)
    return h


def titre3(document, texte):
    h = document.add_heading(texte, level=3)
    h.runs[0].font.color.rgb = BLEU_MOYEN
    h.runs[0].font.size = Pt(11)
    return h


def para(document, texte, italique=False, taille=10.5, justifie=True):
    p = document.add_paragraph()
    if justifie:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(texte)
    run.italic = italique
    run.font.size = Pt(taille)
    return p


def puce(document, texte, taille=10):
    p = document.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(texte)
    run.font.size = Pt(taille)
    return p


def numero(document, texte, taille=10):
    p = document.add_paragraph(style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(texte)
    run.font.size = Pt(taille)
    return p


def encadre(document, titre_encadre, texte):
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cellule = table.rows[0].cells[0]
    ombrer(cellule, "DCE6F1")
    cellule.text = ""
    p = cellule.paragraphs[0]
    run = p.add_run(titre_encadre + " — ")
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = BLEU
    run2 = p.add_run(texte)
    run2.font.size = Pt(9.5)
    document.add_paragraph()


def code(document, lignes):
    for ligne in lignes:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.6)
        run = p.add_run(ligne)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x33)
    document.add_paragraph()


def pied_de_page(document, texte):
    for section in document.sections:
        paragraphe = section.footer.paragraphs[0]
        paragraphe.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraphe.text = ""
        run = paragraphe.add_run(texte)
        run.font.size = Pt(7.5)
        run.font.color.rgb = GRIS


# ---------------------------------------------------------------------------
# Construction du document
# ---------------------------------------------------------------------------
def construire(chemin_sortie: str) -> str:
    document = Document()
    section = document.sections[0]
    section.left_margin = section.right_margin = Cm(2.2)
    section.top_margin = section.bottom_margin = Cm(2)
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    aujourdhui = date.today()
    date_fr = f"{aujourdhui.day} {MOIS_FR[aujourdhui.month - 1]} {aujourdhui.year}"

    # ---------------------------------------------------------------- Garde
    for _ in range(4):
        document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PLATEFORME " + APP_NAME)
    run.bold = True
    run.font.size = Pt(40)
    run.font.color.rgb = BLEU
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(APP_LONG_NAME)
    run.font.size = Pt(13)
    run.italic = True
    run.font.color.rgb = GRIS
    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nERP de planification et de suivi-évaluation\ndes projets et programmes "
                    "de développement")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = BLEU_MOYEN
    for _ in range(3):
        document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DOCUMENT DE DESCRIPTION FONCTIONNELLE ET TECHNIQUE")
    run.bold = True
    run.font.size = Pt(14)
    for _ in range(5):
        document.add_paragraph()
    tableau(document,
            ["Rubrique", "Information"],
            [["Intitulé de la solution", f"{APP_NAME} — {APP_LONG_NAME}"],
             ["Nature", "Application web et mobile (ERP de suivi-évaluation)"],
             ["Version documentée", APP_VERSION],
             ["Technologies", "Python 3.12 · FastAPI · SQLAlchemy · PostgreSQL · "
                              "JavaScript natif (sans dépendance externe)"],
             ["Dépôt de code source", "https://github.com/Ricard228/sepia-erp"],
             ["Hébergement cible", "Render (service web + base PostgreSQL)"],
             ["Date d'édition", date_fr]],
            largeurs=[5.5, 11], taille=10)

    document.add_page_break()

    # -------------------------------------------------------------- Sommaire
    titre1(document, "Sommaire")
    sommaire = [
        ("1.", "Contexte, justification et finalité de la plateforme"),
        ("2.", "Périmètre fonctionnel"),
        ("3.", "Architecture générale de la solution"),
        ("4.", "Modèle de données"),
        ("5.", "Description détaillée des dix-neuf modules"),
        ("6.", "Méthodologie de suivi-évaluation intégrée"),
        ("7.", "Import de données : Excel, Word, XLSForm, KoboToolbox"),
        ("8.", "Livrables générés automatiquement"),
        ("9.", "Fiches de collecte et questionnaires numériques"),
        ("10.", "Tableaux de bord et connexion Power BI"),
        ("11.", "Utilisateurs, rôles et sécurité"),
        ("12.", "Guide de prise en main"),
        ("13.", "Déploiement sur GitHub et Render"),
        ("14.", "Interface de programmation (API)"),
        ("15.", "Exploitation, maintenance et évolutions"),
        ("A.", "Annexe 1 — Glossaire du suivi-évaluation"),
        ("B.", "Annexe 2 — Arborescence du code source"),
        ("C.", "Annexe 3 — Référentiels paramétrables"),
    ]
    for numero_chapitre, libelle in sommaire:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{numero_chapitre}\t{libelle}")
        run.font.size = Pt(10.5)
    document.add_page_break()

    # ------------------------------------------------- 1. Contexte
    titre1(document, "1. Contexte, justification et finalité de la plateforme")

    titre2(document, "1.1 Le constat")
    para(document,
         "La conduite d'un projet ou d'un programme de développement mobilise un ensemble "
         "d'instruments méthodologiques normalisés : cadre logique, cadre de rendement, cadre de "
         "suivi des indicateurs, registre des risques, chronogramme, plan de travail et budget "
         "annuel, manuel de suivi-évaluation, fiches de collecte, rapports périodiques et "
         "tableaux de bord. Dans la pratique courante des unités de gestion de projet, ces "
         "instruments sont produits et maintenus séparément, le plus souvent sous la forme de "
         "classeurs Excel et de documents Word autonomes.")
    para(document,
         "Cette dispersion engendre quatre difficultés récurrentes. Premièrement, la "
         "redondance de saisie : un même indicateur est ressaisi dans le cadre logique, dans le "
         "cadre de rendement, dans le tableau de suivi et dans chaque rapport, avec un risque "
         "élevé de divergence entre les versions. Deuxièmement, la lourdeur de la consolidation : "
         "la production d'un rapport trimestriel mobilise plusieurs jours-hommes de mise en forme "
         "avant même que ne commence le travail d'analyse. Troisièmement, la faiblesse de la "
         "traçabilité : il devient difficile d'établir qui a modifié quelle valeur, quand et sur "
         "la base de quelle source. Quatrièmement, le retard de l'information : lorsque le "
         "tableau de bord parvient au comité de pilotage, les données qu'il présente sont déjà "
         "anciennes de plusieurs semaines.")

    titre2(document, "1.2 La réponse apportée")
    para(document,
         f"La plateforme {APP_NAME} traite ces difficultés en instaurant une source unique de "
         "vérité. L'ensemble des données de planification et de suivi est saisi une seule fois, "
         "dans une base structurée ; tous les instruments méthodologiques en sont ensuite dérivés "
         "automatiquement. Modifier une cible dans la plateforme met simultanément à jour le "
         "cadre logique, le cadre de rendement, le tableau de bord, le rapport de performance et "
         "le flux Power BI.")
    encadre(document, "Principe directeur",
            "Un cadre logique et un budget suffisent à faire naître l'ensemble du dispositif de "
            "suivi-évaluation. Tout le reste — indicateurs paramétrés, cadres de mesure, "
            "questionnaires, chronogrammes, tableaux de bord, manuel de S&E — en est déduit et "
            "reste synchronisé en permanence.")

    titre2(document, "1.3 Finalités poursuivies")
    for texte in [
        "Réduire radicalement le temps consacré à la production documentaire, au profit du temps "
        "d'analyse et de dialogue de gestion.",
        "Garantir la cohérence méthodologique entre tous les instruments du dispositif, sur toute "
        "la durée du projet.",
        "Rendre l'information de performance disponible en continu, et non plus au rythme des "
        "échéances de rapportage.",
        "Sécuriser la mémoire institutionnelle du projet face à la rotation des équipes, en "
        "documentant chaque indicateur, chaque méthode de collecte et chaque décision.",
        "Faciliter la redevabilité envers les bénéficiaires, la tutelle et les partenaires "
        "financiers, par des livrables normalisés et immédiatement exploitables.",
    ]:
        puce(document, texte)

    titre2(document, "1.4 Utilisateurs visés")
    tableau(document, ["Profil", "Usage principal de la plateforme"],
            [["Responsable suivi-évaluation", "Paramétrage du cadre logique et des indicateurs, "
                                              "conception des outils de collecte, analyse et "
                                              "production des rapports"],
             ["Coordonnateur de projet", "Pilotage, revue des alertes, arbitrages, validation "
                                         "des rapports destinés au comité de pilotage"],
             ["Chef de composante", "Mise à jour de l'avancement des activités et des produits "
                                    "relevant de son périmètre"],
             ["Agent de terrain", "Collecte numérique et saisie des données de réalisation"],
             ["Responsable administratif et financier", "Suivi de l'exécution budgétaire, "
                                                        "engagements et décaissements"],
             ["Bailleur de fonds et tutelle", "Consultation des tableaux de bord et "
                                              "téléchargement des livrables"],
             ["Direction de programme", "Consolidation multi-projets au niveau du portefeuille"]],
            largeurs=[5, 11.5])

    document.add_page_break()

    # ------------------------------------------------- 2. Périmètre
    titre1(document, "2. Périmètre fonctionnel")
    para(document,
         "La plateforme couvre l'intégralité du cycle de gestion axée sur les résultats, depuis "
         "la formulation du cadre logique jusqu'à la production des rapports d'évaluation. Le "
         "tableau ci-après récapitule les quatorze modules fonctionnels.")
    tableau(document, ["Domaine", "Module", "Objet"],
            [["Pilotage", "Tableau de bord", "Indicateurs clés, indice de santé, équité, "
                                             "couverture territoriale, qualité, alertes priorisées"],
             ["Pilotage", "Portefeuille", "Vue consolidée multi-projets, comparaison, duplication"],
             ["Pilotage", "Fiche du projet", "Identification, ancrage institutionnel, théorie du "
                                             "changement, alignement stratégique"],
             ["Planification", "Cadre logique", "Chaîne de résultats hiérarchisée, sources de "
                                                "vérification, hypothèses"],
             ["Planification", "Indicateurs", "Fiches métadonnées complètes, cibles, règle "
                                              "d'agrégation, désagrégations exigées"],
             ["Planification", "Zones d'intervention", "Découpage géographique hiérarchisé, "
                                                       "population, cible de bénéficiaires, "
                                                       "coordonnées, responsable, et carte de "
                                                       "couverture à symboles proportionnels"],
             ["Planification", "Chronogramme et ordonnancement",
              "Gantt, chemin critique et durée du projet, réseau PERT, organigramme des tâches "
              "(WBS), matrice des responsabilités (RACI)"],
             ["Planification", "PTBA et budget", "Lignes budgétaires, ventilation trimestrielle, "
                                                 "engagements et décaissements"],
             ["Collecte et suivi", "Saisie des réalisations", "Saisie temps réel : période, zone, "
                                                              "activité source, ventilation par "
                                                              "sexe, âge et groupe cible"],
             ["Collecte et suivi", "Cadre de suivi (IPTT)", "Grille cibles/réalisations par "
                                                            "période, saisie directe"],
             ["Collecte et suivi", "Fiches et questionnaires", "Conception d'instruments, export "
                                                               "Word et XLSForm"],
             ["Analyse", "Équité et désagrégation", "Ventilation consolidée, indice d'équité de "
                                                    "genre, écart à la parité, désagrégations "
                                                    "manquantes"],
             ["Analyse", "Qualité des indicateurs", "Diagnostic SMART, score du système, actions "
                                                    "correctrices"],
             ["Analyse", "Risques et hypothèses", "Registre coté, matrice 5×5, atténuation, "
                                                  "contingence"],
             ["Rapportage", "Rapports périodiques", "Rapports trimestriels, semestriels et "
                                                    "annuels, avec aperçu à l'écran"],
             ["Rapportage", "Livrables", "Génération des vingt-deux documents du dispositif"],
             ["Rapportage", "Power BI", "Flux temps réel, modèle en étoile, faits désagrégés"],
             ["Données", "Importer", "Chargement depuis Excel, Word, XLSForm, exports Kobo"],
             ["Système", "Administration", "Comptes, rôles, journal d'audit"]],
            largeurs=[3.2, 4.3, 9])

    titre2(document, "2.1 Ce que la plateforme ne fait pas")
    para(document,
         "La délimitation du périmètre est aussi importante que son contenu. La plateforme n'est "
         "ni un logiciel de comptabilité, ni un outil de gestion des ressources humaines, ni un "
         "système d'information géographique. Elle enregistre les montants engagés et décaissés "
         "tels qu'ils lui sont communiqués par le service financier, mais ne tient pas la "
         "comptabilité générale du projet et ne produit pas d'états financiers réglementaires. "
         "Elle stocke des coordonnées GPS collectées sur le terrain, mais ne produit pas de "
         "cartographie thématique. Ces fonctions relèvent d'outils spécialisés, avec lesquels la "
         "plateforme s'articule par l'import et l'export de données.")

    document.add_page_break()

    # ------------------------------------------------- 3. Architecture
    titre1(document, "3. Architecture générale de la solution")

    titre2(document, "3.1 Vue d'ensemble")
    para(document,
         "La solution repose sur une architecture à trois couches, volontairement sobre afin de "
         "garantir un déploiement simple et une maintenance durable dans des contextes où les "
         "compétences informatiques disponibles sont limitées.")
    tableau(document, ["Couche", "Composants", "Technologies"],
            [["Présentation", "Interface web-mobile responsive : navigation, formulaires, "
                              "tableaux, graphiques SVG", "HTML5, CSS3, JavaScript natif "
                                                          "(aucune bibliothèque externe)"],
             ["Application", "API REST : authentification, opérations métier, moteur "
                             "analytique, générateurs de documents",
              "Python 3.12, FastAPI, Uvicorn"],
             ["Persistance", "Base de données relationnelle, vingt-quatre entités",
              "SQLAlchemy 2 ; PostgreSQL en production, SQLite en développement"]],
            largeurs=[3.2, 7.3, 6])

    encadre(document, "Choix structurant",
            "L'interface ne dépend d'aucune bibliothèque JavaScript externe et d'aucun CDN. Les "
            "graphiques — anneaux, barres, courbes, jauges, diagramme de Gantt, matrice des "
            "risques, réseau PERT, organigramme des tâches et carte de couverture — sont produits "
            "en SVG par un module interne. Il n'existe donc aucune étape de compilation : le "
            "déploiement se réduit à l'installation des dépendances Python. Ce choix supprime une "
            "classe entière de pannes (rupture de CDN, incompatibilité de version, vulnérabilité "
            "d'une dépendance transitive) et garantit le fonctionnement de l'application sur un "
            "réseau contraint. Le fond de carte OpenStreetMap constitue la seule ressource "
            "externe, facultative et désactivée automatiquement si elle est inaccessible.")

    titre2(document, "3.2 Organisation du code applicatif")
    tableau(document, ["Composant", "Responsabilité"],
            [["app/main.py", "Assemblage de l'application, montage des routeurs, service de "
                             "l'interface, sonde de disponibilité"],
             ["app/config.py", "Configuration par variables d'environnement et référentiels métier"],
             ["app/database.py", "Moteur et session SQLAlchemy, bascule SQLite / PostgreSQL"],
             ["app/models.py", "Modèle de données : vingt-quatre entités et leurs relations"],
             ["app/security.py", "Hachage PBKDF2-SHA256, jetons signés HMAC, contrôle d'accès "
                                 "hiérarchique par rôle"],
             ["app/crud.py", "Fabrique de routeurs CRUD génériques, sérialisation, coercition "
                             "des types, journal d'audit"],
             ["app/seed.py", "Compte administrateur initial et projet de démonstration"],
             ["app/routers/", "Points d'entrée : authentification, projets, entités métier, "
                              "imports, exports, Power BI"],
             ["app/services/analytics.py", "Moteur de calcul de la performance, agrégations, "
                                           "alertes, consolidation du portefeuille"],
             ["app/services/excel_export.py", "Neuf générateurs de classeurs Excel mis en forme"],
             ["app/services/word_export.py", "Sept générateurs de documents Word"],
             ["app/services/xlsform.py", "Génération de formulaires XLSForm pour KoboToolbox et ODK"],
             ["app/services/importer.py", "Analyseurs tolérants de classeurs Excel et de "
                                          "documents Word"],
             ["static/", "Interface : index.html, feuille de style, quatre modules JavaScript"]],
            largeurs=[5.5, 11])

    titre2(document, "3.3 Sécurité applicative")
    para(document,
         "L'authentification repose exclusivement sur la bibliothèque standard de Python. Les "
         "mots de passe sont hachés en PBKDF2-SHA256 avec 180 000 itérations et un sel aléatoire "
         "de 16 octets par compte. Les jetons de session sont des structures JSON signées en "
         "HMAC-SHA256, comportant l'identité, le rôle et une date d'expiration ; leur durée de "
         "validité est de douze heures par défaut. Ce choix évite les dépendances de "
         "cryptographie compilées, dont l'installation échoue fréquemment sur les plateformes "
         "d'hébergement à ressources restreintes.")

    document.add_page_break()

    # ------------------------------------------------- 4. Modèle de données
    titre1(document, "4. Modèle de données")
    para(document,
         "Le modèle comprend vingt-quatre entités. Sa structure reflète directement la logique du "
         "suivi-évaluation : un projet porte une chaîne de résultats et un découpage territorial ; "
         "chaque résultat porte des indicateurs ; chaque indicateur porte des cibles périodiques et "
         "des réalisations mesurées, elles-mêmes localisées dans une zone, rattachées à l'activité "
         "qui les a produites et ventilées par catégorie de bénéficiaires.")
    tableau(document, ["Entité", "Rôle", "Principaux attributs"],
            [["Project", "Projet ou programme",
              "code, intitulé, secteur, bailleur, agence d'exécution, dates, budget, devise, "
              "théorie du changement, alignement stratégique"],
             ["LogframeElement", "Maillon de la chaîne de résultats",
              "niveau (IMPACT / EFFET / PRODUIT / ACTIVITE), code, énoncé, parent, sources de "
              "vérification, hypothèses, responsable"],
             ["Indicator", "Fiche métadonnée d'indicateur",
              "code, libellé, définition, unité, formule, numérateur, dénominateur, "
              "désagrégations, référence, cible, sens, fréquence, source, méthode, responsable, "
              "coût, test SMART"],
             ["IndicatorTarget", "Cible périodique (jalon)",
              "période, année, dates de début et de fin, valeur cible"],
             ["IndicatorActual", "Réalisation mesurée",
              "période, date de référence, valeur, valeurs désagrégées par catégorie et modalité, "
              "zone d'intervention, activité source, source, agent collecteur, statut et auteur "
              "de la validation"],
             ["Zone", "Zone d'intervention",
              "code, nom, niveau administratif, zone parente, population, cible de bénéficiaires, "
              "latitude, longitude, responsable"],
             ["Stakeholder", "Partie prenante",
              "code, fonction ou structure, organisation, catégorie, contact — colonne de la "
              "matrice RACI"],
             ["RaciAssignment", "Affectation RACI",
              "activité, partie prenante, rôle (R, A, C ou I), commentaire"],
             ["Risk", "Risque du registre",
              "code, catégorie, énoncé, cause, conséquence, probabilité, impact, atténuation, "
              "contingence, risque résiduel, porteur, statut, date de revue"],
             ["Assumption", "Hypothèse critique",
              "code, niveau, énoncé, criticité, statut de validation, méthode de vérification, "
              "responsable"],
             ["Activity", "Activité du chronogramme",
              "code, libellé, résultat rattaché, responsable, partenaires, lieu, dates, "
              "avancement, statut, coûts, jalon, livrable"],
             ["BudgetLine", "Ligne du PTBA",
              "code, libellé, activité, catégorie, unité, quantité, coût unitaire, ventilation "
              "trimestrielle, source de financement, engagé, décaissé"],
             ["Beneficiary", "Groupe de bénéficiaires",
              "intitulé, typologie, zone, critères et méthode de sélection, cible et effectifs "
              "atteints ventilés, ménages, revenu et pauvreté de référence, besoins, contraintes, "
              "bénéfices attendus, mode de participation, vulnérabilité, mécanisme de plainte"],
             ["Partner", "Organisation partenaire",
              "intitulé, type, rôle, contribution financière conventionnée et décaissée, "
              "contribution technique, dates de convention, appréciation de la performance"],
             ["Evaluation", "Exercice évaluatif",
              "intitulé, type, période, équipe, statut, notes et justifications par critère du "
              "CAD, note globale calculée"],
             ["EvaluationRecommendation", "Recommandation d'évaluation",
              "énoncé, criticité, réponse du management, responsable, échéance, taux de mise en "
              "œuvre"],
             ["ImpactStudy", "Étude d'évaluation d'impact",
              "approche, méthode, hypothèse d'identification, règle d'affectation, indicateurs "
              "de résultat, tailles des groupes, grappes, corrélation intra-grappe, effet minimal "
              "détectable, écart-type, puissance, seuil, effet estimé, p-value"],
             ["ApiKey", "Clé d'accès en lecture seule",
              "intitulé, empreinte de la clé, projet, propriétaire, date d'expiration, dernière "
              "utilisation, révocation"],
             ["Form", "Instrument de collecte",
              "code, intitulé, type, population cible, périodicité, indicateurs alimentés, "
              "consignes, version, langue"],
             ["FormQuestion", "Question d'un instrument",
              "section, nom technique, libellé, type, modalités, obligation, contrainte, "
              "logique de saut, calcul, indicateur relié"],
             ["FormSubmission", "Réponse collectée", "date, agent, lieu, période, réponses"],
             ["User", "Compte utilisateur",
              "adresse électronique, nom, empreinte du mot de passe, rôle, organisation, état, "
              "confirmation d'adresse, tentatives infructueuses, verrouillage, date de "
              "changement du mot de passe, seuil de validité des jetons"],
             ["ProjectMember", "Affectation d'un utilisateur à un projet", "projet, utilisateur, rôle"],
             ["AuditLog", "Journal d'audit",
              "horodatage, utilisateur, action, entité, référence, projet, détail"]],
            largeurs=[3, 4, 9.5], taille=8.5)

    titre2(document, "4.1 Règles d'intégrité")
    for texte in [
        "Un élément du cadre logique référence son parent, ce qui permet une profondeur "
        "arbitraire et l'insertion de niveaux intermédiaires sans modification du modèle.",
        "Un indicateur est rattaché à un élément de la chaîne de résultats ; la suppression de "
        "cet élément conserve l'indicateur, qui devient orphelin et reste signalé comme tel.",
        "La suppression d'un projet entraîne la suppression explicite et ordonnée de toutes ses "
        "dépendances, y compris sous SQLite où les contraintes de cascade ne sont pas appliquées "
        "par défaut.",
        "Une réalisation est identifiée par le triplet indicateur / période / zone : deux zones "
        "peuvent déclarer une mesure sur la même période sans écrasement mutuel, et la "
        "consolidation s'opère au moment du calcul, non au moment de la saisie.",
        "Toute création, modification ou suppression est inscrite au journal d'audit avec "
        "l'identité de son auteur.",
        "Les colonnes ajoutées lors d'une mise à jour de l'application sont créées "
        "automatiquement au redémarrage, sans migration manuelle ni perte de données.",
    ]:
        puce(document, texte)

    document.add_page_break()

    # ------------------------------------------------- 5. Modules
    titre1(document, "5. Description détaillée des dix-neuf modules")

    modules = [
        ("5.1 Tableau de bord",
         "Vue d'entrée du projet. Elle présente six indicateurs clés — indice de santé, nombre "
         "d'indicateurs suivis, taux moyen de réalisation, avancement physique, exécution "
         "budgétaire, nombre de risques critiques — puis une jauge de santé décomposée en ses "
         "trois composantes pondérées, la répartition des indicateurs par statut de performance, "
         "la performance moyenne par niveau de résultat, la programmation budgétaire "
         "trimestrielle et l'exécution financière.",
         ["La liste d'alertes priorisées signale, dans l'ordre de gravité : les indicateurs en "
          "situation critique, les indicateurs clés non renseignés, les risques critiques ou "
          "élevés encore ouverts, les activités dont l'échéance est dépassée et les hypothèses "
          "invalidées.",
          "L'écart entre la performance réalisée et le pourcentage de temps écoulé fournit une "
          "lecture immédiate de l'avance ou du retard du projet.",
          "Le tableau de bord est imprimable en l'état et exportable en classeur Excel "
          "comportant graphiques et alertes."]),
        ("5.2 Portefeuille",
         "Consolidation de l'ensemble des projets. Chaque projet y est caractérisé par son "
         "indice de santé, son budget, son taux moyen de réalisation, son avancement physique, "
         "son taux d'exécution budgétaire et son nombre de risques critiques.",
         ["La création d'un projet ouvre un formulaire structuré en quatre sections : "
          "identification, ancrage institutionnel, cycle de vie et finances, cadrage stratégique.",
          "La duplication reprend la structure complète d'un projet — cadre logique, indicateurs, "
          "activités, budget, risques, hypothèses — sans les réalisations : elle sert à créer une "
          "phase 2 ou à décliner un projet type sur une nouvelle zone."]),
        ("5.3 Cadre logique",
         "Édition arborescente de la chaîne de résultats. Chaque maillon affiche son niveau, son "
         "code, son énoncé, ses sources de vérification, ses hypothèses et son responsable, ainsi "
         "que les indicateurs qui lui sont rattachés avec leur référence, leur cible et leur "
         "statut de performance.",
         ["Un bouton par élément permet d'ajouter un résultat de niveau inférieur, ce qui rend la "
          "construction de la chaîne à la fois guidée et rapide.",
          "Un second bouton crée directement un indicateur rattaché au résultat, en préremplissant "
          "son niveau.",
          "L'export produit la matrice en Excel (format A3 paysage, prêt à imprimer) et en Word "
          "(insérable dans un document de projet)."]),
        ("5.4 Indicateurs",
         "Gestion des fiches métadonnées. Le formulaire couvre vingt-quatre attributs répartis en "
         "cinq sections : identification, mesure, référence et cible, collecte, qualité.",
         ["La vue liste offre une recherche plein texte et un double filtre par niveau de "
          "résultat et par statut de performance.",
          "La fiche de suivi d'un indicateur présente la courbe des cibles et des réalisations, "
          "l'historique des mesures et les cibles périodiques, avec saisie et suppression directes.",
          "La génération automatique des cibles périodiques interpole linéairement entre la valeur "
          "de référence et la cible finale, selon une granularité trimestrielle, semestrielle ou "
          "annuelle."]),
        ("5.5 Cadre de suivi des indicateurs (IPTT)",
         "Grille de saisie croisant les indicateurs en lignes et les périodes en colonnes, chaque "
         "période affichant sa cible et sa réalisation.",
         ["Les cellules de réalisation sont directement modifiables ; toute cellule modifiée est "
          "signalée visuellement et l'ensemble des saisies est enregistré en une seule action.",
          "Ce mode de saisie reproduit l'ergonomie du tableur à laquelle les équipes de S&E sont "
          "habituées, tout en garantissant l'intégrité de la base.",
          "L'export Excel reprend la grille avec une mise en forme conditionnelle en trois "
          "couleurs sur la colonne de progression."]),
        ("5.6 Risques et hypothèses",
         "Registre des risques coté sur une échelle de 1 à 5 en probabilité et en impact, assorti "
         "d'une matrice de criticité 5×5 renseignée du nombre de risques par case, et suivi "
         "distinct des hypothèses critiques du cadre logique.",
         ["Chaque risque documente sa cause, sa conséquence sur les résultats, ses mesures "
          "d'atténuation, son plan de contingence, son porteur, son statut et sa date de revue.",
          "Le risque résiduel — probabilité et impact après application des mesures d'atténuation "
          "— est saisi séparément, ce qui permet de mesurer l'efficacité du traitement.",
          "Les hypothèses suivent un cycle de validation en quatre états : non vérifiée, "
          "partiellement vérifiée, vérifiée, invalidée. Une hypothèse invalidée déclenche une "
          "alerte, car elle remet en cause la logique d'intervention."]),
        ("5.7 Chronogramme et ordonnancement",
         "Module organisé en cinq onglets : diagramme de Gantt, chemin critique et réseau PERT, "
         "organigramme des tâches, matrice RACI et liste des activités. Un bandeau permanent "
         "affiche la durée calculée du projet, le nombre d'activités critiques, la marge moyenne "
         "et l'avancement physique.",
         ["Diagramme de Gantt mensuel construit dynamiquement : bleu pour les activités planifiées "
          "ou en cours, vert pour les activités achevées, rouge pour celles dont l'échéance est "
          "dépassée ; les jalons sont matérialisés par un losange et la date du jour par une ligne "
          "verticale. La part remplie de chaque barre représente l'avancement déclaré.",
          "Le chemin critique est matérialisé sur le diagramme de Gantt : les barres concernées "
          "sont cerclées de rouge et reliées par une courbe continue, tandis que les autres liens "
          "d'antécédence sont tracés en gris fin. La séquence déterminante se lit ainsi d'un coup "
          "d'œil, sans changer de vue.",
          "Le chemin critique identifie la séquence d'activités sans marge : tout retard sur l'une "
          "d'elles décale d'autant la date d'achèvement du projet. Son coût et son avancement "
          "moyen sont calculés.",
          "Une courbe en S complète l'analyse : elle confronte l'engagement de ressources programmé "
          "et l'engagement réalisé, mois par mois.",
          "Le diagramme de Gantt, le réseau PERT, l'organigramme des tâches, la courbe "
          "d'avancement et la carte de couverture s'exportent en image PNG haute définition ou en "
          "SVG vectoriel, directement insérables dans un rapport.",
          "Le réseau PERT présente chaque activité sous forme de nœud portant sa durée, ses dates "
          "au plus tôt et au plus tard et sa marge ; les activités d'un même rang peuvent être "
          "conduites en parallèle.",
          "L'organigramme des tâches décompose le projet en composantes, sous-composantes et lots "
          "de travail, avec codification automatique et consolidation ascendante des coûts.",
          "La matrice RACI, éditable cellule par cellule, attribue à chaque activité un rôle par "
          "partie prenante et contrôle la cohérence de l'ensemble.",
          "Une section récapitule les activités en retard, avec le nombre de jours de dépassement "
          "et le responsable concerné."]),
        ("5.8 PTBA et budget",
         "Saisie et suivi des lignes budgétaires. Chaque ligne combine une quantité, un coût "
         "unitaire et un nombre de répétitions, dont le produit constitue le montant planifié ; "
         "elle porte une ventilation sur les quatre trimestres, une source de financement, un "
         "montant engagé et un montant décaissé.",
         ["Les rattachements à une activité et, par elle, à un résultat du cadre logique "
          "permettent d'analyser le budget par produit et par effet.",
          "Les synthèses présentent la répartition par catégorie de dépense, la programmation "
          "trimestrielle et, lorsque le projet est pluriannuel, l'exécution par exercice.",
          "L'export produit un PTBA détaillé avec formules de totalisation actives et une feuille "
          "de synthèse graphique."]),
        ("5.9 Fiches et questionnaires",
         "Concepteur d'instruments de collecte. Chaque question est décrite par sa section, son "
         "nom technique, son libellé, son type parmi douze, ses modalités de réponse, son "
         "caractère obligatoire, sa contrainte de saisie, sa condition d'affichage, sa formule de "
         "calcul et l'indicateur qu'elle alimente.",
         ["Les contraintes et les logiques de saut sont exprimées dans la syntaxe XLSForm, ce qui "
          "garantit leur transposition fidèle dans KoboToolbox et ODK Collect.",
          "Chaque instrument s'exporte simultanément en questionnaire Word — mis en page pour "
          "l'administration papier, avec cases à cocher et zones de réponse — et en XLSForm "
          "téléversable tel quel sur un serveur de collecte.",
          "Le lien entre une question et un indicateur permet la réinjection automatique des "
          "données collectées."]),
        ("5.10 Saisie des réalisations",
         "Écran de renseignement des indicateurs conçu pour l'usage quotidien, y compris sur "
         "téléphone. La liste des indicateurs affiche pour chacun la dernière période mesurée, la "
         "valeur consolidée, le nombre de mesures agrégées et le statut de performance ; des "
         "filtres isolent les indicateurs clés, ceux qui ne sont pas encore renseignés ou ceux "
         "dont le statut appelle une actualisation.",
         ["Le formulaire de saisie porte la période, la date de référence, la zone d'intervention, "
          "l'activité qui a produit la donnée, la source et le statut de validation.",
          "Lorsque l'indicateur exige une désagrégation, une grille de ventilation apparaît : une "
          "ligne par catégorie, une case par modalité. Le total de la première catégorie renseignée "
          "alimente automatiquement la valeur globale, ce qui supprime la double saisie et "
          "garantit la cohérence entre le total et sa ventilation.",
          "Une mesure est identifiée par le triplet indicateur / période / zone : deux zones "
          "peuvent donc déclarer une réalisation sur la même période sans écrasement.",
          "Le taux d'atteinte de la période est recalculé et affiché immédiatement après "
          "l'enregistrement."]),
        ("5.11 Équité et désagrégation",
         "Analyse de l'inclusivité effective des interventions. Le module consolide toutes les "
         "ventilations saisies, calcule l'indice d'équité de genre au niveau du projet, de chaque "
         "indicateur et de chaque zone, et signale les indicateurs pour lesquels une désagrégation "
         "est exigée mais n'a pas été renseignée.",
         ["Le taux de désagrégation — part des indicateurs effectivement ventilés parmi ceux qui "
          "l'exigent — constitue un indicateur de qualité du dispositif lui-même.",
          "Pour chaque catégorie, la répartition est présentée en anneau et en tableau, avec la "
          "liste des modalités du référentiel qui n'ont pas été renseignées.",
          "L'écart à la parité est exprimé en points ; la parité est considérée atteinte en deçà "
          "de cinq points d'écart."]),
        ("5.12 Zones d'intervention et carte de couverture",
         "Gestion du découpage géographique et consolidation territoriale. Chaque zone porte son "
         "niveau administratif, sa zone parente, sa population, sa cible de bénéficiaires, ses "
         "coordonnées et son responsable. Une carte de couverture restitue visuellement "
         "l'implantation du projet.",
         ["La carte est à symboles proportionnels : la surface de chaque cercle — et non son "
          "rayon — représente les bénéficiaires atteints, sa couleur le taux de couverture de la "
          "cible de la zone, et un demi-disque la part des femmes. Des liens en pointillé relient "
          "chaque zone à sa zone mère.",
          "La projection est celle de Mercator sphérique (EPSG:3857). Le rendu comprend un "
          "graticule gradué en degrés, une échelle métrique et une rose des vents ; il est produit "
          "en SVG, sans bibliothèque cartographique.",
          "Un fond de carte OpenStreetMap peut être superposé par une simple case à cocher. C'est "
          "le seul appel réseau externe de la plateforme ; s'il échoue, il est automatiquement "
          "désactivé et la carte demeure exploitable.",
          "Le taux de couverture rapporte les bénéficiaires atteints à la cible de la zone : il "
          "révèle les déséquilibres géographiques que les totaux nationaux masquent.",
          "La part des femmes est calculée zone par zone, ce qui permet d'identifier les "
          "territoires où l'inclusion des femmes est en retrait.",
          "Un tableau distinct présente les données collectées par activité : il relie la collecte "
          "de données à la mise en œuvre et permet de vérifier qu'une activité déclarée avancée "
          "produit effectivement des réalisations mesurées."]),
        ("5.13 Qualité des indicateurs",
         "Diagnostic SMART du système de mesure. Chaque indicateur est évalué sur cinq critères, "
         "d'abord par un contrôle automatique fondé sur les données réellement saisies, ensuite "
         "par une revue manuelle qui prévaut sur ce contrôle.",
         ["Le score du système est la moyenne des scores individuels ; il est décomposé par "
          "critère, ce qui indique la faiblesse dominante du dispositif.",
          "Chaque critère non satisfait produit une action correctrice nommée et actionnable "
          "(« rédiger la définition opérationnelle », « fixer une échéance de cible »…).",
          "La revue est datée et commentée, ce qui documente la démarche d'amélioration continue "
          "attendue lors des audits de qualité des données."]),
        ("5.14 Rapports périodiques",
         "Production des rapports trimestriels, semestriels et annuels. La période de rapportage "
         "est choisie dans une liste combinant les périodes déjà renseignées et celles déduites du "
         "calendrier du projet.",
         ["Un aperçu s'affiche avant génération : indicateurs de la période, taux moyen "
          "d'atteinte, bénéficiaires ventilés, exécution financière, écarts à traiter et "
          "consolidation par zone.",
          "Une période annuelle englobe automatiquement ses semestres et ses trimestres ; une "
          "période semestrielle englobe ses deux trimestres.",
          "Le rapport produit comporte huit parties, dont un bloc de validation à signer par le "
          "responsable de suivi-évaluation, le coordonnateur et le comité de pilotage."]),
        ("5.15 Bénéficiaires",
         "Caractérisation des groupes visés par l'intervention, sur deux registres. Le registre "
         "quantitatif porte la cible et les effectifs atteints — ventilés femmes, jeunes et "
         "personnes handicapées —, le nombre de ménages et leur taille moyenne, le revenu annuel "
         "et le taux de pauvreté de référence. Le registre qualitatif porte les critères et la "
         "méthode de sélection, les besoins exprimés, les contraintes, les bénéfices attendus, le "
         "mode de participation, le degré de vulnérabilité et le mécanisme de gestion des "
         "plaintes ouvert au groupe.",
         ["Le taux d'atteinte et la part des femmes atteintes sont calculés à partir des "
          "effectifs saisis : ils ne peuvent être renseignés à la main, donc ni arrondis ni "
          "arrangés.",
          "Un indicateur peut être rattaché à un groupe de bénéficiaires. Le rattachement relie "
          "la population visée à la mesure du changement : la fiche du groupe affiche alors les "
          "indicateurs qui le documentent, leur taux d'atteinte et leur part de femmes.",
          "La synthèse signale les groupes ciblés qu'aucun indicateur ne mesure — angle mort "
          "fréquent d'un dispositif de suivi, où une population figure au document de projet "
          "sans qu'aucune donnée n'en rende compte."]),
        ("5.16 Partenaires",
         "Registre des organisations parties au montage : type, rôle dans le dispositif, "
         "contribution financière conventionnée et effectivement décaissée, contribution "
         "technique, échéances de convention et appréciation de la performance.",
         ["Le taux de décaissement est calculé par partenaire et pour l'ensemble du montage : "
          "il révèle les engagements annoncés mais non honorés, première cause de décalage entre "
          "le budget planifié et le budget disponible.",
          "Les conventions arrivant à échéance sont signalées, afin que leur renouvellement soit "
          "engagé avant l'interruption des financements."]),
        ("5.17 Évaluation selon les critères du CAD de l'OCDE",
         "Conduite des exercices évaluatifs — évaluation de référence, à mi-parcours, finale, "
         "ex-post — notés sur les six critères du Comité d'aide au développement : pertinence, "
         "cohérence, efficacité, efficience, impact et durabilité.",
         ["La notation suit une échelle à six niveaux qui écarte délibérément la note médiane "
          "neutre : l'évaluateur doit se prononcer du côté satisfaisant ou insatisfaisant.",
          "Chaque critère s'accompagne de ses points d'examen — les questions auxquelles la note "
          "doit répondre — et d'une justification écrite obligatoire : une note sans "
          "justification est refusée.",
          "La synthèse produit la moyenne par critère sur l'ensemble des évaluations achevées, "
          "ce qui met en évidence la dimension sur laquelle le projet est systématiquement le "
          "plus faible.",
          "Les recommandations portent leur criticité, la réponse du management — acceptée, "
          "partiellement acceptée, rejetée —, le responsable, l'échéance et le taux de mise en "
          "œuvre ; celles qui sont échues et non soldées sont signalées."]),
        ("5.18 Évaluation d'impact",
         "Conception et suivi des études visant à établir l'effet propre du projet, c'est-à-dire "
         "la part du changement qui lui est imputable et non aux évolutions du contexte. Neuf "
         "méthodes sont documentées, chacune avec son hypothèse d'identification, ses conditions "
         "d'application, ses forces et ses limites.",
         ["Les méthodes expérimentales — essai randomisé contrôlé, randomisation par grappes — "
          "reposent sur l'affectation aléatoire, qui rend les groupes comparables en espérance "
          "sur toutes les caractéristiques, observées comme non observées.",
          "Les méthodes quasi-expérimentales — doubles différences, appariement par score de "
          "propension, combinaison des deux, régression sur discontinuité, variables "
          "instrumentales, contrôle synthétique — s'appliquent lorsque la randomisation est "
          "impossible, au prix d'une hypothèse d'identification explicite qui doit être discutée.",
          "La comparaison avant-après est proposée mais présentée pour ce qu'elle est : une "
          "description, non une mesure d'impact, puisqu'elle confond l'effet du projet avec tout "
          "ce qui a changé par ailleurs.",
          "Un calculateur de taille d'échantillon donne l'effectif requis par bras à partir de "
          "l'effet minimal détectable, de l'écart-type de l'indicateur, de la puissance et du "
          "seuil de signification. Lorsque la randomisation porte sur des grappes, il applique "
          "l'effet de plan 1 + (m − 1) × ρ : avec des grappes de 30 unités et une corrélation "
          "intra-grappe de 0,08, l'échantillon requis est multiplié par 3,3.",
          "Le contrôle de puissance compare l'échantillon prévu à l'échantillon requis et "
          "signale une étude sous-dimensionnée, c'est-à-dire une étude qui risque de conclure à "
          "l'absence d'effet alors qu'un effet réel existe."]),
        ("5.19 Administration",
         "Gestion des comptes, des rôles, des accès par projet, des clés d'API et consultation "
         "du journal d'audit.",
         ["Cinq rôles hiérarchisés déterminent les droits : lecteur, opérateur de saisie, "
          "responsable de suivi-évaluation, coordonnateur, administrateur. Le rôle est doublé "
          "d'un rattachement projet par projet : un utilisateur ne voit que les projets dont il "
          "est membre.",
          "L'ouverture d'un compte engendre un mot de passe provisoire et un lien de "
          "confirmation d'adresse, montrés une seule fois à l'administrateur qui le crée. La "
          "connexion reste refusée tant que le lien n'a pas été ouvert.",
          "Les clés d'API sont nominatives, limitées à un projet, en lecture seule, datées et "
          "révocables ; seule leur empreinte est conservée.",
          "Le journal d'audit conserve l'horodatage, l'auteur, l'action, l'entité concernée et "
          "le projet, pour toutes les opérations d'écriture."]),
    ]
    for titre_module, chapeau, points in modules:
        titre2(document, titre_module)
        para(document, chapeau)
        for point in points:
            puce(document, point)

    document.add_page_break()

    # ------------------------------------------------- 6. Méthodologie
    titre1(document, "6. Méthodologie de suivi-évaluation intégrée")
    para(document,
         "Les règles de calcul appliquées par la plateforme ne sont pas des conventions "
         "arbitraires : elles traduisent les pratiques établies de la gestion axée sur les "
         "résultats. Leur explicitation est nécessaire à l'appropriation de l'outil et à la "
         "défense des chiffres produits devant les instances de gouvernance.")

    titre2(document, "6.1 Les deux taux de réalisation")
    para(document,
         "La plateforme calcule deux taux distincts pour chaque indicateur, car ils répondent à "
         "deux questions différentes.")
    tableau(document, ["Taux", "Formule", "Question à laquelle il répond"],
            [["Taux de la période", "réalisé ÷ cible de la même période × 100",
              "Le projet a-t-il tenu son engagement sur la période écoulée ?"],
             ["Progression vers la cible finale",
              "(réalisé − référence) ÷ (cible − référence) × 100",
              "Quelle part du chemin vers la cible de fin de projet a été parcourue ?"]],
            largeurs=[4, 6, 6.5], taille=9.5)
    para(document,
         "Le statut de performance est déterminé par le taux de la période lorsqu'une cible "
         "périodique existe, et par la progression finale à défaut. Ce choix est essentiel : "
         "comparer une réalisation de première année à une cible de fin de projet conduirait à "
         "classer en situation critique un projet parfaitement conforme à sa programmation.")
    encadre(document, "Illustration",
            "Un projet vise 15 000 producteurs formés à l'horizon 2029, avec un jalon de 3 400 "
            "au troisième trimestre 2025. La réalisation au troisième trimestre s'établit à "
            "3 260. Le taux de la période est de 95,9 % — le projet est en bonne voie. La "
            "progression vers la cible finale n'est que de 21,7 %, ce qui est normal en début "
            "d'exécution. Juger l'indicateur sur ce second chiffre reviendrait à le déclarer "
            "critique à tort.")

    titre2(document, "6.2 Indicateurs à progression décroissante")
    para(document,
         "Certains indicateurs — incidence de la pauvreté, prévalence de l'insécurité "
         "alimentaire, taux de pertes post-récolte — s'améliorent lorsqu'ils diminuent. La "
         "plateforme inverse alors le rapport : le taux de la période devient cible ÷ réalisé, "
         "de sorte qu'une valeur inférieure à la cible produit un taux supérieur à 100 %. La "
         "progression vers la cible finale, fondée sur l'écart à la référence, reste valide sans "
         "modification.")

    titre2(document, "6.3 Statuts de performance")
    tableau(document, ["Statut", "Seuil", "Interprétation opérationnelle"],
            [["Atteint", "≥ 100 %", "L'engagement de la période est tenu ou dépassé"],
             ["En bonne voie", "85 % à 99,9 %", "Écart mineur ne nécessitant pas de mesure "
                                               "corrective immédiate"],
             ["À surveiller", "60 % à 84,9 %", "Écart significatif appelant une analyse causale "
                                               "à la prochaine revue"],
             ["Critique", "< 60 %", "Écart majeur exigeant une mesure corrective documentée"],
             ["Non renseigné", "—", "Absence de mesure : l'indicateur ne peut être apprécié"]],
            largeurs=[3.5, 3, 10], taille=9.5)

    titre2(document, "6.4 Indice de santé du projet")
    para(document,
         "L'indice de santé synthétise la situation du projet en une valeur unique, moyenne "
         "pondérée de trois composantes.")
    tableau(document, ["Composante", "Pondération", "Justification de la pondération"],
            [["Résultats — taux moyen de réalisation des indicateurs", "45 %",
              "La finalité d'un projet est l'atteinte de ses résultats, non la consommation de "
              "ses moyens"],
             ["Exécution physique — avancement moyen des activités", "30 %",
              "Traduit la capacité opérationnelle effective de mise en œuvre"],
             ["Exécution financière — taux de décaissement", "25 %",
              "Contrainte de gestion réelle, mais qui ne préjuge pas de la qualité des résultats"]],
            largeurs=[6.5, 2.5, 7.5], taille=9.5)
    para(document,
         "L'indice est systématiquement rapproché du pourcentage de temps écoulé. Un indice de "
         "40 % à 32 % de la durée du projet traduit une avance ; le même indice à 70 % de la "
         "durée traduit un retard sérieux. C'est cet écart, et non l'indice pris isolément, qui "
         "constitue le signal de gestion.")

    titre2(document, "6.5 Cotation des risques")
    para(document,
         "Chaque risque est coté de 1 à 5 en probabilité et de 1 à 5 en impact. Le produit des "
         "deux notes donne un score de criticité compris entre 1 et 25, qui détermine le niveau "
         "de priorité et la couleur affichée dans la matrice.")
    tableau(document, ["Niveau", "Score", "Conduite à tenir"],
            [["Critique", "15 à 25", "Traitement prioritaire, mesures d'atténuation immédiates, "
                                     "revue mensuelle, information du comité de pilotage"],
             ["Élevé", "10 à 14", "Mesures d'atténuation planifiées et budgétées, revue "
                                  "trimestrielle"],
             ["Modéré", "5 à 9", "Surveillance active, revue semestrielle"],
             ["Faible", "1 à 4", "Acceptation, revue annuelle"]],
            largeurs=[3, 2.5, 11], taille=9.5)
    para(document,
         "Le risque résiduel, coté après application des mesures d'atténuation, permet de vérifier "
         "que le traitement produit l'effet attendu. Un risque dont le score résiduel demeure "
         "critique appelle une réorientation stratégique, non une simple mesure opérationnelle.")

    titre2(document, "6.6 Règle d'agrégation des mesures")
    para(document,
         "Un indicateur suivi sur plusieurs zones d'intervention produit plusieurs mesures pour une "
         "même période. La consolidation ne peut être uniforme : additionner des effectifs de "
         "bénéficiaires est correct, additionner des rendements moyens ne l'est pas. Chaque "
         "indicateur porte donc une règle d'agrégation explicite, déduite de son unité de mesure à "
         "défaut de choix de l'utilisateur.")
    tableau(document, ["Règle", "Application", "Exemples d'indicateurs"],
            [["Somme", "Grandeurs cumulables sur le territoire",
              "Nombre de producteurs formés, tonnes de semences distribuées, hectares aménagés"],
             ["Moyenne", "Taux, ratios, scores et rendements",
              "Rendement moyen du maïs (t/ha), incidence de la pauvreté (%), score de diversité "
              "alimentaire"],
             ["Dernière valeur", "États et stocks appréciés à une date",
              "Nombre d'infrastructures fonctionnelles à la date de la revue"],
             ["Maximum", "Couvertures ne devant pas être cumulées",
              "Portée maximale atteinte sur une zone"]],
            largeurs=[3, 5, 8.5], taille=9.5)
    encadre(document, "Erreur évitée",
            "Sans règle d'agrégation, un indicateur d'effectifs mesuré sur six préfectures "
            "afficherait la valeur de la dernière zone saisie au lieu du total national, ce qui "
            "diviserait mécaniquement la performance affichée par six. Symétriquement, un "
            "rendement moyen mesuré sur six zones afficherait la somme des six rendements, soit "
            "une valeur six fois trop élevée.")

    titre2(document, "6.7 Indice d'équité de genre")
    para(document,
         "À partir de la ventilation par sexe, la plateforme calcule la part des femmes parmi les "
         "bénéficiaires, l'écart à la parité exprimé en points de pourcentage et une appréciation "
         "qualitative. La parité est considérée atteinte lorsque l'écart est inférieur à cinq "
         "points ; en deçà de 45 % de femmes, une sous-représentation est signalée.")
    para(document,
         "Cet indice est produit à quatre niveaux : le projet dans son ensemble, chaque indicateur, "
         "chaque zone d'intervention et chaque activité. Cette granularité est déterminante : un "
         "projet peut afficher une parité satisfaisante au niveau national tout en présentant, sur "
         "certaines zones ou certaines activités, une exclusion marquée des femmes. Seule la "
         "ventilation territoriale rend ce phénomène visible.")

    titre2(document, "6.8 Score de qualité SMART")
    para(document,
         "Chaque indicateur est noté sur cent points, à raison de vingt points par critère "
         "satisfait. Le contrôle est d'abord automatique — il s'appuie sur les informations "
         "effectivement saisies dans la fiche — puis peut être corrigé par une revue manuelle "
         "documentée, qui prévaut sur le contrôle automatique.")
    tableau(document, ["Critère", "Contrôle automatique appliqué"],
            [["Spécifique", "Le libellé et la définition opérationnelle sont renseignés"],
             ["Mesurable", "L'unité de mesure et le mode de calcul ou la méthode de collecte sont "
                           "renseignés"],
             ["Atteignable", "Une valeur de référence et une cible finale distinctes sont "
                             "renseignées"],
             ["Pertinent", "L'indicateur est rattaché à un résultat du cadre logique"],
             ["Temporellement défini", "Une échéance de cible et une fréquence de collecte sont "
                                       "renseignées"]],
            largeurs=[4.5, 12], taille=9.5)
    para(document,
         "Le score du système est la moyenne des scores individuels : excellente au-delà de 90 %, "
         "bonne au-delà de 75 %, acceptable au-delà de 60 %, insuffisante en deçà. Chaque critère "
         "non satisfait produit une action correctrice nommée, ce qui transforme le diagnostic en "
         "plan de travail.")

    titre2(document, "6.9 Indicateurs de résultat et indicateurs de processus")
    para(document,
         "Un indicateur porte une nature : il mesure soit un changement produit chez les "
         "bénéficiaires — c'est un indicateur de résultat —, soit la conduite de l'action "
         "elle-même — c'est un indicateur de processus. Les seconds documentent le taux "
         "d'exécution du plan de travail, les délais de production des rapports ou de passation "
         "des marchés, le taux de participation aux activités programmées ou la complétude des "
         "données transmises.")
    para(document,
         "Cette distinction est méthodologiquement importante : mêler indicateurs de résultat et "
         "de processus dans un même tableau de bord conduit à surestimer la performance, car les "
         "indicateurs de processus sont structurellement plus faciles à atteindre. Leur affichage "
         "est donc commandé par une option activable projet par projet. Lorsqu'elle est "
         "désactivée, les indicateurs de processus restent enregistrés — avec leurs cibles et "
         "leurs mesures — mais sont exclus des tableaux de bord, des analyses et des livrables. "
         "Le nombre d'indicateurs masqués demeure affiché, afin que l'option ne dissimule jamais "
         "l'existence des données.")
    tableau(document, ["Nature", "Ce qu'elle mesure", "Exemples"],
            [["Résultat", "Le changement produit chez les bénéficiaires",
              "Rendement moyen du maïs, revenu agricole par exploitation, incidence de la "
              "pauvreté, nombre de producteurs formés"],
             ["Processus", "La conduite de l'action et la qualité de la gestion",
              "Taux d'exécution du PTBA, délai de production des rapports trimestriels, taux de "
              "participation aux formations, délai de passation des marchés, complétude des "
              "données de suivi"]],
            largeurs=[3, 5.5, 8], taille=9.5)

    titre2(document, "6.10 Ordonnancement : chemin critique et réseau PERT")
    para(document,
         "Les activités du chronogramme portent des antécédents, exprimés en relations fin-début : "
         "une activité ne peut démarrer qu'une fois ses antécédents achevés. Leur durée est soit "
         "imposée, soit déduite des dates de début et de fin.")
    para(document,
         "Le calcul procède en trois temps. Un tri topologique ordonne les activités et détecte "
         "les circuits de dépendances. Une passe avant établit, pour chaque activité, la date de "
         "début et de fin au plus tôt. Une passe arrière établit les dates au plus tard "
         "compatibles avec l'achèvement du projet à la date calculée.")
    tableau(document, ["Notion", "Définition", "Usage en gestion de projet"],
            [["Durée totale du projet", "Date de fin au plus tôt la plus tardive du réseau",
              "Comparée à la date de clôture planifiée ; l'écart est signalé"],
             ["Marge totale", "Retard admissible sans décaler la fin du projet",
              "Une marge nulle rend l'activité critique"],
             ["Marge libre", "Retard admissible sans décaler l'activité suivante",
              "Indique la souplesse réelle dont dispose le responsable de l'activité"],
             ["Chemin critique", "Séquence continue d'activités à marge nulle",
              "Concentre l'attention de la supervision : tout retard s'y répercute intégralement"],
             ["Rang PERT", "Position de l'activité dans le réseau",
              "Les activités d'un même rang sont indépendantes et parallélisables"]],
            largeurs=[3.5, 6, 7], taille=9.5)
    encadre(document, "Contrôle de cohérence",
            "La plateforme signale les antécédents qui s'achèvent après le début planifié de leur "
            "successeur : le calendrier saisi et le lien d'antécédence sont alors contradictoires. "
            "C'est l'incohérence la plus fréquente des chronogrammes construits sous tableur, où "
            "les dates sont saisies indépendamment des liens logiques entre activités.")

    titre2(document, "6.11 Organigramme des tâches (WBS)")
    para(document,
         "L'organigramme des tâches décompose le projet en éléments de plus en plus fins, "
         "jusqu'aux lots de travail élémentaires. La plateforme le déduit de la chaîne de "
         "résultats déjà saisie : les effets constituent les composantes, les produits les "
         "sous-composantes, les activités les lots de travail. Cette dérivation garantit la "
         "cohérence entre le cadre logique et l'organisation opérationnelle — deux instruments "
         "trop souvent construits séparément et divergents.")
    for texte in [
        "La codification est automatique et hiérarchique : 1, 1.1, 1.1.1, 1.1.1.1. Elle peut être "
        "inscrite sur les activités pour être reprise dans les autres livrables.",
        "Les coûts, durées et avancements sont consolidés de bas en haut : chaque niveau totalise "
        "ses descendants, ce qui permet de vérifier la cohérence du budget par composante.",
        "Les activités non rattachées à un produit sont regroupées dans un lot « Gestion, "
        "coordination et suivi-évaluation », conformément à la pratique courante.",
        "Un dictionnaire des lots de travail accompagne l'organigramme : il précise pour chaque "
        "lot son livrable attendu, son responsable, sa durée, son coût et son échéance.",
    ]:
        puce(document, texte)

    titre2(document, "6.12 Matrice des responsabilités (RACI)")
    para(document,
         "La matrice RACI croise les activités et les parties prenantes et attribue à chaque "
         "intersection un rôle. Elle répond à la question qui, en gestion de projet, produit le "
         "plus de blocages : qui décide, qui exécute, qui doit être consulté et qui doit être "
         "informé.")
    tableau(document, ["Rôle", "Signification", "Règle appliquée par la plateforme"],
            [["R — Responsible", "Réalise le travail",
              "Plusieurs réalisateurs sont possibles ; au moins un est exigé par activité"],
             ["A — Accountable", "Approuve et rend compte",
              "Un seul approbateur par activité : la responsabilité ne se partage pas"],
             ["C — Consulted", "Est consulté avant la décision",
              "La consultation est bilatérale et intervient en amont"],
             ["I — Informed", "Est informé après la décision",
              "L'information est unilatérale et intervient en aval"]],
            largeurs=[3.5, 4.5, 8.5], taille=9.5)
    para(document,
         "Deux contrôles automatiques sont appliqués : l'absence d'approbateur laisse une activité "
         "sans responsable identifié, la présence de plusieurs approbateurs dilue la "
         "responsabilité. La charge de chaque acteur est en outre calculée : une partie prenante "
         "qui approuve un nombre disproportionné d'activités constitue un goulot d'étranglement "
         "décisionnel, signalé comme tel.")

    titre2(document, "6.13 Courbe d'avancement (courbe en S)")
    para(document,
         "La courbe en S représente l'engagement cumulé des ressources, mois par mois. Le coût de "
         "chaque activité est réparti linéairement sur sa durée puis cumulé : la courbe programmée "
         "court sur toute la durée du projet et atteint 100 %, tandis que la courbe réalisée "
         "s'arrête au mois en cours, l'avancement déclaré déterminant la part effectivement "
         "consommée.")
    para(document,
         "Sa forme en S est caractéristique de la conduite de projet : montée lente au démarrage, "
         "accélération en phase de croisière, ralentissement à l'approche de la clôture. L'écart "
         "vertical entre les deux courbes à la date du jour mesure l'avance ou le retard "
         "d'exécution, exprimé en points de pourcentage.")

    titre2(document, "6.14 Portabilité et intégrité des données")
    para(document,
         "La sauvegarde JSON produit un fichier autonome contenant l'intégralité d'un projet ou "
         "d'un portefeuille. À l'import, les identifiants sont réattribués par la base d'accueil et "
         "toutes les références internes sont réécrites — parent d'un résultat, zone d'une mesure, "
         "activité d'une ligne budgétaire, acteur d'une affectation RACI. L'échange est donc "
         "indépendant des séquences de clés primaires et fonctionne entre instances distinctes.")
    tableau(document, ["Format", "Contenu", "Réversibilité"],
            [["JSON projet", "Intégralité d'un projet, questionnaires et questions compris",
              "Restitution à l'identique"],
             ["JSON portefeuille", "Tous les projets de l'instance dans un fichier unique",
              "Restitution à l'identique"],
             ["Excel de transfert", "Toutes les données dans la structure du modèle d'import",
              "Réversible, hors questionnaires"]],
            largeurs=[4, 8.5, 4], taille=9.5)
    encadre(document, "Gestion des collisions de code",
            "Si le code du projet importé est déjà utilisé, il est suffixé automatiquement et "
            "l'utilisateur en est averti dans le rapport d'import. L'option « remplacer un projet "
            "existant » produit l'effet inverse : le projet de même code est supprimé au préalable, "
            "ce qui permet de restaurer une sauvegarde par-dessus une version altérée.")

    titre2(document, "6.15 Articulation entre risques et hypothèses")
    para(document,
         "Le cadre logique distingue les hypothèses — conditions externes nécessaires à la "
         "réalisation de la chaîne de résultats — des risques, qui en sont la formulation "
         "négative. La plateforme maintient les deux registres et les relie : une hypothèse dont "
         "le statut passe à « invalidée » signale que la logique d'intervention repose désormais "
         "sur une prémisse fausse, ce qui constitue l'alerte la plus grave que puisse produire un "
         "dispositif de suivi.")

    document.add_page_break()

    # ------------------------------------------------- 7. Imports
    titre1(document, "7. Import de données : Excel, Word, XLSForm, KoboToolbox")

    titre2(document, "7.1 Import depuis Excel")
    para(document,
         "Le module d'import accepte un classeur comportant tout ou partie de huit onglets : "
         "Cadre logique, Indicateurs, Cibles, Réalisations, Activités, Budget, Risques, "
         "Hypothèses. L'appariement des onglets et des colonnes est volontairement tolérant : les "
         "intitulés sont normalisés — passage en minuscules, suppression des accents et de la "
         "ponctuation — puis rapprochés d'une table d'équivalences. Une colonne intitulée "
         "« Situation de référence », « Valeur de référence » ou « Baseline » est reconnue dans "
         "les trois cas.")
    for texte in [
        "Un modèle prérempli, commenté et assorti d'exemples est téléchargeable depuis "
        "l'application. Il constitue le point de départ recommandé.",
        "La hiérarchie du cadre logique est reconstituée à partir d'une colonne « Code parent » : "
        "le produit P1.1 déclare l'effet OS1 comme parent, l'activité A1.1.1 déclare le produit "
        "P1.1. Les rattachements introuvables sont signalés dans le rapport d'import.",
        "Un enregistrement dont le code existe déjà est mis à jour ; les autres sont créés. "
        "L'import est donc idempotent et peut être relancé après correction.",
        "L'option « Remplacer les données existantes » vide préalablement le projet, pour un "
        "rechargement complet.",
        "Le rapport d'import détaille les onglets traités, le nombre d'enregistrements créés par "
        "catégorie et la liste des avertissements.",
    ]:
        puce(document, texte)

    titre2(document, "7.2 Import depuis Word")
    para(document,
         "De nombreux cadres logiques n'existent que sous la forme d'un tableau inséré dans le "
         "document de projet. La plateforme analyse le document, inspecte chaque tableau, "
         "détermine sa nature probable — cadre logique, registre des risques, liste "
         "d'indicateurs, plan d'activités — et présente cet inventaire à l'utilisateur, qui "
         "choisit le tableau à importer.")
    for texte in [
        "Deux structures sont reconnues : la matrice classique à quatre colonnes (logique "
        "d'intervention, indicateurs objectivement vérifiables, sources de vérification, "
        "hypothèses) et le tableau structuré comportant des colonnes explicites de niveau et de "
        "code.",
        "Les niveaux sont déduits des libellés au moyen d'un dictionnaire de correspondances : "
        "« objectif global », « but » et « finalité » désignent l'impact ; « objectif "
        "spécifique », « outcome » et « résultat » désignent l'effet ; « extrant » et « output » "
        "désignent le produit.",
        "Le rattachement hiérarchique est reconstitué par la position : chaque élément est "
        "rattaché au dernier élément rencontré de niveau immédiatement supérieur.",
        "Les indicateurs contenus dans la cellule des IOV sont extraits ligne à ligne, avec "
        "détection du code lorsqu'il précède le libellé.",
    ]:
        puce(document, texte)
    encadre(document, "Recommandation",
            "L'import Word constitue un accélérateur de reprise, non un substitut à la relecture. "
            "Après import, il convient de vérifier les rattachements hiérarchiques et de compléter "
            "les métadonnées des indicateurs — unité, référence, cible, fréquence, source — que la "
            "matrice d'origine ne contient généralement pas.")

    titre2(document, "7.3 Import d'un XLSForm existant")
    para(document,
         "Un questionnaire déjà conçu pour KoboToolbox peut être importé : les feuilles "
         "« survey » et « choices » sont lues, les groupes deviennent des sections, les "
         "métadonnées techniques sont écartées et les contraintes conservées. Le questionnaire "
         "devient alors modifiable dans la plateforme et exportable en Word.")

    titre2(document, "7.4 Sauvegarde et transfert de projets entiers")
    para(document,
         "Au-delà de l'import de données, la plateforme permet de sauvegarder et de transférer des "
         "projets complets. Trois usages sont couverts : la sauvegarde de sécurité avant une "
         "modification lourde, le transfert d'un projet d'une instance à une autre — d'un poste "
         "local vers le serveur, ou entre deux organisations —, et le travail hors ligne dans un "
         "tableur suivi d'un rechargement.")
    for texte in [
        "La sauvegarde JSON d'un projet contient l'ensemble de ses données, questionnaires "
        "compris, et se recharge à l'identique.",
        "La sauvegarde JSON du portefeuille regroupe tous les projets de l'instance dans un "
        "fichier unique, ce qui constitue une copie de sécurité complète de la plateforme.",
        "Le classeur Excel de transfert reprend la structure du modèle d'import : il se retravaille "
        "dans un tableur, se transmet par courriel et se recharge tel quel.",
        "Le rapport d'import détaille, projet par projet, le nombre d'enregistrements recréés dans "
        "chaque catégorie et les éventuels avertissements.",
    ]:
        puce(document, texte)

    titre2(document, "7.5 Réinjection des données collectées")
    para(document,
         "L'export XLSX produit par KoboToolbox est réimportable. Les colonnes portant le nom "
         "technique d'une question reliée à un indicateur alimentent automatiquement les "
         "réalisations de cet indicateur. L'agrégation est choisie selon l'unité de mesure : "
         "somme pour les effectifs et les volumes, moyenne pour les pourcentages, scores, ratios "
         "et indices. Les valeurs produites sont enregistrées au statut « brouillon » et doivent "
         "être validées par le responsable de suivi-évaluation avant d'être considérées comme "
         "définitives.")

    document.add_page_break()

    # ------------------------------------------------- 8. Livrables
    titre1(document, "8. Livrables générés automatiquement")
    para(document,
         "Vingt-huit livrables sont produits à la demande, à partir des données saisies. Tous sont "
         "modifiables après téléchargement, ce qui préserve la liberté rédactionnelle des équipes "
         "tout en supprimant le travail de mise en forme.")
    tableau(document, ["Livrable", "Format", "Contenu et usage"],
            [["Cadre logique", "Excel", "Matrice A3 paysage, indicateurs agrégés par résultat, "
                                        "annexe des hypothèses critiques"],
             ["Cadre logique", "Word", "Matrice à quatre colonnes avec bandeaux de niveau colorés, "
                                       "insérable dans le document de projet"],
             ["Cadre de rendement", "Excel", "Taux de période, progression finale, statuts "
                                             "colorés, sources, méthodes, coûts, filtres actifs"],
             ["Cadre de rendement", "Word", "Version rédactionnelle pour rapport officiel"],
             ["Cadre de suivi des indicateurs (IPTT)", "Excel",
              "Cibles et réalisations par période, taux par période, mise en forme conditionnelle"],
             ["Chronogramme", "Excel", "Diagramme de Gantt mensuel coloré selon l'avancement, "
                                       "avec légende"],
             ["Chemin critique et réseau PERT", "Excel",
              "Ordonnancement CPM : dates au plus tôt et au plus tard, marges, durée du "
              "projet, activités par rang PERT"],
             ["Organigramme des tâches (WBS)", "Excel",
              "Décomposition codifiée, coûts consolidés, dictionnaire des lots de travail"],
             ["Matrice des responsabilités (RACI)", "Excel",
              "Matrice activités × acteurs, charge par partie prenante, contrôle de cohérence"],
             ["Organisation et ordonnancement", "Word",
              "Document réunissant l'organigramme des tâches, le chemin critique, le réseau "
              "PERT et la matrice RACI"],
             ["Projet complet (sauvegarde SEPIA)", "JSON",
              "Sauvegarde intégrale et réversible du projet, questionnaires compris, "
              "rechargeable sur une autre instance"],
             ["Portefeuille complet", "JSON",
              "Sauvegarde de tous les projets de l'instance dans un fichier unique"],
             ["Projet au format d'import", "Excel",
              "Toutes les données du projet dans la structure du modèle d'import, "
              "retravaillable dans un tableur puis rechargeable"],
             ["Plan de travail et budget annuel", "Excel",
              "Budget détaillé, ventilation trimestrielle, formules de totalisation, synthèse "
              "graphique par catégorie"],
             ["Registre des risques", "Excel", "Registre coté trié par criticité et matrice 5×5 "
                                               "renseignée"],
             ["Plan de gestion des risques", "Word", "Registre, matrice et plans de contingence "
                                                     "rédigés"],
             ["Fiches métadonnées des indicateurs", "Word",
              "Une fiche documentée par indicateur, avec série des cibles et réalisations"],
             ["Plan et manuel de suivi-évaluation", "Word",
              "Document maître en quinze chapitres, entièrement alimenté par les données du "
              "projet"],
             ["Rapport de performance", "Word", "Résumé exécutif, tableau des indicateurs, "
                                                "exécution physique et financière, alertes, "
                                                "canevas de mesures correctrices"],
             ["Rapport trimestriel de suivi", "Word",
              "Rapport périodé en huit parties, avec analyse d'équité et consolidation par zone"],
             ["Rapport semestriel d'avancement", "Word",
              "Même structure, sur un semestre de rapportage"],
             ["Rapport annuel de performance", "Word",
              "Bilan annuel consolidé, prêt pour le comité de pilotage et le bailleur"],
             ["Analyse d'équité et données désagrégées", "Excel",
              "Ventilation par catégorie, indice d'équité de genre, détail indicateur × modalité, "
              "graphique de répartition par sexe"],
             ["Consolidation par zone d'intervention", "Excel",
              "Bénéficiaires et indicateurs par zone, taux de couverture, coordonnées "
              "cartographiables dans Power BI ou un SIG, collecte par activité"],
             ["Revue qualité SMART", "Excel",
              "Diagnostic critère par critère, score du système, actions correctrices"],
             ["Bénéficiaires : ciblage et caractérisation", "Excel",
              "Quantification et atteinte par groupe, caractérisation qualitative, indicateurs "
              "rattachés à chaque groupe"],
             ["Partenaires : engagements et performance", "Excel",
              "Contributions conventionnées et décaissées, taux de décaissement, échéances de "
              "convention"],
             ["Évaluation CAD-OCDE", "Excel",
              "Notes et justifications par critère, moyennes consolidées, registre des "
              "recommandations et de leur mise en œuvre"],
             ["Rapport d'évaluation CAD-OCDE", "Word",
              "Rapport structuré critère par critère, échelle de notation, justifications et "
              "suivi des recommandations"],
             ["Protocole d'évaluation d'impact", "Word",
              "Méthode et hypothèse d'identification, groupes de traitement et de comparaison, "
              "contrôle de puissance, résultats et signification statistique"],
             ["Tableau de bord", "Excel", "Indicateurs clés, graphiques natifs Excel, feuille "
                                          "d'alertes, détail des indicateurs"],
             ["Jeu de données Power BI", "Excel", "Modèle en étoile, dimension calendrier, notice "
                                                  "de branchement et mesures DAX"],
             ["Questionnaires", "Word + XLSForm",
              "Version papier mise en page et version numérique téléversable"],
             ["Modèle d'import", "Excel", "Classeur type commenté, huit onglets avec exemples"],
             ["Dossier complet", "ZIP", "Archive de l'ensemble des livrables, organisée par "
                                        "format, avec notice d'utilisation"]],
            largeurs=[5, 2.6, 8.9], taille=8.5)

    titre2(document, "8.1 Le rapport périodique de suivi")
    para(document,
         "C'est le livrable le plus sollicité d'un dispositif de suivi-évaluation, et celui dont la "
         "production mobilise habituellement le plus de temps. La plateforme le compose en huit "
         "parties à partir des seules données de la période choisie.")
    tableau(document, ["Partie", "Contenu produit automatiquement"],
            [["1. Résumé exécutif", "Nombre d'indicateurs mesurés, taux moyen d'atteinte des "
                                    "cibles de la période, répartition par statut, avancement "
                                    "physique, exécution financière, bénéficiaires ventilés"],
             ["2. Performance des indicateurs", "Tableau code, indicateur, unité, cible de la "
                                                "période, réalisé, taux, statut coloré, source, "
                                                "responsable"],
             ["3. Analyse des données désagrégées", "Une sous-section par catégorie de "
                                                    "ventilation, avec effectifs, parts et total ; "
                                                    "indice d'équité de genre commenté"],
             ["4. Consolidation par zone", "Bénéficiaires atteints, cible, couverture et part des "
                                           "femmes par zone ; détail des réalisations par zone et "
                                           "par indicateur"],
             ["5. Exécution physique et financière", "Activités programmées, achevées, avancement "
                                                     "moyen, budget programmé et décaissé ; "
                                                     "données collectées par activité"],
             ["6. Difficultés et mesures correctrices", "Tableau des écarts constatés, à compléter "
                                                        "en réunion de revue (cause, mesure, "
                                                        "responsable, échéance)"],
             ["7. Qualité du dispositif", "Score SMART du système, taux de désagrégation, "
                                          "indicateurs à reprendre et actions recommandées"],
             ["8. Validation", "Bloc de signatures : rédacteur, vérificateur, approbateur"]],
            largeurs=[5, 11.5], taille=9)
    para(document,
         "Le périmètre temporel est déduit du libellé de la période : une période annuelle englobe "
         "ses semestres et ses trimestres, une période semestrielle englobe ses deux trimestres. "
         "L'utilisateur peut visualiser à l'écran le contenu du rapport avant de le générer.")

    titre2(document, "8.2 Le manuel de suivi-évaluation")
    para(document,
         "Ce livrable mérite une mention particulière : il constitue habituellement le document "
         "le plus coûteux à produire d'un dispositif de S&E, et le premier à devenir obsolète. La "
         "plateforme le génère en quinze chapitres — introduction, présentation du projet, cadre "
         "conceptuel et définitions, chaîne de résultats, système d'indicateurs, cadre de mesure "
         "du rendement, dispositif de collecte, gestion des risques et des hypothèses, "
         "planification opérationnelle, rapportage et diffusion, évaluations et études, "
         "dispositif organisationnel, assurance qualité des données, apprentissage et gestion des "
         "connaissances, budget du dispositif — avec les données réelles du projet insérées dans "
         "les tableaux. Il peut être régénéré à chaque revue annuelle, ce qui garantit qu'il "
         "reflète en permanence le dispositif effectivement en vigueur.")

    document.add_page_break()

    # ------------------------------------------------- 9. Collecte
    titre1(document, "9. Fiches de collecte et questionnaires numériques")

    titre2(document, "9.1 Types de questions pris en charge")
    tableau(document, ["Type", "Usage", "Rendu papier"],
            [["text", "Réponse ouverte courte ou longue", "Lignes pointillées"],
             ["integer", "Effectif, comptage", "Cases de saisie chiffre par chiffre"],
             ["decimal", "Superficie, rendement, montant", "Cases de saisie"],
             ["select_one", "Choix unique", "Cercles à cocher avec codes"],
             ["select_multiple", "Choix multiples", "Carrés à cocher avec codes"],
             ["date", "Date d'événement", "Gabarit JJ/MM/AAAA"],
             ["time", "Heure", "Zone de saisie"],
             ["geopoint", "Localisation GPS", "Champs latitude et longitude"],
             ["calculate", "Valeur dérivée (ex. rendement)", "Non imprimé"],
             ["note", "Consigne à l'enquêteur", "Texte en italique"],
             ["image", "Photographie de preuve", "Mention de la pièce jointe"],
             ["barcode", "Code-barres, identifiant", "Zone de saisie"]],
            largeurs=[3.5, 7.5, 5.5], taille=9)

    titre2(document, "9.2 Contrôles de qualité intégrés au formulaire")
    para(document,
         "La qualité d'une donnée se joue au moment de sa saisie. La plateforme permet de "
         "définir, pour chaque question, une contrainte de validité et le message affiché "
         "lorsqu'elle est violée. Ces contrôles sont transposés tels quels dans le XLSForm et "
         "s'appliquent donc sur le terminal de l'enquêteur, y compris hors connexion.")
    code(document, [
        "Contrainte d'âge          . >= 15 and . <= 110",
        "Contrainte de superficie  . >= 0 and . <= 50",
        "Logique de saut           ${consentement} = '1'",
        "Calcul de rendement       if(${superficie} > 0,",
        "                             ${production} div (${superficie} * 1000), 0)",
    ])

    titre2(document, "9.3 Structure du XLSForm produit")
    para(document,
         "Le classeur généré comporte les trois feuilles attendues par la norme XLSForm — "
         "survey, choices, settings — ainsi qu'une notice de déploiement. La plateforme ajoute "
         "automatiquement les métadonnées de collecte (start, end, today, deviceid) et un groupe "
         "d'identification comprenant la date, l'enquêteur, la localité et le point GPS. Les noms "
         "de variables sont normalisés — minuscules, sans accent ni espace, unicité garantie — "
         "afin de respecter les contraintes d'ODK.")
    para(document,
         "Le déploiement s'effectue en trois étapes : ouvrir KoboToolbox, créer un projet à "
         "partir d'un fichier XLSForm téléversé, puis déployer. Les enquêteurs collectent ensuite "
         "avec KoboCollect ou ODK Collect, en mode hors ligne, et synchronisent lorsqu'une "
         "connexion est disponible.")

    document.add_page_break()

    # ------------------------------------------------- 10. BI
    titre1(document, "10. Tableaux de bord et connexion Power BI")

    titre2(document, "10.1 Tableaux de bord intégrés")
    para(document,
         "Les tableaux de bord de la plateforme sont conçus pour la lecture immédiate. Ils "
         "reposent sur six graphiques produits en SVG : jauge de santé, anneau de répartition par "
         "statut, barres de performance par niveau de résultat, colonnes de programmation "
         "budgétaire, diagramme de Gantt et matrice des risques. Aucun de ces graphiques ne "
         "dépend d'une bibliothèque externe, ce qui garantit leur affichage en toutes "
         "circonstances, y compris sur un réseau filtrant les CDN.")

    titre2(document, "10.2 Tableau de bord Excel automatisé")
    para(document,
         "Le classeur généré comporte une feuille de synthèse avec six indicateurs clés et trois "
         "graphiques natifs Excel — camembert, barres, colonnes — alimentés par une feuille de "
         "données, une feuille d'alertes et une feuille de détail des indicateurs avec filtres et "
         "mise en forme conditionnelle. Étant construit sur des graphiques Excel natifs, il reste "
         "modifiable et réutilisable par les équipes.")

    titre2(document, "10.3 Connexion Power BI")
    para(document, "Deux méthodes de branchement sont proposées.")
    titre3(document, "Méthode 1 — Flux web à actualisation directe")
    numero(document, "Ouvrir Power BI Desktop, puis Accueil > Obtenir des données > Web.")
    numero(document, "Coller l'URL du jeu de données fournie par la vue « Power BI » de la "
                     "plateforme, qui comprend le jeton d'accès de l'utilisateur.")
    numero(document, "Dans l'éditeur Power Query, développer la colonne « tables », puis chaque "
                     "table souhaitée.")
    numero(document, "Créer les relations entre tables de dimensions et tables de faits dans la "
                     "vue Modèle.")
    para(document,
         "Chaque table est également exposée individuellement, en JSON et en CSV, ce qui permet "
         "de créer une requête distincte par table — approche généralement préférable pour la "
         "maintenance du rapport.", taille=10)

    titre3(document, "Méthode 2 — Classeur structuré")
    para(document,
         "Lorsque le poste d'analyse n'a pas accès à la plateforme, le classeur « Jeu de données "
         "Power BI » fournit le même modèle sous forme de fichier, accompagné d'une notice "
         "détaillant les relations à créer et les mesures DAX recommandées.")

    titre2(document, "10.4 Cartographie")
    para(document,
         "La plateforme intègre une carte de couverture des zones d'intervention, décrite au "
         "chapitre 5.12. Pour une analyse cartographique plus poussée, les coordonnées de chaque "
         "zone sont exportées dans le classeur « Consolidation par zone » et exposées dans la "
         "table Dim_Zone du flux Power BI : le visuel Carte de Power BI, comme un système "
         "d'information géographique tel que QGIS, peut les exploiter directement, en les croisant "
         "avec les bénéficiaires atteints ou le taux de couverture.")

    titre2(document, "10.5 Modèle en étoile exposé")
    tableau(document, ["Table", "Nature", "Contenu"],
            [["Dim_Projet", "Dimension", "Identification et caractéristiques du projet"],
             ["Dim_Resultat", "Dimension", "Chaîne de résultats hiérarchisée"],
             ["Dim_Indicateur", "Dimension", "Métadonnées complètes des indicateurs"],
             ["Dim_Zone", "Dimension", "Zones d'intervention, population, cible de bénéficiaires, "
                                       "latitude et longitude"],
             ["Dim_Calendrier", "Dimension", "Table de dates mensuelle et trimestrielle"],
             ["Fait_Cible", "Fait", "Cibles périodiques"],
             ["Fait_Realisation", "Fait", "Réalisations mesurées, localisées et rattachées à une "
                                          "activité, taux et statut de performance"],
             ["Fait_Desagregation", "Fait", "Table dépliée : une ligne par modalité de "
                                            "désagrégation (sexe, âge, groupe cible)"],
             ["Fait_Activite", "Fait", "Activités, avancement et coûts"],
             ["Fait_Budget", "Fait", "Lignes budgétaires, engagements et décaissements"],
             ["Fait_Risque", "Fait", "Risques cotés et niveaux de criticité"]],
            largeurs=[4, 3, 9.5], taille=9.5)
    para(document,
         "La table Fait_Desagregation est volontairement dépliée : elle s'utilise directement dans "
         "un histogramme empilé segmenté par sexe ou par groupe cible, sans transformation "
         "préalable. Les coordonnées portées par Dim_Zone permettent un visuel cartographique "
         "immédiat.")

    titre2(document, "10.6 Mesures DAX recommandées")
    code(document, [
        "Taux de réalisation =",
        "    DIVIDE(SUM(Fait_Realisation[ValeurRealisee]), SUM(Fait_Cible[ValeurCible]))",
        "",
        "Taux d'exécution budgétaire =",
        "    DIVIDE(SUM(Fait_Budget[Decaisse]), SUM(Fait_Budget[TotalPlanifie]))",
        "",
        "Avancement physique moyen = AVERAGE(Fait_Activite[Avancement])",
        "",
        "Risques critiques =",
        "    CALCULATE(COUNTROWS(Fait_Risque), Fait_Risque[Niveau] = \"Critique\")",
        "",
        "Bénéficiaires femmes =",
        "    CALCULATE(SUM(Fait_Desagregation[Valeur]),",
        "              Fait_Desagregation[Categorie] = \"Sexe\",",
        "              Fait_Desagregation[Modalite] = \"Femme\")",
        "",
        "Part des femmes =",
        "    DIVIDE([Bénéficiaires femmes],",
        "           CALCULATE(SUM(Fait_Desagregation[Valeur]),",
        "                     Fait_Desagregation[Categorie] = \"Sexe\"))",
        "",
        "Taux de couverture de zone =",
        "    DIVIDE(SUM(Fait_Realisation[ValeurRealisee]), SUM(Dim_Zone[CibleBeneficiaires]))",
    ])

    document.add_page_break()

    # ------------------------------------------------- 11. Sécurité
    titre1(document, "11. Utilisateurs, rôles et sécurité")

    titre2(document, "11.1 Rôles et droits")
    tableau(document, ["Rôle", "Droits accordés"],
            [["Lecteur", "Consultation et téléchargement des livrables, sur les seuls projets "
                         "dont il est membre"],
             ["Opérateur de saisie", "Droits du lecteur, plus création et modification des "
                                     "données : indicateurs, réalisations, activités, lignes "
                                     "budgétaires, risques, questionnaires"],
             ["Responsable suivi-évaluation", "Droits précédents, plus création et paramétrage "
                                              "de projets et consultation du journal d'audit"],
             ["Coordonnateur", "Droits précédents, dans une logique de pilotage et de validation"],
             ["Administrateur", "Accès complet, gestion des comptes, des rôles, des accès par "
                                "projet et des clés d'API"]],
            largeurs=[4.5, 12])
    para(document,
         "Le rôle détermine ce qu'un utilisateur peut faire ; le rattachement au projet détermine "
         "sur quoi il peut le faire. Les deux contrôles sont indépendants et cumulatifs, et tous "
         "deux sont appliqués par le serveur : l'interface masque des écrans selon le rôle, mais "
         "cette décision est cosmétique et ne conditionne aucun accès.")

    titre2(document, "11.2 Principes de conception retenus")
    for intitule, texte in [
        ("Aucun secret dans le dépôt",
         "Il n'existe ni clé, ni jeton, ni mot de passe inscrit dans le code ou dans les fichiers "
         "de configuration versionnés. En production, l'absence de clé de signature empêche le "
         "démarrage plutôt que de retomber silencieusement sur une valeur connue de tous. Si "
         "aucun mot de passe administrateur n'est fourni, la plateforme en engendre un "
         "aléatoirement et l'inscrit une seule fois dans les journaux de démarrage."),
        ("Le navigateur ne décide de rien",
         "Chaque appel revérifie côté serveur l'authentification, le rôle et l'appartenance de "
         "l'objet manipulé à un projet autorisé. Un identifiant deviné dans une URL ne donne "
         "rien. Les projets étant énumérables par construction, un accès refusé répond 404 et "
         "non 403 : la réponse ne confirme pas l'existence de la ressource."),
        ("Le serveur ne fait confiance à rien de ce qu'il reçoit",
         "Les champs modifiables sont définis par liste blanche : un client qui ajoute un rôle, "
         "un état de compte ou une empreinte de mot de passe à sa requête voit ces champs "
         "ignorés. Aucune requête n'est construite par concaténation de texte, et les jokers des "
         "recherches sont échappés."),
        ("Tout texte affiché est échappé",
         "Les libellés, énoncés, commentaires et recommandations issus des données sont échappés "
         "avant insertion dans le HTML ou le SVG. Un énoncé d'indicateur contenant du balisage "
         "s'affiche comme du texte. La politique de sécurité du contenu interdit par ailleurs "
         "tout script en ligne et tout script d'origine externe."),
    ]:
        puce(document, f"{intitule} — {texte}")

    titre2(document, "11.3 Mesures appliquées")
    tableau(document, ["Domaine", "Mesure"],
            [["Mots de passe",
              "PBKDF2-SHA256, 240 000 itérations, sel aléatoire par compte. Politique imposée : "
              "12 caractères au minimum, au moins trois classes de caractères, refus des mots de "
              "passe courants, des suites de caractères et de tout mot de passe contenant le nom "
              "ou l'adresse du compte."],
             ["Sessions",
              "Jeton signé HMAC-SHA256 déposé dans un cookie HttpOnly, SameSite=Strict et Secure "
              "en production : le JavaScript ne peut pas le lire. Aucun jeton en stockage local. "
              "Chaque compte peut fermer toutes ses sessions d'un appel."],
             ["Force brute",
              "Verrouillage progressif du compte après échecs répétés. Message d'échec unique et "
              "comparaison à durée constante même lorsque le compte n'existe pas : l'API ne "
              "permet pas d'énumérer les comptes."],
             ["Limitation de débit",
              "Fenêtre glissante par adresse et par catégorie d'appel, avec des quotas distincts "
              "pour l'authentification, les écritures, les exports et la lecture. Plafond de "
              "taille sur le corps des requêtes."],
             ["En-têtes",
              "Politique de sécurité du contenu, interdiction d'inclusion dans un cadre, "
              "politique de référent, HSTS en production, absence de mise en cache des réponses "
              "d'API."],
             ["Origines",
              "Partage entre origines fermé par défaut : sans configuration explicite, aucune "
              "origine tierce n'est acceptée."],
             ["Comptes",
              "Aucune inscription libre : les comptes sont ouverts par un administrateur. "
              "L'adresse doit être confirmée par un lien à jeton — à usage unique, avec une "
              "réponse indistincte en cas de jeton faux — avant toute connexion. Changement du "
              "mot de passe imposé au premier accès."],
             ["Clés d'API",
              "Le flux Power BI n'utilise plus de jeton de session dans l'URL mais des clés "
              "nominatives, en lecture seule, limitées à un projet, datées, révocables et "
              "stockées hachées : leur valeur n'est affichée qu'à la création."],
             ["Téléversements",
              "Plafond de taille, contrôle du nombre magique du fichier et non de son extension, "
              "refus des archives dont le ratio de décompression dépasse le seuil, analyse en "
              "mémoire sans écriture sur le disque."],
             ["Erreurs",
              "Aucune trace d'exécution renvoyée au client : une erreur produit un identifiant "
              "de corrélation affiché à l'utilisateur et une entrée complète dans les journaux "
              "du serveur."],
             ["Dépendances",
              "Dix dépendances directes, toutes épinglées à une version précise et vérifiées, "
              "aucune dépendance JavaScript. Revue trimestrielle recommandée."]],
            largeurs=[3.5, 13])

    titre2(document, "11.4 Jeu de vérification")
    para(document,
         "Les mesures ci-dessus sont éprouvées par un jeu de cinquante-et-un contrôles automatisés "
         "exécutés contre l'application réelle, sur une base temporaire, par la commande "
         "« python scripts/verifier_securite.py » : refus des identifiants erronés avec message "
         "indistinct, attributs du cookie de session, présence et contenu des en-têtes, refus "
         "systématique des appels anonymes, tentative d'élévation de privilège par le corps de la "
         "requête, cycle complet de confirmation d'adresse avec rejeu d'un jeton déjà consommé, "
         "cloisonnement d'un lecteur non membre, politique de mot de passe, téléversements piégés, "
         "jokers dans la recherche et fermeture effective de la session. La commande sort en "
         "erreur dès qu'un contrôle échoue : elle peut être intégrée à une chaîne d'intégration "
         "continue.")

    titre2(document, "11.5 Traçabilité")
    para(document,
         "Toute opération d'écriture — création, modification, suppression, import — est inscrite "
         "au journal d'audit avec son horodatage, l'adresse de son auteur, l'entité concernée et "
         "le projet. Ce journal, consultable par les responsables de suivi-évaluation et les "
         "administrateurs, constitue l'élément de preuve requis lors des audits de qualité des "
         "données et des missions de supervision.")

    titre2(document, "11.6 Reprendre la main sur le compte d'administration")
    para(document,
         "Mot de passe perdu, compte verrouillé par des tentatives infructueuses, désactivé ou "
         "rétrogradé par erreur : la plateforme n'offre aucun point d'entrée réseau de "
         "réinitialisation. Un tel point d'entrée serait une porte dérobée permanente, exposée à "
         "quiconque connaît l'adresse du service. La reprise de main passe donc par une preuve "
         "d'autorité réelle : l'accès au tableau de bord d'hébergement, ou au serveur lui-même.")
    tableau(document, ["Situation", "Marche à suivre"],
            [["Hébergement sans accès shell (Render, plan gratuit compris)",
              "Définir la variable SEPIA_ADMIN_RESET à 1 — et facultativement "
              "SEPIA_ADMIN_PASSWORD pour choisir le mot de passe —, redéployer, relever le mot de "
              "passe dans les journaux de démarrage à la mention « COMPTE ADMINISTRATEUR "
              "RÉINITIALISÉ », puis retirer la variable : tant qu'elle est présente, chaque "
              "redémarrage réinitialise le compte."],
             ["Accès au serveur ou à la base",
              "Exécuter « python scripts/reinitialiser_admin.py ». Le script demande "
              "confirmation, agit sur la base désignée par DATABASE_URL et affiche le mot de "
              "passe une seule fois."],
             ["Créer un second administrateur plutôt que réinitialiser",
              "Donner une nouvelle valeur à SEPIA_ADMIN_EMAIL : l'amorçage ne trouve pas ce "
              "compte et le crée. L'ancien reste en place et pourra être supprimé depuis la vue "
              "Administration."]],
            largeurs=[4.5, 12])
    para(document,
         "Dans tous les cas, le compte est recréé s'il a disparu et remis en état s'il existe : "
         "nouveau mot de passe, rôle d'administrateur rétabli, compte réactivé, verrouillage et "
         "tentatives infructueuses effacés, adresse considérée comme confirmée. Le changement du "
         "mot de passe est exigé à la connexion suivante et toutes les sessions ouvertes sont "
         "fermées — si le mot de passe a été perdu, rien ne permet d'exclure qu'il l'ait été au "
         "profit de quelqu'un d'autre.")

    titre2(document, "11.7 Recommandations de mise en production")
    for texte in [
        "Renseigner la clé de signature des jetons et le mot de passe administrateur avant "
        "l'ouverture du service ; le blueprint de déploiement engendre la première "
        "automatiquement.",
        "Restreindre les origines autorisées au seul domaine de l'application si des appels "
        "entre domaines sont nécessaires ; sans configuration, aucune origine tierce n'est "
        "acceptée.",
        "Révoquer les clés d'API dès qu'un utilisateur quitte le projet : elles survivent à la "
        "fermeture de sa session.",
        "Créer un compte nominatif par utilisateur, plutôt qu'un compte partagé, afin que le "
        "journal d'audit conserve sa valeur probante.",
        "Attribuer le rôle strictement nécessaire à chaque utilisateur, et réserver le rôle "
        "d'administrateur à deux personnes au plus.",
        "Vérifier que les sauvegardes automatiques de la base de données sont actives et tester "
        "périodiquement la procédure de restauration.",
    ]:
        puce(document, texte)

    document.add_page_break()

    # ------------------------------------------------- 12. Prise en main
    titre1(document, "12. Guide de prise en main")
    para(document,
         "La séquence ci-après permet de rendre un projet pleinement opérationnel dans la "
         "plateforme. Elle suppose que le cadre logique et le budget du projet ont été validés.")

    titre2(document, "12.1 Les deux projets d'exemple")
    para(document,
         "Une instance neuve est livrée avec deux projets complets, destinés à servir de cas "
         "pratiques. Ils permettent de parcourir chaque écran avec des données cohérentes avant "
         "de saisir son propre projet, et de comparer ce que l'on a renseigné à ce qu'un "
         "dispositif complet contient.")
    tableau(document, ["Projet", "Secteur", "Contenu renseigné"],
            [["PADRA-2025 — Programme d'appui au développement rural et à l'agriculture",
              "Agriculture, sécurité alimentaire",
              "Cadre logique complet, zones d'intervention, indicateurs désagrégés, activités "
              "ordonnancées avec antécédences, budget, risques, hypothèses, parties prenantes et "
              "matrice RACI, instruments de collecte"],
             ["PASSE-2026 — Programme de santé scolaire et de scolarisation des filles",
              "Éducation, santé",
              "Huit zones, trois effets et cinq produits, cinq groupes de bénéficiaires, six "
              "partenaires, dix-sept indicateurs dont trois de processus, trente-sept cibles et "
              "cent une réalisations réparties par zone et par sexe, treize activités enchaînées, "
              "quatorze lignes budgétaires, sept risques, six hypothèses, treize parties "
              "prenantes et soixante-quinze affectations RACI, deux évaluations CAD assorties de "
              "quatre recommandations, deux études d'impact — un essai randomisé par grappes et "
              "une combinaison doubles différences et appariement —, deux questionnaires"]],
            largeurs=[5, 3, 8.5], taille=9)
    para(document,
         "Le second projet renseigne l'ensemble des rubriques de la plateforme, y compris celles "
         "que les dispositifs réels laissent le plus souvent vides. La variable "
         "SEPIA_SEED_DEMO fixée à 0 permet de démarrer sur une instance vierge.")

    titre2(document, "12.2 Séquence de mise en œuvre")
    etapes = [
        ("Créer le projet",
         "Vue Portefeuille, bouton « Nouveau projet ». Renseigner le code, l'intitulé, le "
         "bailleur, l'agence d'exécution, les dates et le budget. La théorie du changement et "
         "l'approche de suivi-évaluation peuvent être complétées ultérieurement : elles "
         "alimenteront le manuel de S&E."),
        ("Charger le cadre logique",
         "Deux voies. Si le cadre logique existe sous forme de tableau, télécharger le modèle "
         "d'import depuis la vue Importer, y transposer les données, puis charger le classeur. "
         "Sinon, construire l'arborescence directement dans la vue Cadre logique, du niveau "
         "impact vers les activités."),
        ("Déclarer les zones d'intervention",
         "Vue Zones d'intervention. Saisir le découpage géographique du projet, du niveau le plus "
         "large vers le plus fin, en renseignant pour chaque zone sa population, sa cible de "
         "bénéficiaires, son responsable et ses coordonnées géographiques — ces dernières "
         "conditionnent l'affichage de la carte de couverture. Cette étape commande la "
         "consolidation territoriale et le calcul des taux de couverture ; elle doit précéder la "
         "saisie des réalisations."),
        ("Paramétrer les indicateurs",
         "Pour chaque indicateur, compléter la fiche métadonnée : définition opérationnelle, "
         "unité, mode de calcul, désagrégations exigées (au minimum le sexe), valeur de référence "
         "et sa date, cible finale et son échéance, sens de progression, règle d'agrégation, "
         "fréquence, source, méthode de collecte, responsable. Marquer comme « indicateur clé » "
         "ceux qui seront présentés au comité de pilotage. Vérifier ensuite le score SMART dans la "
         "vue Qualité des indicateurs et traiter les actions correctrices recommandées."),
        ("Générer les cibles périodiques",
         "Depuis la fiche de suivi de chaque indicateur, utiliser « Générer les cibles "
         "périodiques » en choisissant la granularité. La plateforme interpole linéairement entre "
         "la référence et la cible finale ; les valeurs obtenues doivent ensuite être ajustées "
         "pour tenir compte du rythme réel de montée en charge du projet."),
        ("Saisir le chronogramme et ordonnancer",
         "Créer les activités, en les rattachant aux produits du cadre logique, avec dates, "
         "responsables et coûts prévus. Marquer les jalons. Renseigner les antécédents de "
         "chaque activité : ils déclenchent le calcul du chemin critique, de la durée du "
         "projet et du réseau PERT. Vérifier ensuite l'organigramme des tâches et lancer la "
         "codification WBS."),
        ("Établir la matrice des responsabilités",
         "Recenser les parties prenantes du projet dans l'onglet RACI, puis attribuer à "
         "chacune un rôle sur chaque activité : un seul approbateur, au moins un réalisateur. "
         "Traiter les anomalies signalées avant de soumettre la matrice au comité de "
         "pilotage."),
        ("Saisir le PTBA",
         "Créer les lignes budgétaires, rattachées aux activités, avec quantité, coût unitaire et "
         "ventilation trimestrielle. Cette ventilation alimente la programmation présentée au "
         "tableau de bord."),
        ("Constituer le registre des risques",
         "Identifier les risques par catégorie, les coter en probabilité et en impact, documenter "
         "les mesures d'atténuation et les plans de contingence, désigner un porteur et fixer une "
         "date de revue. Renseigner en parallèle les hypothèses critiques du cadre logique."),
        ("Concevoir les instruments de collecte",
         "Créer les fiches et questionnaires, structurer les questions par section, définir les "
         "modalités, les contraintes de saisie et les logiques de saut, puis relier les questions "
         "aux indicateurs qu'elles alimentent."),
        ("Déployer la collecte",
         "Exporter chaque instrument en XLSForm, le téléverser dans KoboToolbox, le déployer, "
         "puis former les enquêteurs. Exporter également la version Word pour les contextes où "
         "l'administration papier reste nécessaire."),
        ("Alimenter le suivi",
         "À chaque échéance de collecte, saisir les réalisations dans la vue Saisie des "
         "réalisations en renseignant la période, la zone, l'activité source et la ventilation par "
         "sexe et par groupe cible ; ou réimporter l'export KoboToolbox. La vue Cadre de suivi "
         "permet une saisie en masse pour les indicateurs non désagrégés. Mettre à jour "
         "l'avancement des activités et les montants engagés et décaissés."),
        ("Analyser et décider",
         "Consulter le tableau de bord, traiter les alertes par ordre de gravité, vérifier "
         "l'équilibre territorial dans la vue Zones et l'inclusivité dans la vue Équité, "
         "documenter les causes des écarts et arrêter les mesures correctrices lors de la revue "
         "périodique."),
        ("Évaluer",
         "Ouvrir un exercice évaluatif dans la vue Évaluation CAD-OCDE, noter chacun des six "
         "critères en justifiant la note au regard des points d'examen, puis enregistrer les "
         "recommandations avec leur criticité, la réponse du management, un responsable et une "
         "échéance. Le suivi de leur mise en œuvre se fait ensuite depuis la même vue : une "
         "recommandation échue et non soldée y est signalée."),
        ("Mesurer l'impact",
         "Concevoir l'étude dans la vue Évaluation d'impact avant le démarrage des activités : "
         "choisir la méthode et expliciter son hypothèse d'identification, définir la règle "
         "d'affectation, puis dimensionner l'échantillon à l'aide du calculateur — en renseignant "
         "l'écart-type de l'indicateur de résultat et, en cas de randomisation par grappes, la "
         "taille des grappes et la corrélation intra-grappe. Une étude dimensionnée après coup ne "
         "peut plus l'être : l'affectation a déjà eu lieu."),
        ("Produire les rapports et livrables",
         "Depuis la vue Rapports périodiques, choisir la période et générer le rapport "
         "trimestriel, semestriel ou annuel après en avoir vérifié l'aperçu. Depuis la vue "
         "Livrables, générer les autres documents attendus ou télécharger le dossier complet au "
         "format ZIP avant chaque comité de pilotage."),
    ]
    for index, (titre_etape, description) in enumerate(etapes, start=1):
        titre3(document, f"Étape {index} — {titre_etape}")
        para(document, description, taille=10)

    encadre(document, "Ordre recommandé",
            "Le cadre logique doit précéder les indicateurs, qui doivent précéder les cibles "
            "périodiques ; les zones et les activités doivent précéder les réalisations, qui s'y "
            "rattachent ; les activités doivent précéder les lignes budgétaires. Cet ordre "
            "garantit que chaque élément trouve son rattachement au moment de sa création et "
            "évite les reprises ultérieures.")

    document.add_page_break()

    # ------------------------------------------------- 13. Déploiement
    titre1(document, "13. Déploiement sur GitHub et Render")

    titre2(document, "13.1 Dépôt de code")
    para(document,
         "Le code source est hébergé sur GitHub. Le dépôt contient l'application, l'interface, "
         "la documentation et le fichier de description d'infrastructure destiné à Render.")
    code(document, [
        "git clone https://github.com/Ricard228/sepia-erp.git",
        "cd sepia-erp",
        "python -m venv .venv && .venv\\Scripts\\activate",
        "pip install -r requirements.txt",
        "uvicorn app.main:app --reload --port 8000",
    ])

    titre2(document, "13.2 Déploiement par blueprint")
    numero(document, "Se connecter à Render, choisir New puis Blueprint.")
    numero(document, "Sélectionner le dépôt sepia-erp. Render lit le fichier render.yaml et "
                     "provisionne simultanément le service web et la base PostgreSQL.")
    numero(document, "Renseigner la variable SEPIA_ADMIN_PASSWORD dans le tableau de bord Render.")
    numero(document, "Lancer le déploiement. La sonde /api/sante confirme la disponibilité du "
                     "service.")

    titre2(document, "13.3 Déploiement manuel")
    tableau(document, ["Paramètre", "Valeur"],
            [["Environnement", "Python 3"],
             ["Commande de construction", "pip install -r requirements.txt"],
             ["Commande de démarrage", "uvicorn app.main:app --host 0.0.0.0 --port $PORT"],
             ["Chemin de la sonde de santé", "/api/sante"]],
            largeurs=[5.5, 11])
    para(document,
         "Créer ensuite une base PostgreSQL sur Render et lier sa chaîne de connexion à la "
         "variable DATABASE_URL du service web.")

    titre2(document, "13.4 Variables d'environnement")
    tableau(document, ["Variable", "Rôle", "Valeur par défaut"],
            [["DATABASE_URL", "Chaîne de connexion PostgreSQL", "SQLite local"],
             ["SEPIA_ENV", "« production » active les garde-fous bloquants : cookie Secure, "
                           "HSTS, erreurs non détaillées",
              "production si DATABASE_URL est définie"],
             ["SEPIA_SECRET_KEY", "Clé de signature des jetons",
              "obligatoire en production : le démarrage est refusé sans elle"],
             ["SEPIA_ADMIN_EMAIL", "Compte administrateur initial", "admin@sepia.org"],
             ["SEPIA_ADMIN_PASSWORD", "Mot de passe initial",
              "aucun : engendré aléatoirement et journalisé une seule fois"],
             ["SEPIA_TOKEN_TTL", "Durée de validité des jetons, en secondes", "43200"],
             ["SEPIA_SEED_DEMO", "Chargement des projets d'exemple", "1"],
             ["SEPIA_ADMIN_RESET", "Réinitialise le compte d'administration au démarrage",
              "vide : aucune réinitialisation"],
             ["SEPIA_CORS_ORIGINS", "Origines autorisées pour les appels entre domaines",
              "vide : aucune origine tierce"]],
            largeurs=[5, 7.5, 4], taille=9.5)

    encadre(document, "Point de vigilance",
            "Sur le plan gratuit de Render, le disque du service web n'est pas persistant : les "
            "fichiers écrits sont perdus à chaque redéploiement et à chaque mise en veille. Le "
            "recours à la base PostgreSQL est donc indispensable en production. En développement "
            "local, la base SQLite convient parfaitement.")

    document.add_page_break()

    # ------------------------------------------------- 14. API
    titre1(document, "14. Interface de programmation (API)")
    para(document,
         "L'ensemble des fonctions de la plateforme est accessible par une API REST documentée "
         "automatiquement. La documentation interactive est disponible à l'adresse /api/docs et "
         "permet d'exécuter les appels depuis le navigateur.")
    tableau(document, ["Méthode et chemin", "Fonction"],
            [["POST /api/auth/login", "Authentification, renvoie un jeton de session"],
             ["GET /api/projects", "Liste des projets"],
             ["GET /api/dashboard/{id}", "Tableau de bord complet d'un projet"],
             ["GET /api/portefeuille", "Consolidation multi-projets"],
             ["GET /api/logframe/tree/{id}", "Arborescence du cadre logique avec indicateurs"],
             ["GET /api/indicateurs/suivi/{id}", "Grille IPTT cibles/réalisations"],
             ["POST /api/indicators/{id}/saisie", "Saisie ou mise à jour d'une réalisation"],
             ["POST /api/projects/{id}/periodes", "Génération automatique des cibles périodiques"],
             ["POST /api/imports/excel/{id}", "Import d'un classeur de projet"],
             ["POST /api/imports/word/analyser", "Analyse des tableaux d'un document Word"],
             ["POST /api/imports/kobo/{form_id}", "Réinjection de données collectées"],
             ["GET /api/exports/{id}/{livrable}", "Téléchargement d'un livrable"],
             ["GET /api/exports/{id}/dossier-complet", "Archive ZIP de tous les livrables"],
             ["GET /api/powerbi/{id}/dataset", "Flux de données pour Power BI"],
             ["GET /api/sante", "Sonde de disponibilité"]],
            largeurs=[7, 9.5], taille=9)
    para(document,
         "Les entités métier — cadre logique, indicateurs, cibles, réalisations, risques, "
         "hypothèses, activités, lignes budgétaires, formulaires, questions — disposent chacune "
         "d'un jeu complet d'opérations de création, lecture, mise à jour, suppression et "
         "création en lot, engendrées par une fabrique commune. Cette approche garantit "
         "l'homogénéité du comportement de l'API et réduit la surface de défaut.")

    document.add_page_break()

    # ------------------------------------------------- 15. Exploitation
    titre1(document, "15. Exploitation, maintenance et évolutions")

    titre2(document, "15.1 Exploitation courante")
    tableau(document, ["Périodicité", "Opération"],
            [["Continue", "Saisie des réalisations et mise à jour de l'avancement des activités"],
             ["Mensuelle", "Mise à jour des engagements et décaissements ; revue des activités "
                           "en retard"],
             ["Trimestrielle", "Génération du rapport de performance ; revue du registre des "
                               "risques ; validation des données collectées"],
             ["Semestrielle", "Atelier de revue de performance ; actualisation des hypothèses ; "
                              "dossier complet pour le comité de pilotage"],
             ["Annuelle", "Enquête de suivi des effets ; élaboration du PTBA de l'exercice "
                          "suivant ; régénération du manuel de suivi-évaluation ; audit de la "
                          "qualité des données"]],
            largeurs=[3.5, 13])

    titre2(document, "15.2 Maintenance technique")
    for texte in [
        "Mettre à jour périodiquement les dépendances Python, en vérifiant au préalable la "
        "compatibilité par un déploiement de test.",
        "Surveiller la sonde de santé et les journaux applicatifs fournis par Render.",
        "Vérifier la présence et l'intégrité des sauvegardes de la base de données.",
        "Renouveler la clé de signature des jetons en cas de suspicion de compromission ; cette "
        "opération invalide toutes les sessions en cours.",
    ]:
        puce(document, texte)

    titre2(document, "15.3 Évolutions envisageables")
    para(document,
         "L'architecture retenue autorise plusieurs extensions sans refonte, par ordre croissant "
         "d'effort de mise en œuvre.")
    tableau(document, ["Évolution", "Apport attendu", "Effort"],
            [["Connexion directe à l'API KoboToolbox", "Suppression de l'étape manuelle "
                                                       "d'export-import des données collectées",
              "Modéré"],
             ["Module de cartographie des interventions", "Visualisation géographique des "
                                                          "réalisations à partir des points GPS "
                                                          "déjà collectés", "Modéré"],
             ["Notifications par courriel et messagerie", "Alerte automatique des responsables en "
                                                          "cas d'échéance ou de dérive", "Modéré"],
             ["Gestion documentaire", "Rattachement des pièces justificatives aux activités et "
                                      "aux réalisations", "Modéré"],
             ["Application mobile hors ligne native", "Collecte autonome sans dépendance à un "
                                                      "outil tiers", "Important"],
             ["Module d'évaluation d'impact", "Analyse contrefactuelle intégrée (groupes de "
                                              "comparaison, appariement)", "Important"]],
            largeurs=[5.5, 8, 3], taille=9.5)

    document.add_page_break()

    # ------------------------------------------------- Annexes
    titre1(document, "Annexe 1 — Glossaire du suivi-évaluation")
    tableau(document, ["Terme", "Définition"],
            [["Cadre logique", "Matrice de planification reliant la logique d'intervention, les "
                               "indicateurs objectivement vérifiables, les sources de "
                               "vérification et les hypothèses"],
             ["Cadre de rendement", "Instrument précisant, pour chaque indicateur, la source, la "
                                    "méthode, la fréquence, le responsable et le coût de la mesure"],
             ["IPTT", "Indicator Performance Tracking Table : tableau de suivi croisant "
                      "indicateurs et périodes, en cibles et en réalisations"],
             ["Chaîne de résultats", "Enchaînement intrants, activités, produits, effets, impact"],
             ["Impact", "Changement de long terme auquel le projet contribue sans en être seul "
                        "responsable"],
             ["Effet", "Changement de comportement, de capacité ou de performance directement "
                       "attribuable au projet"],
             ["Produit", "Bien ou service livré par le projet, sous son contrôle direct"],
             ["Référence (baseline)", "Valeur de l'indicateur avant le démarrage des activités"],
             ["Cible", "Valeur attendue de l'indicateur à une échéance déterminée"],
             ["Jalon", "Cible intermédiaire fixée pour une période donnée"],
             ["Hypothèse", "Condition externe nécessaire à la chaîne de résultats, hors du "
                           "contrôle du projet"],
             ["Risque", "Événement incertain dont la survenue affecterait l'atteinte des "
                        "résultats"],
             ["Désagrégation", "Ventilation d'un indicateur par catégorie : sexe, âge, milieu, "
                               "groupe cible, situation de handicap, niveau de vulnérabilité"],
             ["Règle d'agrégation", "Mode de consolidation des mesures d'une même période "
                                    "collectées sur plusieurs zones : somme, moyenne, dernière "
                                    "valeur ou maximum"],
             ["Indice d'équité de genre", "Part des femmes parmi les bénéficiaires et écart à la "
                                          "parité, exprimé en points de pourcentage"],
             ["Taux de couverture", "Rapport entre les bénéficiaires atteints sur une zone et la "
                                    "cible fixée pour cette zone"],
             ["Zone d'intervention", "Unité géographique de mise en œuvre et de consolidation des "
                                     "données, organisée en hiérarchie administrative"],
             ["Carte à symboles proportionnels", "Carte où la surface du symbole représente une "
                                                 "quantité et sa couleur une intensité relative"],
             ["Graticule", "Réseau des lignes de latitude et de longitude servant de repère "
                           "géographique sur une carte"],
             ["Chemin critique", "Séquence continue d'activités sans marge : tout retard sur "
                                 "l'une d'elles décale la fin du projet"],
             ["Marge totale", "Retard admissible sur une activité sans décaler l'achèvement du "
                              "projet"],
             ["Marge libre", "Retard admissible sans décaler l'activité suivante"],
             ["PERT", "Program Evaluation and Review Technique : représentation en réseau des "
                      "activités et de leurs antécédences"],
             ["WBS", "Work Breakdown Structure, organigramme des tâches : décomposition "
                     "hiérarchique du projet en lots de travail"],
             ["Lot de travail", "Élément élémentaire du WBS, doté d'un livrable, d'un "
                                "responsable, d'une durée et d'un coût"],
             ["RACI", "Matrice des responsabilités : qui réalise (R), approuve (A), est consulté "
                      "(C) et est informé (I)"],
             ["Indicateur de processus", "Indicateur mesurant la conduite de l'action (exécution, "
                                         "délais, participation) plutôt que le changement produit"],
             ["GAR", "Gestion axée sur les résultats : approche centrée sur l'atteinte de "
                     "résultats mesurables"],
             ["CAD", "Comité d'aide au développement de l'OCDE, dont les six critères — "
                     "pertinence, cohérence, efficacité, efficience, impact, durabilité — "
                     "constituent le cadre de référence de l'évaluation"],
             ["Contrefactuel", "Ce qui serait advenu en l'absence du projet. Il n'est jamais "
                               "observable : toute évaluation d'impact consiste à l'estimer"],
             ["Groupe de comparaison", "Population non exposée au projet, servant à estimer le "
                                       "contrefactuel"],
             ["Essai randomisé contrôlé", "Étude où l'affectation au traitement est tirée au "
                                          "sort, ce qui rend les groupes comparables en espérance "
                                          "sur toutes les caractéristiques"],
             ["Doubles différences", "Comparaison de l'évolution du groupe traité et de celle du "
                                     "groupe de comparaison, ce qui neutralise les différences de "
                                     "niveau initial constantes dans le temps"],
             ["Score de propension", "Probabilité estimée d'être traité compte tenu des "
                                     "caractéristiques observées ; l'appariement sur ce score "
                                     "constitue un groupe de comparaison plausible"],
             ["Hypothèse d'identification", "Condition sous laquelle l'écart mesuré s'interprète "
                                            "comme un effet causal. Elle doit être énoncée et "
                                            "discutée : c'est elle qui fonde la validité de "
                                            "l'étude"],
             ["Puissance statistique", "Probabilité de détecter un effet lorsqu'il existe "
                                       "réellement. Une étude sous-dimensionnée conclut à "
                                       "l'absence d'effet faute d'observations, non faute d'effet"],
             ["Effet de plan", "Facteur multipliant la taille d'échantillon requise lorsque "
                               "l'affectation porte sur des grappes plutôt que sur des individus, "
                               "égal à 1 + (m − 1) × ρ"],
             ["Corrélation intra-grappe", "Part de la variance de l'indicateur attribuable aux "
                                          "différences entre grappes ; plus elle est élevée, plus "
                                          "l'échantillon requis augmente"],
             ["SMART", "Critères de qualité d'un indicateur : spécifique, mesurable, atteignable, "
                       "pertinent, temporellement défini"],
             ["PTBA", "Plan de travail et budget annuel"],
             ["XLSForm", "Norme de description de questionnaires numériques sous forme de "
                         "classeur, utilisée par ODK et KoboToolbox"],
             ["Critères du CAD", "Pertinence, cohérence, efficacité, efficience, impact, "
                                 "durabilité"]],
            largeurs=[4, 12.5], taille=9)

    document.add_page_break()
    titre1(document, "Annexe 2 — Arborescence du code source")
    code(document, [
        "sepia-erp/",
        "├── app/",
        "│   ├── main.py                  Application FastAPI",
        "│   ├── config.py                Configuration et référentiels",
        "│   ├── database.py              Moteur SQLAlchemy",
        "│   ├── models.py                Vingt-quatre entités du modèle de données",
        "│   ├── security.py              Mots de passe, sessions, clés d'API",
        "│   ├── middleware.py            En-têtes, limitation de débit, erreurs",
        "│   ├── crud.py                  Fabrique de routeurs CRUD et contrôle d'accès",
        "│   ├── seed.py                  Compte initial et projets d'exemple",
        "│   ├── seed_sante.py            Second projet d'exemple, intégralement renseigné",
        "│   ├── routers/",
        "│   │   ├── auth.py              Authentification et comptes",
        "│   │   ├── projects.py          Projets, tableaux de bord, référentiels",
        "│   │   ├── entities.py          Entités métier",
        "│   │   ├── evaluations.py       Bénéficiaires, partenaires, CAD, impact",
        "│   │   ├── imports.py           Imports Excel, Word, XLSForm, Kobo",
        "│   │   ├── exports.py           Génération des livrables",
        "│   │   └── powerbi.py           Flux de business intelligence",
        "│   └── services/",
        "│       ├── analytics.py         Moteur de performance",
        "│       ├── planning.py          Chemin critique, PERT, courbe en S, WBS, RACI",
        "│       ├── evaluation.py        Notation CAD, bénéficiaires, partenaires, impact",
        "│       ├── portability.py       Export et import JSON (projet, portefeuille)",
        "│       ├── excel_export.py      Générateurs Excel",
        "│       ├── word_export.py       Générateurs Word",
        "│       ├── evaluation_export.py Livrables des modules d'évaluation",
        "│       ├── xlsform.py           Générateur XLSForm",
        "│       └── importer.py          Analyseurs Excel et Word",
        "├── static/",
        "│   ├── index.html               Structure de l'interface",
        "│   ├── css/app.css              Feuille de style responsive",
        "│   └── js/",
        "│       ├── core.js              État, API, composants d'interface",
        "│       ├── charts.js            Graphiques SVG",
        "│       ├── views.js             Vues de planification, de suivi et d'analyse",
        "│       ├── views-evaluation.js  Bénéficiaires, partenaires, CAD, impact",
        "│       └── app.js               Navigation et cycle de vie",
        "├── scripts/",
        "│   ├── generer_documentation.py Production du présent document",
        "│   ├── verifier_securite.py     Jeu de vérification des garde-fous",
        "│   └── reinitialiser_admin.py   Reprise de main sur le compte admin",
        "├── docs/                        Documentation Word",
        "├── requirements.txt             Dépendances Python",
        "├── render.yaml                  Description d'infrastructure Render",
        "└── README.md                    Notice technique",
    ])

    document.add_page_break()
    titre1(document, "Annexe 3 — Référentiels paramétrables")
    para(document,
         "Les listes de valeurs proposées dans les formulaires sont centralisées et modifiables "
         "dans le fichier de configuration de l'application.")
    tableau(document, ["Référentiel", "Valeurs proposées"],
            [["Niveaux du cadre logique", "Impact, Effet, Produit, Activité"],
             ["Fréquences de collecte", "Mensuelle, trimestrielle, semestrielle, annuelle, "
                                        "ponctuelle, mi-parcours, finale"],
             ["Types d'indicateur", "Quantitatif, qualitatif, composite, proxy"],
             ["Statuts de projet", "Identification, formulation, en cours, suspendu, clôturé"],
             ["Catégories de risque", "Politique et gouvernance, sécuritaire, financier et "
                                      "budgétaire, opérationnel, technique, environnemental et "
                                      "climatique, social et genre, sanitaire, institutionnel et "
                                      "capacités, réputationnel"],
             ["Statuts d'activité", "Planifiée, en cours, achevée, retardée, annulée"],
             ["Statuts de risque", "Ouvert, maîtrisé, clos, survenu"],
             ["Catégories budgétaires", "Personnel, équipements, formations, prestations, missions "
                                        "et déplacements, fonctionnement, investissements, "
                                        "communication, suivi-évaluation, imprévus"],
             ["Catégories de désagrégation et modalités",
              "Sexe (Femme, Homme) · Âge (moins de 18 ans, 18 à 35, 36 à 59, 60 et plus) · "
              "Milieu (urbain, rural) · Groupe cible (producteur, transformatrice, jeune, femme "
              "cheffe de ménage, personne en situation de handicap, personne déplacée, autre) · "
              "Situation de handicap · Niveau de vulnérabilité · Statut d'occupation"],
             ["Niveaux de zone", "Pays, région, préfecture, commune, canton, village, site"],
             ["Règles d'agrégation", "Somme, moyenne, dernière valeur, maximum"],
             ["Types de rapport périodique", "Trimestriel, semestriel, annuel"],
             ["Nature des indicateurs", "Résultat, processus"],
             ["Rôles RACI", "R (réalise), A (approuve), C (consulté), I (informé)"],
             ["Catégories de parties prenantes",
              "Interne, tutelle, partenaire d'exécution, prestataire, bailleur, bénéficiaire, "
              "collectivité"],
             ["Unités de mesure", "Nombre, pourcentage, ratio, score, indice, tonne, hectare, "
                                  "kilomètre, FCFA, USD, EUR, jour, mois, t/ha, kg, litre"],
             ["Types de formulaire", "Questionnaire, fiche de suivi, grille d'entretien, grille "
                                     "de focus group, fiche de présence, fiche d'observation"],
             ["Statuts d'hypothèse", "Non vérifiée, partiellement vérifiée, vérifiée, invalidée"]],
            largeurs=[4.5, 12], taille=9)

    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"— Fin du document —\n{APP_NAME} version {APP_VERSION} · édité le {date_fr}")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRIS

    pied_de_page(document, f"{APP_NAME} — Documentation fonctionnelle et technique — version "
                           f"{APP_VERSION}")

    os.makedirs(os.path.dirname(chemin_sortie), exist_ok=True)
    document.save(chemin_sortie)
    return chemin_sortie


if __name__ == "__main__":
    sortie = sys.argv[1] if len(sys.argv) > 1 else os.path.join(
        RACINE, "docs", "SEPIA_Documentation_plateforme.docx")
    chemin = construire(sortie)
    taille = os.path.getsize(chemin)
    print(f"Document généré : {chemin} ({taille // 1024} Ko)")
