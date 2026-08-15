"""Import de cadres logiques, indicateurs, activités, budgets et risques
depuis des classeurs Excel ou des documents Word."""
import re
import unicodedata
from datetime import date, datetime
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from openpyxl import load_workbook
from sqlalchemy.orm import Session

from ..config import CATEGORIES_DESAGREGATION, MODALITES_DESAGREGATION
from ..models import (Activity, Assumption, BudgetLine, Indicator, IndicatorActual,
                      IndicatorTarget, LogframeElement, Project, Risk, Zone)


# ---------------------------------------------------------------------------
# Normalisation
# ---------------------------------------------------------------------------
def normaliser(texte: Any) -> str:
    """Minuscule, sans accent ni ponctuation — pour apparier les intitulés de colonnes."""
    if texte is None:
        return ""
    texte = str(texte).strip().lower()
    texte = unicodedata.normalize("NFD", texte)
    texte = "".join(c for c in texte if unicodedata.category(c) != "Mn")
    return re.sub(r"[^a-z0-9]+", " ", texte).strip()


def _valeur_numerique(valeur: Any) -> Optional[float]:
    if valeur is None or valeur == "":
        return None
    if isinstance(valeur, (int, float)):
        return float(valeur)
    texte = str(valeur).strip().replace("%", "").replace(" ", "").replace(" ", "")
    texte = texte.replace(",", ".")
    texte = re.sub(r"[^0-9.\-]", "", texte)
    if texte in ("", "-", "."):
        return None
    try:
        return float(texte)
    except ValueError:
        return None


def _valeur_date(valeur: Any) -> Optional[date]:
    if valeur is None or valeur == "":
        return None
    if isinstance(valeur, datetime):
        return valeur.date()
    if isinstance(valeur, date):
        return valeur
    texte = str(valeur).strip()
    for motif in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%m/%d/%Y", "%Y/%m/%d", "%d.%m.%Y"):
        try:
            return datetime.strptime(texte[:10], motif).date()
        except ValueError:
            continue
    return None


def _booleen(valeur: Any) -> bool:
    return normaliser(valeur) in ("oui", "yes", "1", "true", "vrai", "x")


def _liste(valeur: Any) -> List[str]:
    if not valeur:
        return []
    return [p.strip() for p in re.split(r"[;,/|]", str(valeur)) if p.strip()]


NIVEAUX_CANONIQUES = {
    "impact": "IMPACT", "objectif global": "IMPACT", "og": "IMPACT", "but": "IMPACT",
    "finalite": "IMPACT", "goal": "IMPACT",
    "effet": "EFFET", "outcome": "EFFET", "objectif specifique": "EFFET", "os": "EFFET",
    "resultat": "EFFET", "objectif du projet": "EFFET",
    "produit": "PRODUIT", "output": "PRODUIT", "extrant": "PRODUIT", "resultat intermediaire": "PRODUIT",
    "activite": "ACTIVITE", "activity": "ACTIVITE", "action": "ACTIVITE", "intrant": "ACTIVITE",
}


def deviner_niveau(texte: Any) -> Optional[str]:
    n = normaliser(texte)
    if not n:
        return None
    if n.upper() in ("IMPACT", "EFFET", "PRODUIT", "ACTIVITE"):
        return n.upper()
    for cle, niveau in NIVEAUX_CANONIQUES.items():
        if n.startswith(cle) or cle in n:
            return niveau
    return None


# ---------------------------------------------------------------------------
# Appariement des colonnes
# ---------------------------------------------------------------------------
def _index_colonnes(entetes: List[Any], correspondances: Dict[str, List[str]]) -> Dict[str, int]:
    """Associe un champ logique à l'indice de colonne, par appariement souple."""
    normalisees = [normaliser(e) for e in entetes]
    resultat: Dict[str, int] = {}
    for champ, alias in correspondances.items():
        for index, entete in enumerate(normalisees):
            if not entete:
                continue
            if any(entete == normaliser(a) for a in alias):
                resultat[champ] = index
                break
        if champ in resultat:
            continue
        for index, entete in enumerate(normalisees):
            if entete and any(normaliser(a) in entete for a in alias):
                resultat[champ] = index
                break
    return resultat


def _cellule(ligne: Tuple, colonnes: Dict[str, int], champ: str) -> Any:
    index = colonnes.get(champ)
    if index is None or index >= len(ligne):
        return None
    valeur = ligne[index]
    if isinstance(valeur, str):
        valeur = valeur.strip()
    return valeur or None


# ---------------------------------------------------------------------------
# Import Excel
# ---------------------------------------------------------------------------
CORRESPONDANCES = {
    "cadre_logique": {
        "level": ["niveau", "niveau de resultat", "level"],
        "code": ["code", "reference", "num", "numero"],
        "parent": ["code parent", "parent", "rattachement", "resultat parent"],
        "statement": ["enonce du resultat", "enonce", "logique d intervention", "intitule",
                      "libelle", "resultat", "description"],
        "mov": ["sources de verification", "source de verification", "moyens de verification", "mov"],
        "assumptions": ["hypotheses", "hypotheses risques", "conditions critiques", "risques"],
        "responsible": ["responsable", "structure responsable"],
    },
    "indicateurs": {
        "code": ["code", "code indicateur", "reference"],
        "name": ["libelle de l indicateur", "indicateur", "libelle", "intitule", "nom"],
        "element": ["code resultat", "resultat", "rattachement", "code parent", "niveau resultat"],
        "level": ["niveau"],
        "indicator_type": ["type", "type d indicateur"],
        "unit": ["unite", "unite de mesure"],
        "formula": ["mode de calcul", "formule", "methode de calcul"],
        "disaggregation": ["desagregation", "ventilation"],
        "baseline_value": ["reference", "valeur de reference", "baseline", "situation de reference"],
        "baseline_date": ["date reference", "date de reference"],
        "target_value": ["cible finale", "cible", "valeur cible", "target"],
        "target_date": ["date cible", "echeance"],
        "direction": ["sens", "sens de progression"],
        "frequency": ["frequence", "periodicite"],
        "data_source": ["source de donnees", "source des donnees", "source"],
        "collection_method": ["methode de collecte", "mode de collecte", "outil de collecte"],
        "responsible": ["responsable"],
        "is_key": ["indicateur cle", "kpi", "cle"],
        "definition": ["definition", "definition operationnelle"],
        "cost_estimate": ["cout estime", "cout"],
    },
    "cibles": {
        "indicator": ["code indicateur", "indicateur", "code"],
        "period_label": ["periode", "libelle periode"],
        "year": ["annee"],
        "period_start": ["debut periode", "date debut"],
        "period_end": ["fin periode", "date fin"],
        "target_value": ["valeur cible", "cible", "valeur"],
    },
    "realisations": {
        "indicator": ["code indicateur", "indicateur", "code"],
        "period_label": ["periode", "libelle periode"],
        "year": ["annee"],
        "reference_date": ["date de reference", "date"],
        "value": ["valeur realisee", "realise", "valeur"],
        "source": ["source"],
        "collected_by": ["collecte par", "agent", "enqueteur"],
        "validation_status": ["statut", "validation"],
        "zone": ["zone", "zone d intervention", "localite", "site"],
        "activity": ["code activite", "activite"],
        "desagregation": ["desagregation", "ventilation", "detail desagrege"],
    },
    "zones": {
        "code": ["code"],
        "name": ["zone", "nom de la zone", "libelle", "nom", "intitule"],
        "level": ["niveau", "type de zone"],
        "parent": ["zone parente", "code parent", "parent", "rattachement"],
        "population": ["population"],
        "beneficiaries_target": ["cible beneficiaires", "beneficiaires cibles",
                                 "cible de beneficiaires", "beneficiaires"],
        "latitude": ["latitude"],
        "longitude": ["longitude"],
        "responsible": ["responsable"],
    },
    "activites": {
        "code": ["code", "code activite"],
        "name": ["libelle de l activite", "activite", "libelle", "intitule"],
        "element": ["code resultat", "resultat", "produit", "rattachement"],
        "responsible": ["responsable"],
        "partners": ["partenaires", "partenaire"],
        "location": ["lieu", "zone", "localisation"],
        "start_date": ["date debut", "debut"],
        "end_date": ["date fin", "fin"],
        "progress": ["avancement", "taux d avancement", "avancement %"],
        "status": ["statut", "etat"],
        "planned_cost": ["cout prevu", "budget", "cout"],
        "year": ["annee"],
        "milestone": ["jalon"],
        "deliverable": ["livrable", "produit attendu"],
    },
    "budget": {
        "code": ["code", "code ligne"],
        "label": ["libelle de la ligne", "libelle", "designation", "intitule", "ligne budgetaire"],
        "activity": ["code activite", "activite"],
        "category": ["categorie", "rubrique", "nature"],
        "unit": ["unite"],
        "quantity": ["quantite", "qte"],
        "unit_cost": ["cout unitaire", "prix unitaire", "pu"],
        "frequency_count": ["nombre", "nb", "frequence"],
        "q1": ["t1", "trimestre 1"], "q2": ["t2", "trimestre 2"],
        "q3": ["t3", "trimestre 3"], "q4": ["t4", "trimestre 4"],
        "funding_source": ["source de financement", "financement", "bailleur"],
        "year": ["annee", "exercice"],
        "committed": ["engage", "montant engage"],
        "disbursed": ["decaisse", "montant decaisse", "depense"],
    },
    "risques": {
        "code": ["code"],
        "category": ["categorie", "type de risque"],
        "title": ["risque identifie", "risque", "libelle", "intitule", "description"],
        "cause": ["cause", "causes"],
        "consequence": ["consequence", "effet", "impact potentiel"],
        "probability": ["probabilite", "probabilite 1 5", "p"],
        "impact": ["impact", "impact 1 5", "gravite", "i"],
        "mitigation": ["mesures d attenuation", "attenuation", "mitigation", "mesures"],
        "contingency": ["plan de contingence", "contingence", "plan b"],
        "owner": ["responsable", "porteur"],
        "status": ["statut", "etat"],
        "review_date": ["date de revue", "revue", "echeance"],
    },
    "hypotheses": {
        "code": ["code"],
        "level": ["niveau"],
        "statement": ["enonce de l hypothese", "hypothese", "enonce", "libelle"],
        "criticality": ["criticite", "importance"],
        "validation_status": ["statut de validation", "statut", "validation"],
        "verification_method": ["methode de verification", "verification", "moyen de verification"],
        "responsible": ["responsable"],
        "review_date": ["date de revue", "revue"],
    },
}

ALIAS_FEUILLES = {
    "cadre_logique": ["cadre logique", "cadre logique", "logframe", "matrice", "cadre de resultats",
                      "chaine de resultats", "resultats"],
    "indicateurs": ["indicateurs", "indicateur", "iov", "liste des indicateurs"],
    "cibles": ["cibles", "cible", "targets", "jalons"],
    "realisations": ["realisations", "realisation", "actuals", "donnees", "valeurs"],
    "activites": ["activites", "activite", "chronogramme", "plan d action", "planning"],
    "budget": ["budget", "ptba", "budget detaille", "lignes budgetaires", "cout"],
    "risques": ["risques", "risque", "registre des risques", "matrice des risques"],
    "hypotheses": ["hypotheses", "hypothese", "conditions critiques"],
    "zones": ["zones", "zone", "zones d intervention", "localites", "sites",
              "zone d intervention"],
}


# ---------------------------------------------------------------------------
# Désagrégation : lecture des colonnes « Catégorie - Modalité »
# ---------------------------------------------------------------------------
def _colonnes_desagregation(entetes: List[Any]) -> Dict[int, Tuple[str, str]]:
    """Repère les colonnes du type « Sexe - Femme » ou « Groupe cible : Jeune ».

    Retourne {indice de colonne: (catégorie, modalité)}. La catégorie est
    rapprochée du référentiel afin de conserver l'orthographe canonique.
    """
    canoniques = {normaliser(c): c for c in CATEGORIES_DESAGREGATION}
    resultat: Dict[int, Tuple[str, str]] = {}
    for index, entete in enumerate(entetes):
        if not isinstance(entete, str):
            continue
        for separateur in (" - ", " – ", " : ", ":", "|", " / "):
            if separateur not in entete:
                continue
            gauche, droite = entete.split(separateur, 1)
            categorie = canoniques.get(normaliser(gauche))
            if categorie and droite.strip():
                modalite = droite.strip()
                referentiel = MODALITES_DESAGREGATION.get(categorie, [])
                correspondance = next(
                    (m for m in referentiel if normaliser(m) == normaliser(modalite)), None)
                resultat[index] = (categorie, correspondance or modalite)
                break
    return resultat


def _lire_desagregation(ligne: Tuple, colonnes_desagregation: Dict[int, Tuple[str, str]],
                        texte_libre: Any) -> Dict[str, Dict[str, float]]:
    """Construit le dictionnaire de valeurs désagrégées d'une ligne.

    Deux écritures sont acceptées : une colonne par modalité, ou une colonne
    unique au format « Sexe:Femme=120;Sexe:Homme=95 ».
    """
    valeurs: Dict[str, Dict[str, float]] = {}
    for index, (categorie, modalite) in colonnes_desagregation.items():
        if index >= len(ligne):
            continue
        nombre = _valeur_numerique(ligne[index])
        if nombre is None:
            continue
        valeurs.setdefault(categorie, {})[modalite] = nombre
    if texte_libre:
        canoniques = {normaliser(c): c for c in CATEGORIES_DESAGREGATION}
        for morceau in re.split(r"[;\n]+", str(texte_libre)):
            correspondance = re.match(r"\s*([^:=]+)\s*:\s*([^=]+)\s*=\s*(.+)\s*$", morceau)
            if not correspondance:
                continue
            categorie = canoniques.get(normaliser(correspondance.group(1)),
                                       correspondance.group(1).strip())
            nombre = _valeur_numerique(correspondance.group(3))
            if nombre is not None:
                valeurs.setdefault(categorie, {})[correspondance.group(2).strip()] = nombre
    return valeurs


def _trouver_feuille(classeur, cle: str):
    alias = [normaliser(a) for a in ALIAS_FEUILLES[cle]]
    for nom in classeur.sheetnames:
        n = normaliser(nom)
        if n in alias:
            return classeur[nom]
    for nom in classeur.sheetnames:
        n = normaliser(nom)
        if any(a in n or n in a for a in alias if a):
            return classeur[nom]
    return None


def _lignes_utiles(feuille) -> Tuple[List[Any], List[Tuple]]:
    """Retourne (entêtes, lignes) en localisant la ligne d'en-tête réelle."""
    lignes = list(feuille.iter_rows(values_only=True))
    if not lignes:
        return [], []
    index_entete = 0
    meilleur_score = -1
    for index, ligne in enumerate(lignes[:8]):
        score = sum(1 for c in ligne if isinstance(c, str) and c.strip())
        if score > meilleur_score:
            meilleur_score, index_entete = score, index
    entetes = list(lignes[index_entete])
    corps = [l for l in lignes[index_entete + 1:] if any(c not in (None, "") for c in l)]
    return entetes, corps


def importer_excel(db: Session, contenu: bytes, project: Project,
                   remplacer: bool = False) -> Dict[str, Any]:
    """Importe un classeur au format du modèle SEPIA (ou proche) dans un projet."""
    classeur = load_workbook(BytesIO(contenu), data_only=True)
    rapport: Dict[str, Any] = {"projet": project.code, "feuilles_traitees": [],
                               "crees": {}, "ignores": [], "avertissements": []}

    if remplacer:
        for modele in (IndicatorActual, IndicatorTarget):
            db.query(modele).filter(modele.indicator_id.in_(
                db.query(Indicator.id).filter(Indicator.project_id == project.id))).delete(
                synchronize_session=False)
        for modele in (BudgetLine, Activity, Indicator, Assumption, Risk, LogframeElement, Zone):
            db.query(modele).filter(modele.project_id == project.id).delete(synchronize_session=False)
        db.flush()

    def compter(cle: str, nombre: int):
        rapport["crees"][cle] = rapport["crees"].get(cle, 0) + nombre

    # --- 1. Cadre logique -------------------------------------------------
    elements_par_code: Dict[str, LogframeElement] = {
        e.code: e for e in db.query(LogframeElement).filter(
            LogframeElement.project_id == project.id).all() if e.code}
    feuille = _trouver_feuille(classeur, "cadre_logique")
    if feuille is not None:
        entetes, corps = _lignes_utiles(feuille)
        colonnes = _index_colonnes(entetes, CORRESPONDANCES["cadre_logique"])
        rapport["feuilles_traitees"].append(feuille.title)
        en_attente: List[Tuple[LogframeElement, str]] = []
        nombre = 0
        for position, ligne in enumerate(corps):
            enonce = _cellule(ligne, colonnes, "statement")
            niveau = deviner_niveau(_cellule(ligne, colonnes, "level"))
            if not enonce:
                continue
            if not niveau:
                niveau = deviner_niveau(_cellule(ligne, colonnes, "code")) or "PRODUIT"
            code = _cellule(ligne, colonnes, "code")
            element = elements_par_code.get(str(code)) if code else None
            if element is None:
                element = LogframeElement(project_id=project.id)
                db.add(element)
                nombre += 1
            element.level = niveau
            element.code = str(code) if code else None
            element.statement = str(enonce)
            element.means_of_verification = _cellule(ligne, colonnes, "mov")
            element.assumptions = _cellule(ligne, colonnes, "assumptions")
            element.responsible = _cellule(ligne, colonnes, "responsible")
            element.order_index = position
            if code:
                elements_par_code[str(code)] = element
            parent = _cellule(ligne, colonnes, "parent")
            if parent:
                en_attente.append((element, str(parent)))
        db.flush()
        for element, code_parent in en_attente:
            parent = elements_par_code.get(code_parent)
            if parent is not None and parent.id != element.id:
                element.parent_id = parent.id
            else:
                rapport["avertissements"].append(
                    f"Cadre logique : parent « {code_parent} » introuvable pour « "
                    f"{(element.code or element.statement)[:40]} ».")
        db.flush()
        compter("resultats", nombre)

    # --- 2. Indicateurs ---------------------------------------------------
    indicateurs_par_code: Dict[str, Indicator] = {
        i.code: i for i in db.query(Indicator).filter(Indicator.project_id == project.id).all()
        if i.code}
    feuille = _trouver_feuille(classeur, "indicateurs")
    if feuille is not None:
        entetes, corps = _lignes_utiles(feuille)
        colonnes = _index_colonnes(entetes, CORRESPONDANCES["indicateurs"])
        rapport["feuilles_traitees"].append(feuille.title)
        nombre = 0
        for ligne in corps:
            libelle = _cellule(ligne, colonnes, "name")
            if not libelle:
                continue
            code = _cellule(ligne, colonnes, "code")
            indicateur = indicateurs_par_code.get(str(code)) if code else None
            if indicateur is None:
                indicateur = Indicator(project_id=project.id)
                db.add(indicateur)
                nombre += 1
            indicateur.code = str(code) if code else None
            indicateur.name = str(libelle)
            code_element = _cellule(ligne, colonnes, "element")
            element = elements_par_code.get(str(code_element)) if code_element else None
            if element is not None:
                indicateur.element_id = element.id
                indicateur.level = element.level
            niveau = deviner_niveau(_cellule(ligne, colonnes, "level"))
            if niveau:
                indicateur.level = niveau
            indicateur.definition = _cellule(ligne, colonnes, "definition")
            indicateur.indicator_type = _cellule(ligne, colonnes, "indicator_type") or "Quantitatif"
            indicateur.unit = _cellule(ligne, colonnes, "unit")
            indicateur.formula = _cellule(ligne, colonnes, "formula")
            indicateur.disaggregation = _liste(_cellule(ligne, colonnes, "disaggregation"))
            indicateur.baseline_value = _valeur_numerique(_cellule(ligne, colonnes, "baseline_value"))
            indicateur.baseline_date = _valeur_date(_cellule(ligne, colonnes, "baseline_date"))
            indicateur.target_value = _valeur_numerique(_cellule(ligne, colonnes, "target_value"))
            indicateur.target_date = _valeur_date(_cellule(ligne, colonnes, "target_date"))
            indicateur.direction = _cellule(ligne, colonnes, "direction") or "Croissant"
            indicateur.frequency = _cellule(ligne, colonnes, "frequency") or "Trimestrielle"
            indicateur.data_source = _cellule(ligne, colonnes, "data_source")
            indicateur.collection_method = _cellule(ligne, colonnes, "collection_method")
            indicateur.responsible = _cellule(ligne, colonnes, "responsible")
            indicateur.cost_estimate = _valeur_numerique(_cellule(ligne, colonnes, "cost_estimate"))
            indicateur.is_key = _booleen(_cellule(ligne, colonnes, "is_key"))
            if code:
                indicateurs_par_code[str(code)] = indicateur
        db.flush()
        compter("indicateurs", nombre)

    # --- 2 bis. Zones d'intervention -------------------------------------
    zones_par_code: Dict[str, Zone] = {}
    zones_par_nom: Dict[str, Zone] = {}
    for zone in db.query(Zone).filter(Zone.project_id == project.id).all():
        if zone.code:
            zones_par_code[str(zone.code)] = zone
        zones_par_nom[normaliser(zone.name)] = zone
    feuille = _trouver_feuille(classeur, "zones")
    if feuille is not None:
        entetes, corps = _lignes_utiles(feuille)
        colonnes = _index_colonnes(entetes, CORRESPONDANCES["zones"])
        rapport["feuilles_traitees"].append(feuille.title)
        nombre = 0
        en_attente_zones: List[Tuple[Zone, str]] = []
        for position, ligne in enumerate(corps):
            nom = _cellule(ligne, colonnes, "name")
            if not nom:
                continue
            code = _cellule(ligne, colonnes, "code")
            zone = zones_par_code.get(str(code)) if code else zones_par_nom.get(normaliser(nom))
            if zone is None:
                zone = Zone(project_id=project.id)
                db.add(zone)
                nombre += 1
            zone.code = str(code) if code else None
            zone.name = str(nom)
            zone.level = _cellule(ligne, colonnes, "level") or "Région"
            population = _valeur_numerique(_cellule(ligne, colonnes, "population"))
            zone.population = int(population) if population else None
            cible = _valeur_numerique(_cellule(ligne, colonnes, "beneficiaries_target"))
            zone.beneficiaries_target = int(cible) if cible else None
            zone.latitude = _valeur_numerique(_cellule(ligne, colonnes, "latitude"))
            zone.longitude = _valeur_numerique(_cellule(ligne, colonnes, "longitude"))
            zone.responsible = _cellule(ligne, colonnes, "responsible")
            zone.order_index = position
            db.flush()
            if code:
                zones_par_code[str(code)] = zone
            zones_par_nom[normaliser(nom)] = zone
            parent = _cellule(ligne, colonnes, "parent")
            if parent:
                en_attente_zones.append((zone, str(parent)))
        for zone, code_parent in en_attente_zones:
            parent = zones_par_code.get(code_parent) or zones_par_nom.get(normaliser(code_parent))
            if parent is not None and parent.id != zone.id:
                zone.parent_id = parent.id
            else:
                rapport["avertissements"].append(
                    f"Zones : zone parente « {code_parent} » introuvable pour « {zone.name} ».")
        db.flush()
        compter("zones", nombre)

    # --- 2 ter. Activités -------------------------------------------------
    # Importées avant les réalisations : celles-ci peuvent référencer l'activité
    # qui les a produites, ce qui permet la consolidation par activité.
    activites_par_code: Dict[str, Activity] = {
        a.code: a for a in db.query(Activity).filter(Activity.project_id == project.id).all()
        if a.code}
    feuille = _trouver_feuille(classeur, "activites")
    if feuille is not None:
        entetes, corps = _lignes_utiles(feuille)
        colonnes = _index_colonnes(entetes, CORRESPONDANCES["activites"])
        rapport["feuilles_traitees"].append(feuille.title)
        nombre = 0
        for position, ligne in enumerate(corps):
            libelle = _cellule(ligne, colonnes, "name")
            if not libelle:
                continue
            code = _cellule(ligne, colonnes, "code")
            activite = activites_par_code.get(str(code)) if code else None
            if activite is None:
                activite = Activity(project_id=project.id)
                db.add(activite)
                nombre += 1
            activite.code = str(code) if code else None
            activite.name = str(libelle)
            code_element = _cellule(ligne, colonnes, "element")
            element = elements_par_code.get(str(code_element)) if code_element else None
            if element is not None:
                activite.element_id = element.id
            activite.responsible = _cellule(ligne, colonnes, "responsible")
            activite.partners = _cellule(ligne, colonnes, "partners")
            activite.location = _cellule(ligne, colonnes, "location")
            activite.start_date = _valeur_date(_cellule(ligne, colonnes, "start_date"))
            activite.end_date = _valeur_date(_cellule(ligne, colonnes, "end_date"))
            activite.progress = _valeur_numerique(_cellule(ligne, colonnes, "progress")) or 0
            activite.status = _cellule(ligne, colonnes, "status") or "Planifiée"
            activite.planned_cost = _valeur_numerique(_cellule(ligne, colonnes, "planned_cost")) or 0
            annee = _valeur_numerique(_cellule(ligne, colonnes, "year"))
            activite.year = int(annee) if annee else (
                activite.start_date.year if activite.start_date else None)
            activite.milestone = _booleen(_cellule(ligne, colonnes, "milestone"))
            activite.deliverable = _cellule(ligne, colonnes, "deliverable")
            activite.order_index = position
            if code:
                activites_par_code[str(code)] = activite
        db.flush()
        compter("activites", nombre)

    # --- 3. Cibles et réalisations ---------------------------------------
    for cle, modele, champ_valeur in (("cibles", IndicatorTarget, "target_value"),
                                      ("realisations", IndicatorActual, "value")):
        feuille = _trouver_feuille(classeur, cle)
        if feuille is None:
            continue
        entetes, corps = _lignes_utiles(feuille)
        colonnes = _index_colonnes(entetes, CORRESPONDANCES[cle])
        colonnes_desagregation = _colonnes_desagregation(entetes) if cle == "realisations" else {}
        rapport["feuilles_traitees"].append(feuille.title)
        nombre = 0
        for ligne in corps:
            code = _cellule(ligne, colonnes, "indicator")
            indicateur = indicateurs_par_code.get(str(code)) if code else None
            valeur = _valeur_numerique(_cellule(ligne, colonnes, champ_valeur))
            periode = _cellule(ligne, colonnes, "period_label")
            if indicateur is None or valeur is None or not periode:
                if code and indicateur is None:
                    rapport["avertissements"].append(
                        f"{cle.capitalize()} : indicateur « {code} » introuvable.")
                continue
            objet = modele(indicator_id=indicateur.id, period_label=str(periode))
            annee = _valeur_numerique(_cellule(ligne, colonnes, "year"))
            objet.year = int(annee) if annee else None
            if cle == "cibles":
                objet.target_value = valeur
                objet.period_start = _valeur_date(_cellule(ligne, colonnes, "period_start"))
                objet.period_end = _valeur_date(_cellule(ligne, colonnes, "period_end"))
            else:
                objet.value = valeur
                objet.reference_date = _valeur_date(_cellule(ligne, colonnes, "reference_date"))
                objet.source = _cellule(ligne, colonnes, "source")
                objet.collected_by = _cellule(ligne, colonnes, "collected_by")
                objet.validation_status = _cellule(ligne, colonnes, "validation_status") or "Validé"
                nom_zone = _cellule(ligne, colonnes, "zone")
                if nom_zone:
                    zone = zones_par_code.get(str(nom_zone)) or \
                        zones_par_nom.get(normaliser(nom_zone))
                    if zone is not None:
                        objet.zone_id = zone.id
                    else:
                        rapport["avertissements"].append(
                            f"Réalisations : zone « {nom_zone} » introuvable ; la mesure est "
                            f"enregistrée sans localisation.")
                code_activite = _cellule(ligne, colonnes, "activity")
                if code_activite:
                    activite = db.query(Activity).filter(
                        Activity.project_id == project.id,
                        Activity.code == str(code_activite)).first()
                    if activite is not None:
                        objet.activity_id = activite.id
                desagregation = _lire_desagregation(
                    ligne, colonnes_desagregation, _cellule(ligne, colonnes, "desagregation"))
                if desagregation:
                    objet.disaggregated_values = desagregation
            db.add(objet)
            nombre += 1
        db.flush()
        compter(cle, nombre)

    # --- 4. Budget --------------------------------------------------------
    feuille = _trouver_feuille(classeur, "budget")
    if feuille is not None:
        entetes, corps = _lignes_utiles(feuille)
        colonnes = _index_colonnes(entetes, CORRESPONDANCES["budget"])
        rapport["feuilles_traitees"].append(feuille.title)
        nombre = 0
        for ligne in corps:
            libelle = _cellule(ligne, colonnes, "label")
            if not libelle:
                continue
            budget = BudgetLine(project_id=project.id, label=str(libelle))
            budget.code = _cellule(ligne, colonnes, "code")
            code_activite = _cellule(ligne, colonnes, "activity")
            activite = activites_par_code.get(str(code_activite)) if code_activite else None
            if activite is not None:
                budget.activity_id = activite.id
            budget.category = _cellule(ligne, colonnes, "category")
            budget.unit = _cellule(ligne, colonnes, "unit")
            budget.quantity = _valeur_numerique(_cellule(ligne, colonnes, "quantity")) or 1
            budget.unit_cost = _valeur_numerique(_cellule(ligne, colonnes, "unit_cost")) or 0
            budget.frequency_count = _valeur_numerique(_cellule(ligne, colonnes, "frequency_count")) or 1
            for trimestre in ("q1", "q2", "q3", "q4"):
                setattr(budget, trimestre, _valeur_numerique(_cellule(ligne, colonnes, trimestre)) or 0)
            budget.funding_source = _cellule(ligne, colonnes, "funding_source")
            annee = _valeur_numerique(_cellule(ligne, colonnes, "year"))
            budget.year = int(annee) if annee else None
            budget.committed = _valeur_numerique(_cellule(ligne, colonnes, "committed")) or 0
            budget.disbursed = _valeur_numerique(_cellule(ligne, colonnes, "disbursed")) or 0
            db.add(budget)
            nombre += 1
        db.flush()
        compter("lignes_budgetaires", nombre)

    # --- 6. Risques -------------------------------------------------------
    feuille = _trouver_feuille(classeur, "risques")
    if feuille is not None:
        entetes, corps = _lignes_utiles(feuille)
        colonnes = _index_colonnes(entetes, CORRESPONDANCES["risques"])
        rapport["feuilles_traitees"].append(feuille.title)
        nombre = 0
        for ligne in corps:
            titre = _cellule(ligne, colonnes, "title")
            if not titre:
                continue
            risque = Risk(project_id=project.id, title=str(titre))
            risque.code = _cellule(ligne, colonnes, "code")
            risque.category = _cellule(ligne, colonnes, "category")
            risque.cause = _cellule(ligne, colonnes, "cause")
            risque.consequence = _cellule(ligne, colonnes, "consequence")
            risque.probability = int(_valeur_numerique(_cellule(ligne, colonnes, "probability")) or 3)
            risque.impact = int(_valeur_numerique(_cellule(ligne, colonnes, "impact")) or 3)
            risque.mitigation = _cellule(ligne, colonnes, "mitigation")
            risque.contingency = _cellule(ligne, colonnes, "contingency")
            risque.owner = _cellule(ligne, colonnes, "owner")
            risque.status = _cellule(ligne, colonnes, "status") or "Ouvert"
            risque.review_date = _valeur_date(_cellule(ligne, colonnes, "review_date"))
            db.add(risque)
            nombre += 1
        db.flush()
        compter("risques", nombre)

    # --- 7. Hypothèses ----------------------------------------------------
    feuille = _trouver_feuille(classeur, "hypotheses")
    if feuille is not None:
        entetes, corps = _lignes_utiles(feuille)
        colonnes = _index_colonnes(entetes, CORRESPONDANCES["hypotheses"])
        rapport["feuilles_traitees"].append(feuille.title)
        nombre = 0
        for ligne in corps:
            enonce = _cellule(ligne, colonnes, "statement")
            if not enonce:
                continue
            hypothese = Assumption(project_id=project.id, statement=str(enonce))
            hypothese.code = _cellule(ligne, colonnes, "code")
            hypothese.level = deviner_niveau(_cellule(ligne, colonnes, "level"))
            hypothese.criticality = _cellule(ligne, colonnes, "criticality") or "Moyenne"
            hypothese.validation_status = _cellule(ligne, colonnes, "validation_status") or "Non vérifiée"
            hypothese.verification_method = _cellule(ligne, colonnes, "verification_method")
            hypothese.responsible = _cellule(ligne, colonnes, "responsible")
            hypothese.review_date = _valeur_date(_cellule(ligne, colonnes, "review_date"))
            db.add(hypothese)
            nombre += 1
        db.flush()
        compter("hypotheses", nombre)

    db.commit()
    if not rapport["feuilles_traitees"]:
        rapport["avertissements"].append(
            "Aucune feuille reconnue. Vérifiez les noms d'onglets (Cadre logique, Indicateurs, "
            "Activités, Budget, Risques…) ou utilisez le modèle d'import fourni.")
    return rapport


# ---------------------------------------------------------------------------
# Import Word
# ---------------------------------------------------------------------------
MOTS_CLES_LOGFRAME = ["logique d intervention", "indicateur", "source de verification",
                      "moyens de verification", "hypothese", "niveau", "resultat"]


def _table_en_matrice(table) -> List[List[str]]:
    return [[cellule.text.strip() for cellule in ligne.cells] for ligne in table.rows]


def analyser_word(contenu: bytes) -> Dict[str, Any]:
    """Extrait les tableaux d'un document Word et repère ceux qui ressemblent
    à un cadre logique, à une liste d'indicateurs ou à un registre de risques."""
    document = Document(BytesIO(contenu))
    resultat: Dict[str, Any] = {"tableaux": [], "titre_detecte": None, "paragraphes": 0}
    for paragraphe in document.paragraphs:
        if paragraphe.text.strip():
            resultat["paragraphes"] += 1
            if resultat["titre_detecte"] is None and paragraphe.style.name.startswith("Title"):
                resultat["titre_detecte"] = paragraphe.text.strip()
    for index, table in enumerate(document.tables):
        matrice = _table_en_matrice(table)
        if len(matrice) < 2:
            continue
        entetes = [normaliser(c) for c in matrice[0]]
        score = sum(1 for e in entetes for mot in MOTS_CLES_LOGFRAME if mot in e)
        nature = "inconnu"
        if any("logique d intervention" in e for e in entetes) or score >= 3:
            nature = "cadre_logique"
        elif any("risque" in e for e in entetes):
            nature = "risques"
        elif any("indicateur" in e for e in entetes):
            nature = "indicateurs"
        elif any("activite" in e for e in entetes):
            nature = "activites"
        resultat["tableaux"].append({
            "index": index, "nature": nature, "lignes": len(matrice),
            "colonnes": len(matrice[0]), "entetes": matrice[0],
            "apercu": matrice[1:4],
        })
    resultat["nb_tableaux"] = len(document.tables)
    return resultat


def importer_word(db: Session, contenu: bytes, project: Project,
                  index_tableau: Optional[int] = None) -> Dict[str, Any]:
    """Importe un cadre logique depuis un tableau d'un document Word.

    Deux structures sont reconnues :
      * matrice classique à 4 colonnes (logique d'intervention / IOV / sources / hypothèses),
        les niveaux étant détectés à partir des libellés de la première colonne ;
      * tableau structuré comportant une colonne « Niveau » et une colonne « Code ».
    """
    document = Document(BytesIO(contenu))
    rapport: Dict[str, Any] = {"resultats_crees": 0, "indicateurs_crees": 0,
                               "risques_crees": 0, "avertissements": []}
    tables = document.tables
    if not tables:
        rapport["avertissements"].append("Le document ne contient aucun tableau exploitable.")
        return rapport

    indices = [index_tableau] if index_tableau is not None else range(len(tables))
    ordre_niveau = {"IMPACT": 0, "EFFET": 1, "PRODUIT": 2, "ACTIVITE": 3}
    dernier_par_niveau: Dict[str, LogframeElement] = {}
    compteurs = {"IMPACT": 0, "EFFET": 0, "PRODUIT": 0, "ACTIVITE": 0}
    position = 0

    for index in indices:
        if index >= len(tables):
            continue
        matrice = _table_en_matrice(tables[index])
        if len(matrice) < 2:
            continue
        entetes = [normaliser(c) for c in matrice[0]]
        colonnes = _index_colonnes(matrice[0], CORRESPONDANCES["cadre_logique"])
        est_cadre = ("statement" in colonnes and
                     (any("logique d intervention" in e for e in entetes) or
                      "level" in colonnes or
                      any("indicateur" in e for e in entetes)))
        est_risque = any("risque" in e for e in entetes) and any(
            "probabilite" in e or "impact" in e for e in entetes)

        if est_risque:
            colonnes_risque = _index_colonnes(matrice[0], CORRESPONDANCES["risques"])
            for ligne in matrice[1:]:
                titre = _cellule(tuple(ligne), colonnes_risque, "title")
                if not titre:
                    continue
                db.add(Risk(
                    project_id=project.id, title=titre,
                    code=_cellule(tuple(ligne), colonnes_risque, "code"),
                    category=_cellule(tuple(ligne), colonnes_risque, "category"),
                    cause=_cellule(tuple(ligne), colonnes_risque, "cause"),
                    consequence=_cellule(tuple(ligne), colonnes_risque, "consequence"),
                    probability=int(_valeur_numerique(
                        _cellule(tuple(ligne), colonnes_risque, "probability")) or 3),
                    impact=int(_valeur_numerique(
                        _cellule(tuple(ligne), colonnes_risque, "impact")) or 3),
                    mitigation=_cellule(tuple(ligne), colonnes_risque, "mitigation"),
                    owner=_cellule(tuple(ligne), colonnes_risque, "owner"),
                ))
                rapport["risques_crees"] += 1
            continue

        if not est_cadre:
            continue

        index_iov = next((i for i, e in enumerate(entetes) if "indicateur" in e), None)
        for ligne in matrice[1:]:
            cellules = tuple(ligne)
            enonce = _cellule(cellules, colonnes, "statement")
            if not enonce or len(enonce) < 4:
                continue
            niveau = deviner_niveau(_cellule(cellules, colonnes, "level")) or deviner_niveau(enonce)
            # Ligne de titre de niveau (une seule cellule remplie) : sert de contexte
            cellules_remplies = [c for c in ligne if c.strip()]
            if niveau and len(set(cellules_remplies)) <= 1:
                dernier_par_niveau["_courant"] = niveau
                continue
            if not niveau:
                niveau = dernier_par_niveau.get("_courant") or "PRODUIT"
            compteurs[niveau] = compteurs.get(niveau, 0) + 1
            code = _cellule(cellules, colonnes, "code") or \
                {"IMPACT": "OG", "EFFET": "OS", "PRODUIT": "P", "ACTIVITE": "A"}[niveau] + \
                str(compteurs[niveau])
            element = LogframeElement(
                project_id=project.id, level=niveau, code=str(code)[:30],
                statement=enonce,
                means_of_verification=_cellule(cellules, colonnes, "mov"),
                assumptions=_cellule(cellules, colonnes, "assumptions"),
                order_index=position,
            )
            position += 1
            # Rattachement au dernier élément de niveau supérieur
            for niveau_parent in ("PRODUIT", "EFFET", "IMPACT"):
                if ordre_niveau[niveau_parent] < ordre_niveau[niveau] and \
                        niveau_parent in dernier_par_niveau:
                    element.parent = dernier_par_niveau[niveau_parent]
                    break
            db.add(element)
            db.flush()
            dernier_par_niveau[niveau] = element
            rapport["resultats_crees"] += 1

            # Extraction des indicateurs contenus dans la cellule « IOV »
            if index_iov is not None and index_iov < len(ligne):
                for morceau in re.split(r"[\n•;]+", ligne[index_iov]):
                    libelle = morceau.strip(" -–\t")
                    if len(libelle) < 6:
                        continue
                    correspondance = re.match(r"^([A-Z0-9][A-Z0-9\.\-]{0,12})[\s\):-]+(.+)$", libelle)
                    code_ind, texte = (correspondance.group(1), correspondance.group(2)) \
                        if correspondance else (None, libelle)
                    db.add(Indicator(project_id=project.id, element_id=element.id, level=niveau,
                                     code=code_ind, name=texte[:500]))
                    rapport["indicateurs_crees"] += 1
    db.commit()
    if not any((rapport["resultats_crees"], rapport["indicateurs_crees"], rapport["risques_crees"])):
        rapport["avertissements"].append(
            "Aucun tableau n'a pu être interprété comme un cadre logique. Utilisez l'aperçu "
            "d'analyse pour sélectionner explicitement le tableau à importer.")
    return rapport
