"""Génération des livrables Word : cadre logique, cadre de rendement, plan et manuel
de suivi-évaluation, registre des risques, fiches d'indicateurs, questionnaires."""
from datetime import date
from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from sqlalchemy.orm import Session

from ..config import APP_LONG_NAME, APP_NAME, LIBELLES_NIVEAUX
from ..models import (Activity, Assumption, BudgetLine, Form, Indicator, LogframeElement, Project, Risk)
from . import analytics
from .excel_export import _elements_tries

BLEU = RGBColor(0x1F, 0x4E, 0x79)
GRIS = RGBColor(0x55, 0x55, 0x55)
COULEURS_NIVEAU = {"IMPACT": "1F4E79", "EFFET": "2E75B6", "PRODUIT": "5B9BD5", "ACTIVITE": "9DC3E6"}
COULEURS_SEVERITE = {"Critique": "D93025", "Élevé": "EA8600", "Modéré": "F9A825", "Faible": "0F9D58"}


# ---------------------------------------------------------------------------
# Utilitaires de mise en forme
# ---------------------------------------------------------------------------
def _ombrer(cellule, couleur_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), couleur_hex)
    cellule._tc.get_or_add_tcPr().append(shd)


def _texte_cellule(cellule, texte: str, *, gras: bool = False, taille: int = 9,
                   blanc: bool = False, centre: bool = False) -> None:
    cellule.text = ""
    paragraphe = cellule.paragraphs[0]
    if centre:
        paragraphe.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraphe.add_run("" if texte is None else str(texte))
    run.bold = gras
    run.font.size = Pt(taille)
    if blanc:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _paysage(section) -> None:
    largeur, hauteur = section.page_width, section.page_height
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = max(largeur, hauteur)
    section.page_height = min(largeur, hauteur)
    section.left_margin = section.right_margin = Cm(1.5)
    section.top_margin = section.bottom_margin = Cm(1.5)


def _tableau(document, entetes: List[str], largeurs: Optional[List[float]] = None,
             couleur_entete: str = "1F4E79"):
    table = document.add_table(rows=1, cols=len(entetes))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, titre in enumerate(entetes):
        cellule = table.rows[0].cells[index]
        _ombrer(cellule, couleur_entete)
        _texte_cellule(cellule, titre, gras=True, blanc=True, centre=True, taille=9)
        if largeurs:
            cellule.width = Cm(largeurs[index])
    return table


def _en_tete_document(document, projet: Project, titre: str, sous_titre: str = "") -> None:
    titre_par = document.add_paragraph()
    titre_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titre_par.add_run(titre)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = BLEU
    identite = document.add_paragraph()
    identite.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = identite.add_run(f"{projet.code} — {projet.title}")
    run.bold = True
    run.font.size = Pt(12)
    if sous_titre:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(sous_titre)
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = GRIS
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Bailleur : {projet.donor or 'N/A'} | Agence d'exécution : "
                    f"{projet.executing_agency or 'N/A'} | Période : "
                    f"{_date_fr(projet.start_date)} → {_date_fr(projet.end_date)}\n"
                    f"Document généré le {date.today().strftime('%d/%m/%Y')} par {APP_NAME} — {APP_LONG_NAME}")
    run.font.size = Pt(8)
    run.font.color.rgb = GRIS


def _date_fr(valeur) -> str:
    return valeur.strftime("%d/%m/%Y") if valeur else "N/A"


def _pied_de_page(document, texte: str) -> None:
    for section in document.sections:
        paragraphe = section.footer.paragraphs[0]
        paragraphe.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraphe.text = ""
        run = paragraphe.add_run(texte)
        run.font.size = Pt(7)
        run.font.color.rgb = GRIS


def _sauver(document) -> BytesIO:
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _nombre(valeur, defaut="—") -> str:
    if valeur is None:
        return defaut
    if isinstance(valeur, float) and valeur.is_integer():
        return f"{int(valeur):,}".replace(",", " ")
    if isinstance(valeur, (int, float)):
        return f"{valeur:,.2f}".replace(",", " ").replace(".", ",")
    return str(valeur)


# ---------------------------------------------------------------------------
# 1. Cadre logique
# ---------------------------------------------------------------------------
def cadre_logique_docx(db: Session, project: Project) -> BytesIO:
    document = Document()
    _paysage(document.sections[0])
    _en_tete_document(document, project, "CADRE LOGIQUE",
                      "Matrice de planification du projet (Logical Framework Approach)")
    document.add_paragraph()

    table = _tableau(document, ["Logique d'intervention", "Indicateurs objectivement vérifiables",
                                "Sources de vérification", "Hypothèses / Risques"],
                     largeurs=[7.5, 8.5, 5.5, 5.5])
    indicateurs_par_element: Dict[Any, List[Indicator]] = {}
    for ind in db.query(Indicator).filter(Indicator.project_id == project.id).all():
        indicateurs_par_element.setdefault(ind.element_id, []).append(ind)

    niveau_courant = None
    for element in _elements_tries(db, project.id):
        if element.level != niveau_courant:
            niveau_courant = element.level
            ligne_titre = table.add_row()
            fusion = ligne_titre.cells[0].merge(ligne_titre.cells[3])
            _ombrer(fusion, COULEURS_NIVEAU.get(element.level, "BFBFBF"))
            _texte_cellule(fusion, LIBELLES_NIVEAUX.get(element.level, element.level),
                           gras=True, blanc=True, taille=10)
        ligne = table.add_row()
        _texte_cellule(ligne.cells[0], f"{element.code or ''} — {element.statement}"
                       if element.code else element.statement, gras=True)
        inds = indicateurs_par_element.get(element.id, [])
        texte_iov = "\n".join(
            f"• {i.code or ''} {i.name}"
            f"\n   Réf. : {_nombre(i.baseline_value)} ({_date_fr(i.baseline_date)}) — "
            f"Cible : {_nombre(i.target_value)} ({_date_fr(i.target_date)}) — {i.unit or ''}"
            for i in inds) or "—"
        _texte_cellule(ligne.cells[1], texte_iov)
        sources = element.means_of_verification or "\n".join(
            f"• {i.data_source}" for i in inds if i.data_source) or "—"
        _texte_cellule(ligne.cells[2], sources)
        _texte_cellule(ligne.cells[3], element.assumptions or "—")

    hypotheses = db.query(Assumption).filter(Assumption.project_id == project.id).all()
    if hypotheses:
        document.add_page_break()
        titre = document.add_heading("Registre des hypothèses critiques", level=1)
        titre.runs[0].font.color.rgb = BLEU
        table_h = _tableau(document, ["Code", "Niveau", "Énoncé", "Criticité", "Statut",
                                      "Vérification", "Responsable"],
                           largeurs=[1.5, 2.2, 9, 2, 2.6, 5.5, 3.2])
        for h in hypotheses:
            ligne = table_h.add_row()
            for index, valeur in enumerate([h.code, h.level, h.statement, h.criticality,
                                            h.validation_status, h.verification_method, h.responsible]):
                _texte_cellule(ligne.cells[index], valeur or "—")
    _pied_de_page(document, f"{project.code} — Cadre logique — Généré par {APP_NAME}")
    return _sauver(document)


# ---------------------------------------------------------------------------
# 2. Cadre de rendement
# ---------------------------------------------------------------------------
def cadre_rendement_docx(db: Session, project: Project) -> BytesIO:
    document = Document()
    _paysage(document.sections[0])
    _en_tete_document(document, project, "CADRE DE RENDEMENT",
                      "Performance Measurement Framework — mesure et suivi des résultats")
    document.add_paragraph()

    document.add_paragraph(
        "Le taux affiché mesure l'atteinte du jalon de la période évaluée ; la colonne "
        "« Progr. finale » indique la progression accomplie depuis la valeur de référence vers "
        "la cible de fin de projet.").runs[0].font.size = Pt(8)
    table = _tableau(document, ["Niveau", "Résultat attendu", "Indicateur", "Réf.", "Cible finale",
                                "Période", "Cible période", "Réalisé", "Taux", "Progr. finale",
                                "Statut", "Source de données", "Fréquence", "Responsable"],
                     largeurs=[1.7, 4.2, 5, 1.3, 1.5, 1.5, 1.5, 1.4, 1.3, 1.5, 1.8, 2.8, 1.7, 2.2])
    elements = {e.id: e for e in db.query(LogframeElement).filter(
        LogframeElement.project_id == project.id).all()}
    perf = {p["id"]: p for p in analytics.synthese_indicateurs(db, project.id)["lignes"]}
    ordre = {"IMPACT": 0, "EFFET": 1, "PRODUIT": 2, "ACTIVITE": 3}
    for ind in sorted(db.query(Indicator).filter(Indicator.project_id == project.id).all(),
                      key=lambda i: (ordre.get(i.level, 9), i.code or "")):
        p = perf.get(ind.id, {})
        element = elements.get(ind.element_id)
        ligne = table.add_row()
        _ombrer(ligne.cells[0], COULEURS_NIVEAU.get(ind.level or "", "BFBFBF"))
        _texte_cellule(ligne.cells[0], ind.level or "—", gras=True, blanc=True, centre=True, taille=7.5)
        _texte_cellule(ligne.cells[1], element.statement if element else "—", taille=8)
        _texte_cellule(ligne.cells[2], f"{ind.code or ''} {ind.name}".strip(), taille=8)
        _texte_cellule(ligne.cells[3], _nombre(ind.baseline_value), centre=True, taille=8)
        _texte_cellule(ligne.cells[4], _nombre(ind.target_value), centre=True, taille=8)
        _texte_cellule(ligne.cells[5], p.get("period_label") or "—", centre=True, taille=8)
        _texte_cellule(ligne.cells[6], _nombre(p.get("period_target")), centre=True, taille=8)
        _texte_cellule(ligne.cells[7], _nombre(p.get("actual_value")), centre=True, taille=8)
        taux = p.get("taux")
        _texte_cellule(ligne.cells[8], f"{taux} %" if taux is not None else "—",
                       centre=True, gras=True, taille=8)
        taux_final = p.get("taux_final")
        _texte_cellule(ligne.cells[9], f"{taux_final} %" if taux_final is not None else "—",
                       centre=True, taille=8)
        statut = p.get("statut", "Non renseigné")
        _ombrer(ligne.cells[10], analytics.COULEURS_STATUT.get(statut, "#9AA0A6").lstrip("#"))
        _texte_cellule(ligne.cells[10], statut, centre=True, gras=True, blanc=True, taille=7.5)
        _texte_cellule(ligne.cells[11], ind.data_source or "—", taille=8)
        _texte_cellule(ligne.cells[12], ind.frequency or "—", centre=True, taille=8)
        _texte_cellule(ligne.cells[13], ind.responsible or "—", taille=8)
    _pied_de_page(document, f"{project.code} — Cadre de rendement — Généré par {APP_NAME}")
    return _sauver(document)


# ---------------------------------------------------------------------------
# 3. Registre des risques
# ---------------------------------------------------------------------------
def risques_docx(db: Session, project: Project) -> BytesIO:
    document = Document()
    _paysage(document.sections[0])
    _en_tete_document(document, project, "REGISTRE ET PLAN DE GESTION DES RISQUES",
                      "Identification, cotation, atténuation et suivi des risques du projet")

    synthese = analytics.synthese_risques(db, project.id)
    document.add_paragraph()
    p = document.add_paragraph()
    run = p.add_run(f"Synthèse : {synthese['total']} risques recensés, dont {synthese['critiques']} "
                    f"critiques et {synthese['ouverts']} ouverts. Score moyen : "
                    f"{synthese['score_moyen']}/25.")
    run.bold = True
    run.font.size = Pt(10)

    heading = document.add_heading("1. Matrice de criticité (Probabilité × Impact)", level=1)
    heading.runs[0].font.color.rgb = BLEU
    matrice = _tableau(document, ["Impact ↓ / Probabilité →", "1 — Très faible", "2 — Faible",
                                  "3 — Moyenne", "4 — Forte", "5 — Très forte"],
                       largeurs=[4.5, 3.4, 3.4, 3.4, 3.4, 3.4])
    libelles = {5: "5 — Catastrophique", 4: "4 — Majeur", 3: "3 — Modéré",
                2: "2 — Mineur", 1: "1 — Négligeable"}
    for impact in range(5, 0, -1):
        ligne = matrice.add_row()
        _ombrer(ligne.cells[0], "1F4E79")
        _texte_cellule(ligne.cells[0], libelles[impact], gras=True, blanc=True, centre=True)
        for proba in range(1, 6):
            score = impact * proba
            couleur = "D93025" if score >= 15 else "EA8600" if score >= 10 else \
                      "F9A825" if score >= 5 else "0F9D58"
            nb = synthese["matrice"][impact - 1][proba - 1]
            _ombrer(ligne.cells[proba], couleur)
            _texte_cellule(ligne.cells[proba], f"{nb}" if nb else "", gras=True, blanc=True,
                           centre=True, taille=11)

    document.add_paragraph()
    heading = document.add_heading("2. Registre détaillé des risques", level=1)
    heading.runs[0].font.color.rgb = BLEU
    table = _tableau(document, ["Code", "Catégorie", "Risque", "Cause / Conséquence", "P", "I",
                                "Score", "Niveau", "Mesures d'atténuation", "Responsable", "Statut"],
                     largeurs=[1.3, 2.6, 5.5, 5, 0.9, 0.9, 1.1, 1.8, 6, 2.6, 1.9])
    for r in sorted(db.query(Risk).filter(Risk.project_id == project.id).all(), key=lambda x: -x.score):
        ligne = table.add_row()
        valeurs = [r.code, r.category, r.title,
                   f"Cause : {r.cause or '—'}\nConséquence : {r.consequence or '—'}",
                   r.probability, r.impact, r.score, r.severity, r.mitigation, r.owner, r.status]
        for index, valeur in enumerate(valeurs):
            _texte_cellule(ligne.cells[index], valeur if valeur not in (None, "") else "—",
                           centre=index in (4, 5, 6), taille=8)
        _ombrer(ligne.cells[7], COULEURS_SEVERITE.get(r.severity, "9AA0A6"))
        _texte_cellule(ligne.cells[7], r.severity, gras=True, blanc=True, centre=True, taille=8)

    document.add_page_break()
    heading = document.add_heading("3. Plans de contingence", level=1)
    heading.runs[0].font.color.rgb = BLEU
    for r in db.query(Risk).filter(Risk.project_id == project.id).order_by(Risk.code).all():
        if not r.contingency:
            continue
        p = document.add_paragraph()
        run = p.add_run(f"{r.code or ''} — {r.title} ({r.severity})")
        run.bold = True
        run.font.size = Pt(10)
        document.add_paragraph(r.contingency, style="List Bullet")
    _pied_de_page(document, f"{project.code} — Registre des risques — Généré par {APP_NAME}")
    return _sauver(document)


# ---------------------------------------------------------------------------
# 4. Fiches d'indicateurs (metadata sheets)
# ---------------------------------------------------------------------------
def fiches_indicateurs_docx(db: Session, project: Project,
                            indicator_id: Optional[int] = None) -> BytesIO:
    document = Document()
    section = document.sections[0]
    section.left_margin = section.right_margin = Cm(2)
    _en_tete_document(document, project, "FICHES MÉTADONNÉES DES INDICATEURS",
                      "Documentation opérationnelle de chaque indicateur de suivi")

    requete = db.query(Indicator).filter(Indicator.project_id == project.id)
    if indicator_id:
        requete = requete.filter(Indicator.id == indicator_id)
    indicateurs = requete.order_by(Indicator.code).all()
    elements = {e.id: e for e in db.query(LogframeElement).filter(
        LogframeElement.project_id == project.id).all()}

    for position, ind in enumerate(indicateurs):
        if position:
            document.add_page_break()
        titre = document.add_heading(f"{ind.code or 'IND'} — {ind.name}", level=1)
        titre.runs[0].font.color.rgb = BLEU
        element = elements.get(ind.element_id)
        perf = analytics.indicator_performance(ind)
        rubriques = [
            ("Résultat rattaché", f"{element.code or ''} {element.statement}" if element else "—"),
            ("Niveau du cadre logique", LIBELLES_NIVEAUX.get(ind.level, ind.level or "—")),
            ("Type d'indicateur", ind.indicator_type or "—"),
            ("Définition opérationnelle", ind.definition or "—"),
            ("Unité de mesure", ind.unit or "—"),
            ("Mode de calcul / formule", ind.formula or "—"),
            ("Numérateur", ind.numerator or "—"),
            ("Dénominateur", ind.denominator or "—"),
            ("Désagrégation exigée", ", ".join(ind.disaggregation or []) or "—"),
            ("Valeur de référence", f"{_nombre(ind.baseline_value)} au {_date_fr(ind.baseline_date)}"
                                    f" (source : {ind.baseline_source or 'N/A'})"),
            ("Valeur cible", f"{_nombre(ind.target_value)} au {_date_fr(ind.target_date)}"),
            ("Sens de progression", ind.direction or "—"),
            ("Dernière valeur réalisée", f"{_nombre(perf['actual_value'])} "
                                         f"({perf['period_label'] or 'N/A'})"),
            ("Taux de réalisation", f"{perf['taux']}% — {perf['statut']}"
                                    if perf["taux"] is not None else "Non renseigné"),
            ("Fréquence de collecte", ind.frequency or "—"),
            ("Source de données", ind.data_source or "—"),
            ("Méthode de collecte", ind.collection_method or "—"),
            ("Responsable de la collecte", ind.responsible or "—"),
            ("Niveau de rapportage", ind.reporting_level or "—"),
            ("Coût estimé de la collecte", _nombre(ind.cost_estimate)),
            ("Indicateur clé (KPI)", "Oui" if ind.is_key else "Non"),
            ("Contrôle qualité SMART", _texte_smart(ind.smart_check)),
            ("Observations qualité", ind.quality_note or "—"),
        ]
        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for libelle, valeur in rubriques:
            ligne = table.add_row()
            ligne.cells[0].width = Cm(5.5)
            ligne.cells[1].width = Cm(11.5)
            _ombrer(ligne.cells[0], "DCE6F1")
            _texte_cellule(ligne.cells[0], libelle, gras=True, taille=9)
            _texte_cellule(ligne.cells[1], valeur, taille=9)

        if ind.targets:
            document.add_paragraph()
            p = document.add_paragraph()
            run = p.add_run("Cibles et réalisations par période")
            run.bold = True
            run.font.size = Pt(10)
            serie = analytics.serie_temporelle(ind)
            table_p = _tableau(document, ["Période", "Cible", "Réalisé", "Taux (%)"],
                               largeurs=[4.5, 4, 4, 4])
            for index, periode in enumerate(serie["periodes"]):
                cible = serie["cibles"][index]
                reel = serie["reels"][index]
                ligne = table_p.add_row()
                taux = round(reel / cible * 100, 1) if (cible and reel is not None) else None
                for col, valeur in enumerate([periode, _nombre(cible), _nombre(reel),
                                              f"{taux}%" if taux is not None else "—"]):
                    _texte_cellule(ligne.cells[col], valeur, centre=col > 0, taille=9)
    _pied_de_page(document, f"{project.code} — Fiches indicateurs — Généré par {APP_NAME}")
    return _sauver(document)


def _texte_smart(smart: Optional[Dict[str, Any]]) -> str:
    if not smart:
        return "Non évalué"
    libelles = {"specifique": "Spécifique", "mesurable": "Mesurable", "atteignable": "Atteignable",
                "pertinent": "Pertinent", "temporel": "Temporellement défini"}
    return " | ".join(f"{libelles.get(cle, cle)} : {'✔' if valeur else '✘'}"
                      for cle, valeur in smart.items())


# ---------------------------------------------------------------------------
# 5. Questionnaire / fiche de collecte
# ---------------------------------------------------------------------------
def questionnaire_docx(db: Session, form: Form, project: Project) -> BytesIO:
    document = Document()
    section = document.sections[0]
    section.left_margin = section.right_margin = Cm(2)
    _en_tete_document(document, project, form.name.upper(),
                      f"{form.form_type or 'Questionnaire'} — version {form.version or '1.0'}")

    if form.description:
        document.add_paragraph(form.description)
    encadre = document.add_table(rows=0, cols=2)
    encadre.style = "Table Grid"
    for libelle, valeur in [("Code du formulaire", form.code or "—"),
                            ("Population cible", form.target_respondent or "—"),
                            ("Périodicité de collecte", form.periodicity or "—"),
                            ("Indicateurs renseignés", ", ".join(form.linked_indicators or []) or "—"),
                            ("Langue", form.language or "fr")]:
        ligne = encadre.add_row()
        ligne.cells[0].width = Cm(5)
        _ombrer(ligne.cells[0], "DCE6F1")
        _texte_cellule(ligne.cells[0], libelle, gras=True)
        _texte_cellule(ligne.cells[1], valeur)

    if form.instructions:
        document.add_paragraph()
        titre = document.add_heading("Consignes à l'enquêteur", level=2)
        titre.runs[0].font.color.rgb = BLEU
        document.add_paragraph(form.instructions)

    document.add_paragraph()
    titre = document.add_heading("Identification de la collecte", level=2)
    titre.runs[0].font.color.rgb = BLEU
    identification = _tableau(document, ["Rubrique", "Réponse"], largeurs=[7, 10])
    for rubrique in ["Date de l'entretien", "Région / Préfecture", "Commune / Village",
                     "Nom de l'enquêteur", "Code du questionnaire", "Coordonnées GPS"]:
        ligne = identification.add_row()
        _texte_cellule(ligne.cells[0], rubrique, gras=True)
        _texte_cellule(ligne.cells[1], "")

    section_courante = None
    numero = 0
    for question in form.questions:
        if question.section and question.section != section_courante:
            section_courante = question.section
            document.add_paragraph()
            titre = document.add_heading(section_courante, level=2)
            titre.runs[0].font.color.rgb = BLEU
        if question.question_type == "note":
            p = document.add_paragraph()
            run = p.add_run(question.label)
            run.italic = True
            continue
        numero += 1
        p = document.add_paragraph()
        run = p.add_run(f"Q{numero}. {question.label}")
        run.bold = True
        run.font.size = Pt(10.5)
        if question.required:
            obligatoire = p.add_run("  *obligatoire")
            obligatoire.font.size = Pt(8)
            obligatoire.font.color.rgb = RGBColor(0xD9, 0x30, 0x25)
        if question.hint:
            hint = document.add_paragraph()
            run = hint.add_run(f"({question.hint})")
            run.italic = True
            run.font.size = Pt(8.5)
            run.font.color.rgb = GRIS
        if question.question_type in ("select_one", "select_multiple"):
            marque = "◯" if question.question_type == "select_one" else "☐"
            for choix in question.choices or []:
                libelle = choix.get("label") if isinstance(choix, dict) else str(choix)
                code = choix.get("name") if isinstance(choix, dict) else ""
                item = document.add_paragraph(f"     {marque}  {code}. {libelle}"
                                              if code else f"     {marque}  {libelle}")
                item.paragraph_format.space_after = Pt(2)
        elif question.question_type in ("integer", "decimal"):
            document.add_paragraph("     |__|__|__|__|__|   " + (question.hint or ""))
        elif question.question_type == "date":
            document.add_paragraph("     |__|__| / |__|__| / |__|__|__|__|   (JJ/MM/AAAA)")
        elif question.question_type == "geopoint":
            document.add_paragraph("     Latitude : ____________   Longitude : ____________")
        else:
            document.add_paragraph("     " + "_" * 90)
            document.add_paragraph("     " + "_" * 90)
        if question.constraint:
            p = document.add_paragraph()
            run = p.add_run(f"     Contrôle de cohérence : {question.constraint}")
            run.font.size = Pt(7.5)
            run.font.color.rgb = GRIS
        if question.relevant:
            p = document.add_paragraph()
            run = p.add_run(f"     Ne poser que si : {question.relevant}")
            run.font.size = Pt(7.5)
            run.italic = True
            run.font.color.rgb = GRIS

    document.add_paragraph()
    document.add_paragraph("_" * 100)
    p = document.add_paragraph()
    run = p.add_run("Observations de l'enquêteur :")
    run.bold = True
    for _ in range(3):
        document.add_paragraph("_" * 100)
    p = document.add_paragraph()
    run = p.add_run("\nSignature de l'enquêteur : ______________________        "
                    "Visa du superviseur : ______________________")
    run.font.size = Pt(9)
    _pied_de_page(document, f"{project.code} — {form.name} v{form.version} — Généré par {APP_NAME}")
    return _sauver(document)


# ---------------------------------------------------------------------------
# 6. Plan / manuel de suivi-évaluation (document maître)
# ---------------------------------------------------------------------------
def plan_suivi_evaluation_docx(db: Session, project: Project) -> BytesIO:
    document = Document()
    section = document.sections[0]
    section.left_margin = section.right_margin = Cm(2.2)

    # Page de garde
    for _ in range(5):
        document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(project.executing_agency or "AGENCE D'EXÉCUTION")
    run.bold = True
    run.font.size = Pt(13)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n{project.code} — {project.title}")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = BLEU
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\nPLAN ET MANUEL DE SUIVI-ÉVALUATION")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = BLEU
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Dispositif de suivi, d'évaluation, d'apprentissage et de redevabilité (SEAR)")
    run.italic = True
    run.font.size = Pt(11)
    for _ in range(6):
        document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Bailleur : {project.donor or 'N/A'}\n"
                    f"Période d'exécution : {_date_fr(project.start_date)} — {_date_fr(project.end_date)}\n"
                    f"Budget total : {_nombre(project.total_budget)} {project.currency}\n\n"
                    f"Version générée le {date.today().strftime('%d %B %Y')} par la plateforme "
                    f"{APP_NAME}")
    run.font.size = Pt(10)
    run.font.color.rgb = GRIS
    document.add_page_break()

    # Sommaire
    titre = document.add_heading("SOMMAIRE", level=1)
    titre.runs[0].font.color.rgb = BLEU
    sommaire = [
        "1. Introduction et objet du manuel", "2. Présentation du projet",
        "3. Cadre conceptuel du suivi-évaluation", "4. Chaîne de résultats et cadre logique",
        "5. Système d'indicateurs", "6. Cadre de mesure du rendement",
        "7. Dispositif de collecte des données", "8. Gestion des risques et des hypothèses",
        "9. Planification opérationnelle (chronogramme et PTBA)",
        "10. Rapportage et diffusion de l'information",
        "11. Évaluations et études", "12. Dispositif organisationnel et responsabilités",
        "13. Assurance qualité des données", "14. Apprentissage et gestion des connaissances",
        "15. Budget du dispositif de S&E", "Annexes",
    ]
    for entree in sommaire:
        document.add_paragraph(entree, style="List Number" if entree[0].isdigit() else "List Bullet")
    document.add_page_break()

    def chapitre(numero_titre: str):
        h = document.add_heading(numero_titre, level=1)
        h.runs[0].font.color.rgb = BLEU
        return h

    def sous_chapitre(texte: str):
        h = document.add_heading(texte, level=2)
        h.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        return h

    # --- 1. Introduction
    chapitre("1. Introduction et objet du manuel")
    document.add_paragraph(
        f"Le présent manuel définit le dispositif de suivi-évaluation (S&E) du projet "
        f"« {project.title} ». Il constitue le document de référence opposable à l'ensemble des "
        f"acteurs impliqués dans la mise en œuvre : unité de gestion du projet, partenaires "
        f"d'exécution, prestataires, services techniques déconcentrés et bailleur de fonds.")
    document.add_paragraph(
        "Il poursuit quatre finalités : (i) harmoniser la compréhension des résultats attendus et "
        "des indicateurs de mesure ; (ii) standardiser les procédures de collecte, de traitement et "
        "de validation des données ; (iii) organiser la production et la diffusion des rapports de "
        "performance dans des délais maîtrisés ; (iv) alimenter la prise de décision, "
        "l'apprentissage organisationnel et la redevabilité envers les bénéficiaires et les "
        "partenaires financiers.")
    document.add_paragraph(
        "Le manuel est un document vivant : il est actualisé à chaque revue annuelle du projet et "
        "chaque fois qu'une modification substantielle du cadre logique, du cadre de rendement ou "
        "des modalités de collecte est approuvée par le comité de pilotage.")

    # --- 2. Présentation du projet
    chapitre("2. Présentation du projet")
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    fiche_projet = [
        ("Code du projet", project.code), ("Intitulé", project.title),
        ("Acronyme", project.acronym or "—"), ("Secteur", project.sector or "—"),
        ("Sous-secteur", project.sub_sector or "—"), ("Pays", project.country or "—"),
        ("Zones d'intervention", ", ".join(project.regions or []) or "—"),
        ("Bailleur / PTF", project.donor or "—"),
        ("Agence d'exécution", project.executing_agency or "—"),
        ("Ministère de tutelle", project.supervising_ministry or "—"),
        ("Bénéficiaires", project.beneficiaries or "—"),
        ("Population cible", _nombre(project.target_population)),
        ("Date de démarrage", _date_fr(project.start_date)),
        ("Date de clôture", _date_fr(project.end_date)),
        ("Statut", project.status or "—"),
        ("Budget total", f"{_nombre(project.total_budget)} {project.currency}"),
        ("Contrepartie nationale", f"{_nombre(project.counterpart_budget)} {project.currency}"),
    ]
    for libelle, valeur in fiche_projet:
        ligne = table.add_row()
        ligne.cells[0].width = Cm(5.5)
        _ombrer(ligne.cells[0], "DCE6F1")
        _texte_cellule(ligne.cells[0], libelle, gras=True)
        _texte_cellule(ligne.cells[1], valeur)
    if project.description:
        sous_chapitre("2.1 Description du projet")
        document.add_paragraph(project.description)
    if project.theory_of_change:
        sous_chapitre("2.2 Théorie du changement")
        document.add_paragraph(project.theory_of_change)
    if project.strategic_alignment:
        sous_chapitre("2.3 Alignement stratégique")
        for cle, valeur in (project.strategic_alignment or {}).items():
            document.add_paragraph(f"{cle} : {valeur}", style="List Bullet")

    # --- 3. Cadre conceptuel
    document.add_page_break()
    chapitre("3. Cadre conceptuel du suivi-évaluation")
    sous_chapitre("3.1 Définitions opératoires")
    definitions = [
        ("Suivi", "Processus continu de collecte et d'analyse d'informations permettant de comparer "
                  "l'état d'avancement réel de la mise en œuvre à la programmation initiale, afin "
                  "d'informer la gestion courante."),
        ("Évaluation", "Appréciation systématique et objective, à un moment donné, de la "
                       "pertinence, de la cohérence, de l'efficacité, de l'efficience, de l'impact "
                       "et de la durabilité d'une intervention (critères du CAD de l'OCDE)."),
        ("Gestion axée sur les résultats (GAR)", "Approche de gestion centrée sur l'atteinte de "
                                                 "résultats mesurables plutôt que sur la seule "
                                                 "exécution des activités."),
        ("Chaîne de résultats", "Enchaînement logique intrants → activités → produits → effets → "
                                "impact, reliant les moyens mobilisés aux changements recherchés."),
        ("Indicateur", "Variable quantitative ou qualitative permettant de mesurer objectivement "
                       "les changements induits par l'intervention."),
        ("Référence (baseline)", "Valeur de l'indicateur avant le démarrage des activités, servant "
                                 "de point de comparaison pour apprécier les progrès."),
        ("Cible", "Valeur attendue de l'indicateur à une échéance déterminée."),
        ("Hypothèse", "Condition externe nécessaire à la réalisation de la chaîne de résultats mais "
                      "hors du contrôle direct du projet."),
        ("Risque", "Événement incertain dont la survenue affecterait négativement l'atteinte des "
                   "résultats ; il est coté selon sa probabilité et son impact."),
    ]
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for terme, definition in definitions:
        ligne = table.add_row()
        ligne.cells[0].width = Cm(4.5)
        _ombrer(ligne.cells[0], "DCE6F1")
        _texte_cellule(ligne.cells[0], terme, gras=True)
        _texte_cellule(ligne.cells[1], definition)

    sous_chapitre("3.2 Approche retenue")
    document.add_paragraph(project.me_approach or
        "Le dispositif s'appuie sur l'approche de la gestion axée sur les résultats (GAR) et sur "
        "le cadre logique comme instrument central de planification. Le suivi est organisé à trois "
        "niveaux complémentaires : le suivi de l'exécution physique et financière (activités, "
        "PTBA, décaissements), le suivi des produits (extrants livrés) et le suivi des effets "
        "(changements observés chez les bénéficiaires). L'évaluation intervient aux moments clés du "
        "cycle : évaluation de référence, revue à mi-parcours, évaluation finale et, le cas "
        "échéant, évaluation d'impact ex post.")

    # --- 4. Cadre logique
    document.add_page_break()
    chapitre("4. Chaîne de résultats et cadre logique")
    document.add_paragraph(
        "La matrice ci-après présente la logique d'intervention du projet, les indicateurs "
        "objectivement vérifiables associés à chaque niveau de résultat, les sources de "
        "vérification et les hypothèses critiques.")
    elements = _elements_tries(db, project.id)
    indicateurs_par_element: Dict[Any, List[Indicator]] = {}
    for ind in db.query(Indicator).filter(Indicator.project_id == project.id).all():
        indicateurs_par_element.setdefault(ind.element_id, []).append(ind)
    table = _tableau(document, ["Niveau", "Code", "Énoncé du résultat", "Indicateurs", "Hypothèses"],
                     largeurs=[2.2, 1.5, 5.5, 5, 3])
    for element in elements:
        ligne = table.add_row()
        _ombrer(ligne.cells[0], COULEURS_NIVEAU.get(element.level, "BFBFBF"))
        _texte_cellule(ligne.cells[0], element.level, gras=True, blanc=True, centre=True, taille=8)
        _texte_cellule(ligne.cells[1], element.code or "—", centre=True, taille=8)
        _texte_cellule(ligne.cells[2], element.statement, taille=8)
        _texte_cellule(ligne.cells[3], "\n".join(
            f"• {i.code or ''} {i.name}" for i in indicateurs_par_element.get(element.id, [])) or "—",
            taille=8)
        _texte_cellule(ligne.cells[4], element.assumptions or "—", taille=8)

    # --- 5. Indicateurs
    document.add_page_break()
    chapitre("5. Système d'indicateurs")
    synthese_ind = analytics.synthese_indicateurs(db, project.id)
    document.add_paragraph(
        f"Le système de suivi comporte {synthese_ind['total']} indicateurs, dont "
        f"{len([l for l in synthese_ind['lignes'] if l['is_key']])} indicateurs clés de performance. "
        f"Le taux de couverture (indicateurs effectivement renseignés) s'établit à "
        f"{synthese_ind['taux_couverture']}%.")
    sous_chapitre("5.1 Critères de qualité des indicateurs")
    document.add_paragraph(
        "Chaque indicateur du dispositif est soumis au test SMART : Spécifique (il mesure sans "
        "ambiguïté un aspect précis du résultat), Mesurable (il est quantifiable ou objectivement "
        "appréciable), Atteignable (la cible est réaliste au regard des moyens), Pertinent (il "
        "rend compte du changement recherché) et Temporellement défini (il est assorti d'une "
        "échéance). Les indicateurs sont en outre appréciés selon leur coût de collecte, leur "
        "sensibilité au changement et leur comparabilité dans le temps.")
    sous_chapitre("5.2 Répartition des indicateurs par niveau de résultat")
    table = _tableau(document, ["Niveau de résultat", "Nombre d'indicateurs",
                                "Indicateurs renseignés", "Taux moyen de réalisation"],
                     largeurs=[6, 3.5, 3.5, 4])
    for niveau, valeurs in synthese_ind["par_niveau"].items():
        ligne = table.add_row()
        _texte_cellule(ligne.cells[0], LIBELLES_NIVEAUX.get(niveau, niveau))
        _texte_cellule(ligne.cells[1], valeurs["nombre"], centre=True)
        _texte_cellule(ligne.cells[2], valeurs["renseignes"], centre=True)
        _texte_cellule(ligne.cells[3], f"{valeurs['taux_moyen']}%"
                       if valeurs["taux_moyen"] is not None else "—", centre=True)

    sous_chapitre("5.3 Liste des indicateurs et paramétrage")
    table = _tableau(document, ["Code", "Indicateur", "Unité", "Référence", "Cible", "Fréquence",
                                "Source", "Responsable"],
                     largeurs=[1.4, 5.5, 1.8, 1.6, 1.6, 2, 3, 2.5])
    for ind in db.query(Indicator).filter(Indicator.project_id == project.id).order_by(
            Indicator.code).all():
        ligne = table.add_row()
        for index, valeur in enumerate([ind.code, ind.name, ind.unit, _nombre(ind.baseline_value),
                                        _nombre(ind.target_value), ind.frequency, ind.data_source,
                                        ind.responsible]):
            _texte_cellule(ligne.cells[index], valeur if valeur not in (None, "") else "—", taille=8)

    # --- 6. Cadre de rendement
    document.add_page_break()
    chapitre("6. Cadre de mesure du rendement")
    document.add_paragraph(
        "Le cadre de mesure du rendement précise, pour chaque indicateur, la source de données, la "
        "méthode et la fréquence de collecte, le responsable et le coût prévisionnel de la mesure. "
        "Il est extrait automatiquement de la plateforme et joint en annexe sous forme de tableau "
        "détaillé (fichier « Cadre de rendement »).")
    table = _tableau(document, ["Indicateur", "Source de données", "Méthode de collecte",
                                "Fréquence", "Responsable", "Coût estimé"],
                     largeurs=[5.5, 3.5, 3.2, 2, 2.5, 2])
    for ind in db.query(Indicator).filter(Indicator.project_id == project.id).order_by(
            Indicator.code).all():
        ligne = table.add_row()
        for index, valeur in enumerate([f"{ind.code or ''} {ind.name}", ind.data_source,
                                        ind.collection_method, ind.frequency, ind.responsible,
                                        _nombre(ind.cost_estimate)]):
            _texte_cellule(ligne.cells[index], valeur if valeur not in (None, "") else "—", taille=8)

    # --- 7. Collecte
    document.add_page_break()
    chapitre("7. Dispositif de collecte des données")
    sous_chapitre("7.1 Instruments de collecte")
    formulaires = db.query(Form).filter(Form.project_id == project.id).all()
    if formulaires:
        table = _tableau(document, ["Code", "Instrument", "Type", "Population cible", "Périodicité",
                                    "Nb de questions", "Indicateurs renseignés"],
                         largeurs=[1.5, 4.5, 2.5, 3, 2.2, 1.8, 3.5])
        for f in formulaires:
            ligne = table.add_row()
            for index, valeur in enumerate([f.code, f.name, f.form_type, f.target_respondent,
                                            f.periodicity, len(f.questions),
                                            ", ".join(f.linked_indicators or [])]):
                _texte_cellule(ligne.cells[index], valeur if valeur not in (None, "") else "—", taille=8)
    else:
        document.add_paragraph("Aucun instrument de collecte n'est encore paramétré dans la plateforme.")
    sous_chapitre("7.2 Circuit de l'information")
    for etape in [
        "Collecte primaire sur le terrain par les agents de suivi, au moyen des fiches et "
        "questionnaires numérisés (application mobile KoboToolbox / ODK Collect, ou saisie directe "
        "dans la plateforme SEPIA).",
        "Contrôle de premier niveau par le superviseur de zone : vérification de l'exhaustivité, "
        "de la cohérence interne et du respect de l'échantillon.",
        "Transmission et centralisation des données dans la base de la plateforme, sous un délai "
        "maximal de sept (7) jours après la fin de la période de collecte.",
        "Contrôle de second niveau et validation par le responsable du suivi-évaluation : "
        "traitement des valeurs aberrantes, triangulation avec les sources secondaires.",
        "Calcul automatique des indicateurs et des taux de réalisation par la plateforme.",
        "Analyse, production des tableaux de bord et diffusion aux instances de pilotage.",
    ]:
        document.add_paragraph(etape, style="List Number")
    sous_chapitre("7.3 Échantillonnage et désagrégation")
    document.add_paragraph(
        "Les enquêtes de mesure des effets reposent sur un échantillonnage aléatoire stratifié, "
        "les strates étant définies par zone d'intervention et par catégorie de bénéficiaires. "
        "La taille de l'échantillon est calculée pour un niveau de confiance de 95 % et une marge "
        "d'erreur maximale de 5 %. Toutes les données individuelles sont désagrégées au minimum "
        "par sexe et par tranche d'âge, conformément aux exigences de suivi sensible au genre.")

    # --- 8. Risques
    document.add_page_break()
    chapitre("8. Gestion des risques et des hypothèses")
    synthese_risques = analytics.synthese_risques(db, project.id)
    document.add_paragraph(
        f"Le registre des risques comporte {synthese_risques['total']} risques identifiés, dont "
        f"{synthese_risques['critiques']} de niveau critique. Chaque risque est coté sur une "
        f"échelle de 1 à 5 en probabilité et en impact ; le produit des deux notes détermine le "
        f"score de criticité (échelle de 1 à 25) et le niveau de priorité de traitement.")
    table = _tableau(document, ["Code", "Risque", "Catégorie", "P", "I", "Score", "Niveau",
                                "Atténuation", "Responsable"],
                     largeurs=[1.2, 4.5, 2.5, 0.8, 0.8, 1.1, 1.6, 4.5, 2.5])
    for r in sorted(db.query(Risk).filter(Risk.project_id == project.id).all(), key=lambda x: -x.score):
        ligne = table.add_row()
        for index, valeur in enumerate([r.code, r.title, r.category, r.probability, r.impact,
                                        r.score, r.severity, r.mitigation, r.owner]):
            _texte_cellule(ligne.cells[index], valeur if valeur not in (None, "") else "—",
                           centre=index in (3, 4, 5), taille=8)
        _ombrer(ligne.cells[6], COULEURS_SEVERITE.get(r.severity, "9AA0A6"))
        _texte_cellule(ligne.cells[6], r.severity, gras=True, blanc=True, centre=True, taille=8)
    sous_chapitre("8.1 Suivi des hypothèses critiques")
    hypotheses = db.query(Assumption).filter(Assumption.project_id == project.id).all()
    if hypotheses:
        table = _tableau(document, ["Code", "Hypothèse", "Niveau", "Criticité", "Statut",
                                    "Méthode de vérification", "Responsable"],
                         largeurs=[1.2, 5.5, 1.8, 1.8, 2.2, 4, 2.5])
        for h in hypotheses:
            ligne = table.add_row()
            for index, valeur in enumerate([h.code, h.statement, h.level, h.criticality,
                                            h.validation_status, h.verification_method, h.responsible]):
                _texte_cellule(ligne.cells[index], valeur if valeur not in (None, "") else "—", taille=8)
    else:
        document.add_paragraph("Les hypothèses critiques sont documentées dans la colonne dédiée du "
                               "cadre logique.")

    # --- 9. Planification
    document.add_page_break()
    chapitre("9. Planification opérationnelle")
    activites = db.query(Activity).filter(Activity.project_id == project.id).order_by(
        Activity.code).all()
    synthese_act = analytics.synthese_activites(db, project.id)
    document.add_paragraph(
        f"Le projet compte {synthese_act['total']} activités programmées, avec un taux "
        f"d'avancement physique moyen de {synthese_act['avancement_moyen']}%. "
        f"{synthese_act['nb_en_retard']} activité(s) accusent un retard par rapport au calendrier.")
    if activites:
        table = _tableau(document, ["Code", "Activité", "Responsable", "Début", "Fin",
                                    "Avancement", "Statut", "Coût prévu"],
                         largeurs=[1.4, 5.5, 2.5, 1.8, 1.8, 1.6, 1.8, 2.2])
        for a in activites:
            ligne = table.add_row()
            for index, valeur in enumerate([a.code, a.name, a.responsible, _date_fr(a.start_date),
                                            _date_fr(a.end_date), f"{a.progress or 0}%", a.status,
                                            _nombre(a.planned_cost)]):
                _texte_cellule(ligne.cells[index], valeur if valeur not in (None, "") else "—",
                               centre=index in (3, 4, 5), taille=8)
    sous_chapitre("9.1 Programmation budgétaire")
    budget = analytics.synthese_budget(db, project.id)
    table = _tableau(document, ["Rubrique", f"Montant ({project.currency})"], largeurs=[9, 6])
    for libelle, valeur in [("Budget planifié (PTBA)", budget["planifie"]),
                            ("Montant engagé", budget["engage"]),
                            ("Montant décaissé", budget["decaisse"]),
                            ("Solde disponible", budget["solde"]),
                            ("Taux d'exécution financière", f"{budget['taux_execution']} %"),
                            ("Taux d'engagement", f"{budget['taux_engagement']} %")]:
        ligne = table.add_row()
        _texte_cellule(ligne.cells[0], libelle, gras=True)
        _texte_cellule(ligne.cells[1], _nombre(valeur) if isinstance(valeur, (int, float)) else valeur,
                       centre=True)

    # --- 10. Rapportage
    document.add_page_break()
    chapitre("10. Rapportage et diffusion de l'information")
    table = _tableau(document, ["Produit de rapportage", "Périodicité", "Échéance de production",
                                "Responsable", "Destinataires"],
                     largeurs=[4.5, 2.2, 3.2, 3, 3.5])
    rapports = [
        ("Fiche de suivi mensuel des activités", "Mensuelle", "5 du mois suivant",
         "Assistant S&E", "Coordination"),
        ("Rapport trimestriel de performance", "Trimestrielle", "15 jours après la fin du trimestre",
         "Responsable S&E", "Comité technique, bailleur"),
        ("Tableau de bord des indicateurs", "Trimestrielle", "En continu (plateforme)",
         "Responsable S&E", "Toutes les parties prenantes"),
        ("Rapport semestriel d'avancement", "Semestrielle", "30 jours après la fin du semestre",
         "Coordonnateur", "Comité de pilotage, bailleur"),
        ("Rapport annuel de performance", "Annuelle", "45 jours après la clôture de l'exercice",
         "Coordonnateur", "Comité de pilotage, tutelle, bailleur"),
        ("PTBA de l'exercice suivant", "Annuelle", "31 octobre", "Coordonnateur", "Comité de pilotage"),
        ("Rapport de revue à mi-parcours", "Ponctuelle", "À mi-parcours du projet",
         "Consultant externe", "Comité de pilotage, bailleur"),
        ("Rapport d'évaluation finale", "Ponctuelle", "3 mois avant la clôture",
         "Consultant externe", "Ensemble des parties prenantes"),
    ]
    for valeurs in rapports:
        ligne = table.add_row()
        for index, valeur in enumerate(valeurs):
            _texte_cellule(ligne.cells[index], valeur, taille=8)

    # --- 11. Évaluations
    chapitre("11. Évaluations et études")
    document.add_paragraph(
        "Le dispositif prévoit quatre exercices évaluatifs structurants :")
    for texte in [
        "Étude de référence (baseline) : réalisée dans les six premiers mois d'exécution, elle "
        "établit les valeurs initiales de tous les indicateurs d'effet et d'impact et confirme la "
        "faisabilité des cibles.",
        "Revue à mi-parcours : elle apprécie la pertinence et l'efficacité de la mise en œuvre, "
        "identifie les blocages et propose, le cas échéant, une révision du cadre logique et des "
        "cibles.",
        "Évaluation finale : conduite par une équipe externe indépendante, elle porte sur les six "
        "critères du CAD de l'OCDE (pertinence, cohérence, efficacité, efficience, impact, "
        "durabilité) et formule des recommandations opérationnelles.",
        "Études thématiques et évaluation d'impact : selon les besoins, des études spécifiques "
        "(genre, environnement, chaînes de valeur) et, lorsque le dispositif expérimental le "
        "permet, une évaluation d'impact contrefactuelle sont conduites.",
    ]:
        document.add_paragraph(texte, style="List Bullet")

    # --- 12. Organisation
    document.add_page_break()
    chapitre("12. Dispositif organisationnel et responsabilités")
    table = _tableau(document, ["Acteur", "Responsabilités en matière de S&E", "Périodicité "
                                "d'intervention"], largeurs=[4, 9, 3.5])
    acteurs = [
        ("Comité de pilotage", "Valide le cadre logique, les cibles annuelles et le PTBA ; examine "
                               "les rapports de performance ; arbitre les réorientations "
                               "stratégiques.", "Semestrielle"),
        ("Coordonnateur du projet", "Garantit la mise en œuvre du dispositif de S&E ; valide les "
                                    "rapports avant transmission ; anime le dialogue avec le "
                                    "bailleur.", "Continue"),
        ("Responsable suivi-évaluation", "Administre la plateforme ; conçoit et actualise les "
                                         "outils ; consolide et analyse les données ; produit les "
                                         "rapports et tableaux de bord ; forme les acteurs.",
         "Continue"),
        ("Assistant(e) S&E / statisticien", "Saisit et contrôle les données ; assure la qualité de "
                                            "la base ; prépare les analyses.", "Continue"),
        ("Chefs de composante", "Renseignent les données d'activités et de produits relevant de "
                                "leur composante ; justifient les écarts.", "Mensuelle"),
        ("Agents de terrain / animateurs", "Réalisent la collecte primaire au moyen des fiches et "
                                           "questionnaires ; transmettent les données dans les "
                                           "délais.", "Continue"),
        ("Responsable administratif et financier", "Fournit les données d'exécution budgétaire "
                                                   "(engagements, décaissements).", "Mensuelle"),
        ("Partenaires d'exécution", "Transmettent les données relevant de leur périmètre selon le "
                                    "format convenu.", "Trimestrielle"),
        ("Bailleur de fonds", "Examine les rapports ; commandite les évaluations externes.",
         "Semestrielle"),
    ]
    for valeurs in acteurs:
        ligne = table.add_row()
        for index, valeur in enumerate(valeurs):
            _texte_cellule(ligne.cells[index], valeur, gras=(index == 0), taille=8)

    # --- 13. Qualité des données
    chapitre("13. Assurance qualité des données")
    document.add_paragraph(
        "La qualité des données est appréciée selon cinq dimensions standard : validité, "
        "fiabilité, exhaustivité, précision et actualité (timeliness). Un audit de la qualité des "
        "données (Data Quality Assessment) est conduit au moins une fois par an sur un échantillon "
        "d'indicateurs clés.")
    table = _tableau(document, ["Dimension", "Définition", "Mécanisme de contrôle"],
                     largeurs=[3, 6.5, 7])
    dimensions = [
        ("Validité", "La donnée mesure effectivement ce qu'elle prétend mesurer.",
         "Revue des définitions opérationnelles et des modes de calcul ; test de cohérence avec "
         "les sources secondaires."),
        ("Fiabilité", "La mesure est reproductible d'une période à l'autre et d'un agent à l'autre.",
         "Standardisation des outils, formation des enquêteurs, double saisie sur un échantillon."),
        ("Exhaustivité", "Toutes les unités attendues sont couvertes.",
         "Suivi des taux de réponse et de couverture géographique ; relances automatiques de la "
         "plateforme."),
        ("Précision", "La donnée est exempte d'erreurs de saisie ou de mesure.",
         "Contrôles de cohérence intégrés aux formulaires (contraintes XLSForm), détection des "
         "valeurs aberrantes."),
        ("Actualité", "La donnée est disponible au moment où la décision doit être prise.",
         "Respect des délais du calendrier de rapportage ; alertes automatiques en cas de retard."),
    ]
    for valeurs in dimensions:
        ligne = table.add_row()
        for index, valeur in enumerate(valeurs):
            _texte_cellule(ligne.cells[index], valeur, gras=(index == 0), taille=8)

    # --- 14. Apprentissage
    chapitre("14. Apprentissage et gestion des connaissances")
    document.add_paragraph(
        "Le suivi-évaluation ne se réduit pas à une fonction de contrôle : il alimente un cycle "
        "d'apprentissage continu. Des ateliers de revue de performance sont organisés chaque "
        "semestre avec l'ensemble des acteurs pour analyser collectivement les écarts, identifier "
        "les facteurs explicatifs et formuler des mesures correctrices assorties de responsables "
        "et d'échéances. Les enseignements tirés sont capitalisés dans un registre dédié et "
        "diffusés sous forme de notes techniques.")
    table = _tableau(document, ["Mécanisme d'apprentissage", "Périodicité", "Produit attendu"],
                     largeurs=[6, 3, 6])
    for valeurs in [
        ("Atelier de revue de performance", "Semestrielle", "Plan d'actions correctrices"),
        ("Réunion de coordination technique", "Mensuelle", "Compte rendu et décisions"),
        ("Registre des leçons apprises", "Continue", "Fiches de capitalisation"),
        ("Note technique thématique", "Ponctuelle", "Publication interne / externe"),
        ("Restitution aux bénéficiaires", "Annuelle", "Rapport de redevabilité simplifié"),
    ]:
        ligne = table.add_row()
        for index, valeur in enumerate(valeurs):
            _texte_cellule(ligne.cells[index], valeur, taille=8)

    # --- 15. Budget S&E
    chapitre("15. Budget du dispositif de suivi-évaluation")
    cout_indicateurs = sum(i.cost_estimate or 0 for i in db.query(Indicator).filter(
        Indicator.project_id == project.id).all())
    document.add_paragraph(
        f"Le coût prévisionnel de la collecte des données, estimé indicateur par indicateur, "
        f"s'élève à {_nombre(cout_indicateurs)} {project.currency}. Conformément aux bonnes "
        f"pratiques, une enveloppe comprise entre 3 % et 5 % du budget total du projet est "
        f"consacrée au dispositif de suivi-évaluation, soit un ordre de grandeur de "
        f"{_nombre((project.total_budget or 0) * 0.03)} à "
        f"{_nombre((project.total_budget or 0) * 0.05)} {project.currency}.")
    table = _tableau(document, ["Poste de dépense", "Description", f"Estimation ({project.currency})"],
                     largeurs=[5, 7, 3.5])
    postes = [
        ("Personnel S&E", "Responsable S&E, assistant, agents de saisie", (project.total_budget or 0) * 0.012),
        ("Études et enquêtes", "Baseline, enquêtes de suivi, évaluation finale", (project.total_budget or 0) * 0.015),
        ("Collecte de terrain", "Missions, per diem, carburant, communication", cout_indicateurs),
        ("Équipements et logiciels", "Tablettes, licences, hébergement de la plateforme", (project.total_budget or 0) * 0.003),
        ("Renforcement de capacités", "Formation des acteurs aux outils de S&E", (project.total_budget or 0) * 0.004),
        ("Ateliers de revue et capitalisation", "Revues semestrielles, restitutions", (project.total_budget or 0) * 0.006),
    ]
    for libelle, description, montant in postes:
        ligne = table.add_row()
        _texte_cellule(ligne.cells[0], libelle, gras=True, taille=8)
        _texte_cellule(ligne.cells[1], description, taille=8)
        _texte_cellule(ligne.cells[2], _nombre(round(montant, 0)), centre=True, taille=8)

    # --- Annexes
    document.add_page_break()
    chapitre("Annexes")
    for annexe in [
        "Annexe 1 — Cadre logique détaillé (export Excel/Word de la plateforme)",
        "Annexe 2 — Cadre de rendement complet",
        "Annexe 3 — Cadre de suivi des indicateurs (IPTT) avec cibles périodiques",
        "Annexe 4 — Fiches métadonnées des indicateurs",
        "Annexe 5 — Registre des risques et matrice de criticité",
        "Annexe 6 — Chronogramme d'exécution (diagramme de Gantt)",
        "Annexe 7 — Plan de travail et budget annuel (PTBA)",
        "Annexe 8 — Fiches de collecte et questionnaires (versions Word et XLSForm)",
        "Annexe 9 — Modèles de rapports périodiques",
        "Annexe 10 — Tableau de bord de performance (Excel automatisé et rapport Power BI)",
    ]:
        document.add_paragraph(annexe, style="List Bullet")
    document.add_paragraph()
    p = document.add_paragraph()
    run = p.add_run(f"Document produit automatiquement par la plateforme {APP_NAME} — "
                    f"{APP_LONG_NAME}. Toutes les données présentées sont extraites en temps réel "
                    f"de la base du projet.")
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = GRIS
    _pied_de_page(document, f"{project.code} — Plan et manuel de suivi-évaluation — {APP_NAME}")
    return _sauver(document)


# ---------------------------------------------------------------------------
# 7 bis. Rapport périodique de suivi (trimestriel, semestriel, annuel)
# ---------------------------------------------------------------------------
LIBELLES_RAPPORT = {
    "trimestriel": "RAPPORT TRIMESTRIEL DE SUIVI",
    "semestriel": "RAPPORT SEMESTRIEL D'AVANCEMENT",
    "annuel": "RAPPORT ANNUEL DE PERFORMANCE",
}


def rapport_periodique_docx(db: Session, project: Project, periode: str,
                            type_rapport: str = "trimestriel") -> BytesIO:
    """Rapport de suivi-évaluation portant sur une période de rapportage donnée.

    Structure conforme aux attentes des bailleurs : résumé exécutif, performance
    des indicateurs sur la période, analyse d'équité (genre et groupe cible),
    consolidation par zone d'intervention, exécution physique et financière,
    difficultés et mesures correctrices.
    """
    analyse = analytics.analyse_periode(db, project.id, periode)
    document = Document()
    section = document.sections[0]
    section.left_margin = section.right_margin = Cm(2)
    _en_tete_document(document, project, LIBELLES_RAPPORT.get(type_rapport, "RAPPORT DE SUIVI"),
                      f"Période de rapportage : {periode}")

    # --- 1. Résumé exécutif
    titre = document.add_heading("1. Résumé exécutif", level=1)
    titre.runs[0].font.color.rgb = BLEU
    par_statut = analyse["par_statut"]
    atteints = par_statut.get("Atteint", 0) + par_statut.get("En bonne voie", 0)
    document.add_paragraph(
        f"Sur la période {periode}, {analyse['total_indicateurs']} indicateurs ont fait l'objet "
        f"d'une programmation ou d'une mesure, dont {analyse['renseignes']} effectivement "
        f"renseignés. Le taux moyen d'atteinte des cibles de la période s'établit à "
        f"{analyse['taux_moyen'] if analyse['taux_moyen'] is not None else '—'} %. "
        f"{atteints} indicateur(s) sont atteints ou en bonne voie, "
        f"{par_statut.get('À surveiller', 0)} sont à surveiller et "
        f"{par_statut.get('Critique', 0)} présentent un écart critique.")
    document.add_paragraph(
        f"L'exécution physique des activités atteint {analyse['activites']['avancement_moyen']} % "
        f"en moyenne, avec {analyse['activites']['achevees']} activité(s) achevée(s) sur "
        f"{analyse['activites']['total']}. Le taux d'exécution financière de l'exercice "
        f"{analyse['budget']['annee'] or ''} s'élève à "
        f"{analyse['budget']['taux_execution']} %, soit "
        f"{_nombre(analyse['budget']['decaisse'])} {project.currency} décaissés sur "
        f"{_nombre(analyse['budget']['planifie'])} {project.currency} programmés.")
    equite = analyse["equite_genre"]
    if equite:
        document.add_paragraph(
            f"Les interventions de la période ont touché {_nombre(equite['total'])} bénéficiaires "
            f"identifiés, dont {_nombre(equite['femmes'])} femmes, soit "
            f"{equite['part_femmes']} % de l'effectif ({equite['appreciation'].lower()}).")

    # --- 2. Performance des indicateurs
    titre = document.add_heading("2. Performance des indicateurs sur la période", level=1)
    titre.runs[0].font.color.rgb = BLEU
    table = _tableau(document, ["Code", "Indicateur", "Unité", "Cible période", "Réalisé",
                                "Taux", "Statut", "Source", "Responsable"],
                     largeurs=[1.4, 5.5, 1.5, 1.8, 1.6, 1.3, 1.9, 2.5, 2.2])
    ordre = {"IMPACT": 0, "EFFET": 1, "PRODUIT": 2, "ACTIVITE": 3}
    for ligne_ind in sorted(analyse["lignes"],
                            key=lambda l: (ordre.get(l["level"], 9), l["code"] or "")):
        ligne = table.add_row()
        valeurs = [ligne_ind["code"], ligne_ind["name"], ligne_ind["unit"],
                   _nombre(ligne_ind["cible_periode"]), _nombre(ligne_ind["realise_periode"]),
                   f"{ligne_ind['taux']} %" if ligne_ind["taux"] is not None else "—"]
        for index, valeur in enumerate(valeurs):
            _texte_cellule(ligne.cells[index], valeur if valeur not in (None, "") else "—",
                           centre=index >= 3, taille=8)
        _ombrer(ligne.cells[6], ligne_ind["couleur"].lstrip("#"))
        _texte_cellule(ligne.cells[6], ligne_ind["statut"], gras=True, blanc=True,
                       centre=True, taille=7.5)
        _texte_cellule(ligne.cells[7], ligne_ind["source"] or "—", taille=8)
        _texte_cellule(ligne.cells[8], ligne_ind["responsable"] or "—", taille=8)

    # --- 3. Analyse d'équité
    document.add_page_break()
    titre = document.add_heading("3. Analyse des données désagrégées", level=1)
    titre.runs[0].font.color.rgb = BLEU
    if analyse["desagregation"]:
        document.add_paragraph(
            "Les données collectées sur la période sont ventilées selon les catégories de "
            "désagrégation exigées par le dispositif de suivi. Cette ventilation permet "
            "d'apprécier l'inclusivité effective des interventions.")
        for categorie, bloc in sorted(analyse["desagregation"].items()):
            sous_titre = document.add_heading(f"3.{list(sorted(analyse['desagregation'])).index(categorie) + 1} "
                                              f"Ventilation par « {categorie} »", level=2)
            sous_titre.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
            table = _tableau(document, ["Modalité", "Effectif", "Part (%)"],
                             largeurs=[8, 4, 4])
            total = bloc["total"] or 1
            for modalite, valeur in sorted(bloc["modalites"].items(), key=lambda x: -x[1]):
                ligne = table.add_row()
                _texte_cellule(ligne.cells[0], modalite)
                _texte_cellule(ligne.cells[1], _nombre(valeur), centre=True)
                _texte_cellule(ligne.cells[2], f"{round(valeur / total * 100, 1)} %", centre=True)
            ligne = table.add_row()
            _ombrer(ligne.cells[0], "DCE6F1")
            _texte_cellule(ligne.cells[0], "Total", gras=True)
            _ombrer(ligne.cells[1], "DCE6F1")
            _texte_cellule(ligne.cells[1], _nombre(bloc["total"]), gras=True, centre=True)
            _ombrer(ligne.cells[2], "DCE6F1")
            _texte_cellule(ligne.cells[2], "100 %", gras=True, centre=True)
        if equite:
            p = document.add_paragraph()
            run = p.add_run(
                f"Indice d'équité de genre : les femmes représentent {equite['part_femmes']} % "
                f"des bénéficiaires, soit un écart de {equite['ecart_parite']:+} point(s) par "
                f"rapport à la parité. Appréciation : {equite['appreciation'].lower()}.")
            run.bold = True
            run.font.size = Pt(10)
    else:
        document.add_paragraph(
            "Aucune donnée désagrégée n'a été enregistrée sur la période. Il est recommandé de "
            "renseigner systématiquement la ventilation par sexe et par groupe cible lors de la "
            "saisie des réalisations, conformément au dispositif de suivi.")

    # --- 4. Consolidation par zone
    titre = document.add_heading("4. Consolidation par zone d'intervention", level=1)
    titre.runs[0].font.color.rgb = BLEU
    zones = analyse["zones"]
    zones_actives = [z for z in zones["zones"] if z["nb_mesures"]]
    if zones_actives:
        document.add_paragraph(
            f"{zones['zones_couvertes']} zone(s) sur {zones['nb_zones']} ont fait l'objet d'une "
            f"collecte sur la période, soit un taux de couverture géographique de "
            f"{zones['taux_couverture_zones']} %."
            + (f" {zones['mesures_non_localisees']} mesure(s) n'ont pas été rattachées à une zone."
               if zones["mesures_non_localisees"] else ""))
        table = _tableau(document, ["Code", "Zone", "Niveau", "Bénéficiaires atteints",
                                    "Cible", "Couverture", "Part des femmes", "Mesures"],
                         largeurs=[1.5, 4, 2.2, 2.6, 2, 2, 2.2, 1.5])
        for zone in zones_actives:
            ligne = table.add_row()
            equite_zone = zone.get("equite_genre") or {}
            valeurs = [zone["code"], zone["nom"], zone["niveau"],
                       _nombre(zone["beneficiaires_atteints"]),
                       _nombre(zone["cible_beneficiaires"]),
                       f"{zone['taux_couverture']} %" if zone["taux_couverture"] is not None else "—",
                       f"{equite_zone['part_femmes']} %" if equite_zone else "—",
                       zone["nb_mesures"]]
            for index, valeur in enumerate(valeurs):
                _texte_cellule(ligne.cells[index], valeur if valeur not in (None, "") else "—",
                               centre=index >= 3, taille=8)
        document.add_paragraph()
        p = document.add_paragraph()
        run = p.add_run("Détail des réalisations par zone et par indicateur")
        run.bold = True
        run.font.size = Pt(10)
        table = _tableau(document, ["Zone", "Indicateur", "Unité", "Valeur cumulée"],
                         largeurs=[4, 8, 2, 3])
        for zone in zones_actives:
            for indicateur in zone["indicateurs"]:
                ligne = table.add_row()
                for index, valeur in enumerate([zone["nom"], indicateur["libelle"],
                                                indicateur["unite"] or "—",
                                                _nombre(indicateur["valeur"])]):
                    _texte_cellule(ligne.cells[index], valeur, centre=index == 3, taille=8)
    else:
        document.add_paragraph(
            "Aucune mesure n'a été rattachée à une zone d'intervention sur la période. Le "
            "renseignement de la zone lors de la saisie permet de produire cette consolidation "
            "et d'apprécier l'équilibre géographique des interventions.")

    # --- 5. Exécution physique et financière
    document.add_page_break()
    titre = document.add_heading("5. Exécution physique et financière", level=1)
    titre.runs[0].font.color.rgb = BLEU
    activites_mesurees = analytics.consolidation_par_activite(db, project.id)
    table = _tableau(document, ["Rubrique", "Valeur"], largeurs=[9, 6])
    for libelle, valeur in [
        ("Activités programmées", analyse["activites"]["total"]),
        ("Activités achevées", analyse["activites"]["achevees"]),
        ("Avancement physique moyen", f"{analyse['activites']['avancement_moyen']} %"),
        (f"Budget programmé ({project.currency})", _nombre(analyse["budget"]["planifie"])),
        (f"Montant décaissé ({project.currency})", _nombre(analyse["budget"]["decaisse"])),
        ("Taux d'exécution financière", f"{analyse['budget']['taux_execution']} %"),
    ]:
        ligne = table.add_row()
        _texte_cellule(ligne.cells[0], libelle, gras=True)
        _texte_cellule(ligne.cells[1], valeur, centre=True)
    if activites_mesurees:
        p = document.add_paragraph()
        run = p.add_run("Données collectées par activité")
        run.bold = True
        run.font.size = Pt(10)
        table = _tableau(document, ["Code", "Activité", "Avancement", "Indicateurs renseignés",
                                    "Part des femmes"], largeurs=[1.5, 6, 2, 5, 2.5])
        for activite in activites_mesurees:
            ligne = table.add_row()
            equite_act = activite.get("equite_genre") or {}
            valeurs = [activite["code"], activite["libelle"], f"{activite['avancement']} %",
                       ", ".join(f"{i['code']} = {_nombre(i['valeur'])}"
                                 for i in activite["indicateurs"]) or "—",
                       f"{equite_act['part_femmes']} %" if equite_act else "—"]
            for index, valeur in enumerate(valeurs):
                _texte_cellule(ligne.cells[index], valeur if valeur not in (None, "") else "—",
                               centre=index in (2, 4), taille=8)

    # --- 6. Difficultés et mesures correctrices
    titre = document.add_heading("6. Difficultés rencontrées et mesures correctrices", level=1)
    titre.runs[0].font.color.rgb = BLEU
    if analyse["alertes"]:
        document.add_paragraph(
            "Les écarts significatifs constatés sur la période appellent l'analyse causale et les "
            "mesures correctrices consignées dans le tableau ci-après, à compléter lors de la "
            "réunion de revue de performance.")
        table = _tableau(document, ["Indicateur", "Écart constaté", "Cause identifiée",
                                    "Mesure corrective", "Responsable", "Échéance"],
                         largeurs=[4.5, 2.5, 3.5, 4, 2.5, 2])
        for ligne_ind in analyse["alertes"]:
            ligne = table.add_row()
            _texte_cellule(ligne.cells[0], f"{ligne_ind['code'] or ''} {ligne_ind['name']}",
                           taille=8)
            _texte_cellule(ligne.cells[1],
                           f"{ligne_ind['taux']} % de la cible" if ligne_ind["taux"] is not None
                           else "Non renseigné", centre=True, taille=8)
            for index in (2, 3, 4, 5):
                _texte_cellule(ligne.cells[index], "", taille=8)
    else:
        document.add_paragraph("Aucun écart significatif n'a été constaté sur la période.")

    # --- 7. Qualité des données
    titre = document.add_heading("7. Qualité du dispositif de suivi", level=1)
    titre.runs[0].font.color.rgb = BLEU
    qualite = analytics.synthese_qualite_smart(db, project.id)
    desagregation = analytics.synthese_desagregation(db, project.id, periode)
    document.add_paragraph(
        f"Le score de qualité SMART du système d'indicateurs s'établit à "
        f"{qualite['score_systeme']} % ({qualite['appreciation'].lower()}) : "
        f"{qualite['conformes']} indicateur(s) pleinement conformes et "
        f"{qualite['a_reprendre']} nécessitant une reprise. "
        + (f"Le taux de désagrégation effective des données atteint "
           f"{desagregation['taux_desagregation']} %."
           if desagregation["taux_desagregation"] is not None else ""))
    if qualite["a_reprendre"]:
        table = _tableau(document, ["Indicateur", "Score", "Actions correctrices recommandées"],
                         largeurs=[5, 1.8, 10])
        for ligne_ind in [l for l in qualite["lignes"] if l["score"] < 60][:12]:
            ligne = table.add_row()
            _texte_cellule(ligne.cells[0], f"{ligne_ind['code'] or ''} {ligne_ind['name']}",
                           taille=8)
            _texte_cellule(ligne.cells[1], f"{ligne_ind['score']} %", centre=True, taille=8)
            _texte_cellule(ligne.cells[2], " • ".join(ligne_ind["recommandations"]), taille=8)

    # --- 8. Validation
    titre = document.add_heading("8. Validation du rapport", level=1)
    titre.runs[0].font.color.rgb = BLEU
    table = _tableau(document, ["Fonction", "Nom et prénoms", "Date", "Signature"],
                     largeurs=[4.5, 5, 3, 4])
    for fonction in ["Rédigé par — Responsable suivi-évaluation",
                     "Vérifié par — Coordonnateur du projet",
                     "Approuvé par — Président du comité de pilotage"]:
        ligne = table.add_row()
        _texte_cellule(ligne.cells[0], fonction, gras=True)
        for index in (1, 2, 3):
            _texte_cellule(ligne.cells[index], "")
        ligne.cells[0].paragraphs[0].paragraph_format.space_after = Pt(10)

    _pied_de_page(document, f"{project.code} — {LIBELLES_RAPPORT.get(type_rapport, 'Rapport')} — "
                            f"{periode} — Généré par {APP_NAME}")
    return _sauver(document)


# ---------------------------------------------------------------------------
# 7. Rapport périodique de performance
# ---------------------------------------------------------------------------
def rapport_performance_docx(db: Session, project: Project, periode: str = "") -> BytesIO:
    document = Document()
    section = document.sections[0]
    section.left_margin = section.right_margin = Cm(2)
    tdb = analytics.tableau_de_bord(db, project.id)
    _en_tete_document(document, project, "RAPPORT DE PERFORMANCE",
                      f"Période de référence : {periode or 'depuis le démarrage'}")

    heading = document.add_heading("1. Résumé exécutif", level=1)
    heading.runs[0].font.color.rgb = BLEU
    sante = tdb["sante_globale"]
    document.add_paragraph(
        f"À la date du {date.today().strftime('%d/%m/%Y')}, l'indice de santé global du projet "
        f"s'établit à {sante['score']}% ({sante['statut']}). Le taux moyen de réalisation des "
        f"indicateurs atteint {tdb['indicateurs']['taux_moyen'] or 0}%, pour un avancement "
        f"physique moyen des activités de {tdb['activites']['avancement_moyen']}% et un taux "
        f"d'exécution budgétaire de {tdb['budget']['taux_execution']}%. "
        f"{tdb['temps']['taux_temps'] or 0}% de la durée du projet est écoulée.")
    if sante.get("ecart_calendrier") is not None:
        ecart = sante["ecart_calendrier"]
        appreciation = ("en avance sur le calendrier" if ecart > 5 else
                        "conforme au calendrier" if ecart >= -5 else "en retard sur le calendrier")
        document.add_paragraph(
            f"L'écart entre la performance réalisée et le temps consommé est de {ecart} points : "
            f"le projet est {appreciation}.")

    heading = document.add_heading("2. Indicateurs de performance", level=1)
    heading.runs[0].font.color.rgb = BLEU
    table = _tableau(document, ["Code", "Indicateur", "Référence", "Cible", "Réalisé", "Taux",
                                "Statut"], largeurs=[1.4, 6.5, 1.8, 1.8, 1.8, 1.5, 2.2])
    for ligne_ind in tdb["indicateurs"]["lignes"]:
        ligne = table.add_row()
        valeurs = [ligne_ind["code"], ligne_ind["name"], _nombre(ligne_ind["baseline_value"]),
                   _nombre(ligne_ind["target_value"]), _nombre(ligne_ind["actual_value"]),
                   f"{ligne_ind['taux']}%" if ligne_ind["taux"] is not None else "—"]
        for index, valeur in enumerate(valeurs):
            _texte_cellule(ligne.cells[index], valeur if valeur not in (None, "") else "—",
                           centre=index > 1, taille=8)
        _ombrer(ligne.cells[6], ligne_ind["couleur"].lstrip("#"))
        _texte_cellule(ligne.cells[6], ligne_ind["statut"], gras=True, blanc=True, centre=True, taille=8)

    heading = document.add_heading("3. Exécution physique et financière", level=1)
    heading.runs[0].font.color.rgb = BLEU
    table = _tableau(document, ["Rubrique", "Valeur"], largeurs=[9, 6])
    for libelle, valeur in [
        ("Activités programmées", tdb["activites"]["total"]),
        ("Activités achevées", tdb["activites"]["achevees"]),
        ("Activités en retard", tdb["activites"]["nb_en_retard"]),
        ("Avancement physique moyen", f"{tdb['activites']['avancement_moyen']} %"),
        (f"Budget planifié ({project.currency})", _nombre(tdb["budget"]["planifie"])),
        (f"Montant décaissé ({project.currency})", _nombre(tdb["budget"]["decaisse"])),
        ("Taux d'exécution financière", f"{tdb['budget']['taux_execution']} %"),
    ]:
        ligne = table.add_row()
        _texte_cellule(ligne.cells[0], libelle, gras=True)
        _texte_cellule(ligne.cells[1], valeur, centre=True)

    heading = document.add_heading("4. Points d'attention et alertes", level=1)
    heading.runs[0].font.color.rgb = BLEU
    if tdb["alertes"]:
        for alerte in tdb["alertes"][:25]:
            p = document.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{alerte['type']}] ")
            run.bold = True
            run.font.color.rgb = RGBColor(0xD9, 0x30, 0x25) if alerte["niveau"] == "danger" \
                else RGBColor(0xEA, 0x86, 0x00)
            p.add_run(alerte["message"])
    else:
        document.add_paragraph("Aucune alerte critique n'est signalée sur la période.")

    heading = document.add_heading("5. Mesures correctrices proposées", level=1)
    heading.runs[0].font.color.rgb = BLEU
    table = _tableau(document, ["Constat", "Mesure corrective", "Responsable", "Échéance"],
                     largeurs=[5.5, 5.5, 3, 2.5])
    for alerte in tdb["alertes"][:10]:
        ligne = table.add_row()
        _texte_cellule(ligne.cells[0], alerte["message"], taille=8)
        for index in (1, 2, 3):
            _texte_cellule(ligne.cells[index], "", taille=8)
    document.add_paragraph()
    p = document.add_paragraph()
    run = p.add_run("Tableau à compléter lors de la réunion de revue de performance.")
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = GRIS
    _pied_de_page(document, f"{project.code} — Rapport de performance — Généré par {APP_NAME}")
    return _sauver(document)
