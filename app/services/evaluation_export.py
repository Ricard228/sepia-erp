"""Livrables des modules bénéficiaires, partenaires, évaluation CAD et impact."""
from datetime import date
from io import BytesIO
from typing import Any, List

import xlsxwriter
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from sqlalchemy.orm import Session

from ..config import APP_NAME, CRITERES_CAD, ECHELLE_NOTATION_CAD
from ..models import Project
from . import evaluation as service
from .excel_export import _entete_feuille, _formats
from .word_export import (_date_iso_fr, _en_tete_document, _nombre, _ombrer, _paysage,
                          _pied_de_page, _sauver, _tableau, _texte_cellule, BLEU, GRIS)


# ---------------------------------------------------------------------------
# Bénéficiaires
# ---------------------------------------------------------------------------
def beneficiaires_xlsx(db: Session, project: Project) -> BytesIO:
    synthese = service.synthese_beneficiaires(db, project.id)
    buffer = BytesIO()
    wb = xlsxwriter.Workbook(buffer, {"in_memory": True})
    fmt = _formats(wb)

    ws = wb.add_worksheet("Bénéficiaires")
    ws.set_landscape()
    ws.set_paper(8)
    ligne = _entete_feuille(ws, fmt, project, "GROUPES DE BÉNÉFICIAIRES — CIBLAGE ET ATTEINTE")
    entetes = ["Code", "Groupe", "Catégorie", "Typologie", "Zone", "Vulnérabilité",
               "Ciblé", "Dont femmes", "Dont jeunes", "Dont handicap",
               "Atteint", "Dont femmes", "Dont jeunes", "Dont handicap",
               "Taux d'atteinte (%)", "Part des femmes (%)", "Ménages",
               "Taille du ménage", "Revenu de référence", "Taux de pauvreté (%)",
               "Indicateurs rattachés"]
    largeurs = [8, 34, 12, 20, 20, 14, 11, 12, 12, 12, 11, 12, 12, 12, 14, 14, 11, 14, 16, 14, 14]
    for col, (titre, largeur) in enumerate(zip(entetes, largeurs)):
        ws.write(ligne, col, titre, fmt["entete"])
        ws.set_column(col, col, largeur)
    ws.set_row(ligne, 34)
    ligne += 1
    depart = ligne
    for groupe in synthese["groupes"]:
        valeurs = [groupe["code"] or "", groupe["nom"], groupe["categorie"] or "",
                   groupe["typologie"] or "", groupe["zone"] or "", groupe["vulnerabilite"] or "",
                   groupe["cible_total"], groupe["cible_femmes"], groupe["cible_jeunes"],
                   groupe["cible_handicap"], groupe["atteint_total"], groupe["atteint_femmes"],
                   groupe["atteint_jeunes"], groupe["atteint_handicap"], groupe["taux_atteinte"],
                   groupe["part_femmes_atteintes"], groupe["menages"], groupe["taille_menage"],
                   groupe["revenu_reference"], groupe["taux_pauvrete"], groupe["nb_indicateurs"]]
        for col, valeur in enumerate(valeurs):
            style = fmt["cellule"] if col < 6 else (
                fmt["cellule_c"] if col in (14, 15, 20) else fmt["nombre"])
            ws.write(ligne, col, valeur if valeur is not None else "", style)
        ligne += 1
    if synthese["groupes"]:
        ws.autofilter(depart - 1, 0, ligne - 1, len(entetes) - 1)
        ws.conditional_format(depart, 14, ligne - 1, 14, {
            "type": "3_color_scale", "min_color": "#F4C7C3", "mid_color": "#FCE8B2",
            "max_color": "#B7E1CD"})
    ligne += 1
    for libelle, valeur in [
        ("Effectif total ciblé", synthese["cible_totale"]),
        ("Effectif total atteint", synthese["atteint_total"]),
        ("Taux d'atteinte global (%)", synthese["taux_atteinte_global"]),
        ("Part des femmes ciblée (%)", synthese["part_femmes_ciblee"]),
        ("Part des femmes atteinte (%)", synthese["part_femmes_atteinte"]),
        ("Ménages concernés", synthese["menages"]),
        ("Personnes touchées estimées", synthese["personnes_touchees_estimees"]),
        ("Indicateurs rattachés à un groupe (%)", synthese["taux_rattachement"]),
    ]:
        ws.write(ligne, 0, libelle, fmt["gras"])
        ws.write(ligne, 1, valeur if valeur is not None else "—", fmt["nombre"])
        ligne += 1

    # Caractérisation qualitative, sur une feuille dédiée pour rester lisible.
    ws2 = wb.add_worksheet("Caractérisation")
    l2 = _entete_feuille(ws2, fmt, project, "CARACTÉRISATION QUALITATIVE DES BÉNÉFICIAIRES")
    colonnes = ["Code", "Groupe", "Critères d'éligibilité", "Méthode de ciblage", "Besoins",
                "Contraintes d'accès", "Bénéfices attendus", "Participation",
                "Mécanisme de plainte", "Observations"]
    for col, (titre, largeur) in enumerate(zip(colonnes, [8, 30, 46, 40, 44, 44, 38, 38, 38, 30])):
        ws2.write(l2, col, titre, fmt["entete"])
        ws2.set_column(col, col, largeur)
    ws2.set_row(l2, 32)
    l2 += 1
    for groupe in synthese["groupes"]:
        valeurs = [groupe["code"] or "", groupe["nom"], groupe["criteres_selection"] or "",
                   groupe["methode_ciblage"] or "", groupe["besoins"] or "",
                   groupe["contraintes"] or "", groupe["benefices_attendus"] or "",
                   groupe["participation"] or "", groupe["mecanisme_plainte"] or "",
                   groupe["commentaire"] or ""]
        for col, valeur in enumerate(valeurs):
            ws2.write(l2, col, valeur, fmt["cellule"])
        ws2.set_row(l2, 62)
        l2 += 1

    # Indicateurs rattachés à chaque groupe.
    ws3 = wb.add_worksheet("Indicateurs par groupe")
    l3 = _entete_feuille(ws3, fmt, project, "INDICATEURS RATTACHÉS AUX GROUPES DE BÉNÉFICIAIRES")
    for col, (titre, largeur) in enumerate(zip(
            ["Groupe", "Code indicateur", "Indicateur", "Unité", "Référence", "Cible",
             "Réalisé", "Taux (%)", "Statut"], [30, 14, 48, 12, 12, 12, 12, 11, 16])):
        ws3.write(l3, col, titre, fmt["entete"])
        ws3.set_column(col, col, largeur)
    l3 += 1
    for groupe in synthese["groupes"]:
        for indicateur in groupe["indicateurs"]:
            valeurs = [groupe["nom"], indicateur["code"] or "", indicateur["name"],
                       indicateur["unit"] or "", indicateur["baseline_value"],
                       indicateur["target_value"], indicateur["actual_value"],
                       indicateur["taux"], indicateur["statut"]]
            for col, valeur in enumerate(valeurs):
                ws3.write(l3, col, valeur if valeur is not None else "",
                          fmt["cellule"] if col < 4 else fmt["cellule_c"])
            l3 += 1
    wb.close()
    buffer.seek(0)
    return buffer


def partenaires_xlsx(db: Session, project: Project) -> BytesIO:
    synthese = service.synthese_partenaires(db, project.id)
    buffer = BytesIO()
    wb = xlsxwriter.Workbook(buffer, {"in_memory": True})
    fmt = _formats(wb)
    ws = wb.add_worksheet("Partenaires")
    ws.set_landscape()
    ws.set_paper(8)
    ligne = _entete_feuille(ws, fmt, project, "PARTENAIRES DU PROJET — ENGAGEMENTS ET PERFORMANCE")
    entetes = ["Code", "Partenaire", "Type", "Pays", "Rôle", "Convention", "Début", "Fin",
               "Contribution", f"Engagé ({project.currency})", f"Versé ({project.currency})",
               "Décaissement (%)", "Obligations", "Livrables", "Appréciation", "Commentaire",
               "Risques", "Contact", "Courriel", "Statut"]
    largeurs = [8, 34, 22, 14, 40, 20, 12, 12, 16, 16, 16, 13, 38, 32, 12, 34, 34, 22, 24, 12]
    for col, (titre, largeur) in enumerate(zip(entetes, largeurs)):
        ws.write(ligne, col, titre, fmt["entete"])
        ws.set_column(col, col, largeur)
    ws.set_row(ligne, 34)
    ligne += 1
    depart = ligne
    for p in synthese["partenaires"]:
        valeurs = [p["code"] or "", p["nom"], p["type"] or "", p["pays"] or "", p["role"] or "",
                   p["convention"] or "", p["debut"] or "", p["fin"] or "",
                   p["type_contribution"] or "", p["engage"], p["verse"], p["taux_decaissement"],
                   p["obligations"] or "", p["livrables"] or "", p["note"],
                   p["appreciation"] or "", p["risques"] or "", p["contact"] or "",
                   p["courriel"] or "", p["statut"] or ""]
        for col, valeur in enumerate(valeurs):
            style = fmt["nombre"] if col in (9, 10) else (
                fmt["cellule_c"] if col in (6, 7, 11, 14, 19) else fmt["cellule"])
            ws.write(ligne, col, valeur if valeur is not None else "", style)
        ws.set_row(ligne, 46)
        ligne += 1
    if synthese["partenaires"]:
        ws.autofilter(depart - 1, 0, ligne - 1, len(entetes) - 1)
    ligne += 1
    for libelle, valeur in [
        (f"Engagement total ({project.currency})", synthese["engagement_total"]),
        (f"Montant versé ({project.currency})", synthese["verse_total"]),
        ("Taux de décaissement global (%)", synthese["taux_decaissement_global"]),
        ("Appréciation moyenne (sur 5)", synthese["note_moyenne"]),
        ("Partenaires actifs", synthese["actifs"]),
    ]:
        ws.write(ligne, 0, libelle, fmt["gras"])
        ws.write(ligne, 1, valeur if valeur is not None else "—", fmt["nombre"])
        ligne += 1
    wb.close()
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# Évaluation CAD-OCDE
# ---------------------------------------------------------------------------
def evaluation_cad_xlsx(db: Session, project: Project) -> BytesIO:
    synthese = service.synthese_evaluations(db, project.id)
    buffer = BytesIO()
    wb = xlsxwriter.Workbook(buffer, {"in_memory": True})
    fmt = _formats(wb)

    ws = wb.add_worksheet("Notation CAD")
    ws.set_landscape()
    ligne = _entete_feuille(ws, fmt, project,
                            "ÉVALUATION SELON LES CRITÈRES DU CAD DE L'OCDE")
    ws.set_column(0, 0, 34)
    ws.set_column(1, 8, 18)
    ws.write(ligne, 0, "Évaluation", fmt["entete"])
    for col, critere in enumerate(CRITERES_CAD, start=1):
        ws.write(ligne, col, critere["libelle"], fmt["entete"])
    ws.write(ligne, len(CRITERES_CAD) + 1, "Note globale", fmt["entete"])
    ws.write(ligne, len(CRITERES_CAD) + 2, "Appréciation", fmt["entete"])
    ligne += 1
    couleurs = {e["note"]: e["couleur"] for e in ECHELLE_NOTATION_CAD}
    for evaluation in synthese["evaluations"]:
        ws.write(ligne, 0, f"{evaluation['code'] or ''} {evaluation['titre']}", fmt["cellule"])
        for col, critere in enumerate(evaluation["criteres"], start=1):
            if critere["note"] is None:
                ws.write(ligne, col, "—", fmt["cellule_c"])
                continue
            couleur = couleurs.get(int(round(critere["note"])), "#9AA0A6")
            ws.write(ligne, col, critere["note"], wb.add_format({
                "border": 1, "align": "center", "bold": True, "font_color": "white",
                "bg_color": couleur}))
        ws.write(ligne, len(CRITERES_CAD) + 1,
                 evaluation["note_globale"] if evaluation["note_globale"] is not None else "—",
                 fmt["cellule_c"])
        ws.write(ligne, len(CRITERES_CAD) + 2, evaluation["globale_libelle"], fmt["cellule"])
        ligne += 1

    ligne += 2
    ws.write(ligne, 0, "Échelle de notation", fmt["entete"])
    for index, entree in enumerate(ECHELLE_NOTATION_CAD):
        ws.write(ligne + 1 + index, 0, f"{entree['note']} — {entree['libelle']}",
                 wb.add_format({"border": 1, "bg_color": entree["couleur"],
                                "font_color": "white", "bold": True}))

    ws2 = wb.add_worksheet("Justifications")
    l2 = _entete_feuille(ws2, fmt, project, "JUSTIFICATION DES NOTES PAR CRITÈRE")
    for col, (titre, largeur) in enumerate(zip(
            ["Évaluation", "Critère", "Question évaluative", "Note", "Appréciation",
             "Justification"], [30, 20, 50, 8, 18, 70])):
        ws2.write(l2, col, titre, fmt["entete"])
        ws2.set_column(col, col, largeur)
    l2 += 1
    for evaluation in synthese["evaluations"]:
        for critere in evaluation["criteres"]:
            valeurs = [f"{evaluation['code'] or ''} {evaluation['titre']}", critere["libelle"],
                       critere["question"],
                       critere["note"] if critere["note"] is not None else "—",
                       critere["libelle_note"] if critere.get("libelle_note") else "",
                       critere["justification"] or ""]
            for col, valeur in enumerate(valeurs):
                ws2.write(l2, col, valeur, fmt["cellule_c"] if col == 3 else fmt["cellule"])
            ws2.set_row(l2, 58)
            l2 += 1

    ws3 = wb.add_worksheet("Recommandations")
    l3 = _entete_feuille(ws3, fmt, project, "SUIVI DES RECOMMANDATIONS D'ÉVALUATION")
    for col, (titre, largeur) in enumerate(zip(
            ["Évaluation", "Code", "Critère", "Recommandation", "Priorité", "Responsable",
             "Échéance", "Réponse de la direction", "Statut", "Mise en œuvre (%)", "Preuve"],
            [26, 8, 18, 56, 12, 24, 13, 22, 18, 16, 34])):
        ws3.write(l3, col, titre, fmt["entete"])
        ws3.set_column(col, col, largeur)
    l3 += 1
    depart3 = l3
    for evaluation in synthese["evaluations"]:
        for reco in evaluation["recommandations"]:
            valeurs = [evaluation["code"] or "", reco["code"] or "", reco["critere"] or "",
                       reco["enonce"], reco["priorite"] or "", reco["responsable"] or "",
                       reco["echeance"] or "", reco["reponse_management"] or "",
                       reco["statut"] or "", reco["taux"], reco["preuve"] or ""]
            for col, valeur in enumerate(valeurs):
                ws3.write(l3, col, valeur,
                          fmt["cellule_c"] if col in (4, 6, 7, 8, 9) else fmt["cellule"])
            l3 += 1
    if l3 > depart3:
        ws3.autofilter(depart3 - 1, 0, l3 - 1, 10)
        ws3.conditional_format(depart3, 9, l3 - 1, 9, {
            "type": "3_color_scale", "min_color": "#F4C7C3", "mid_color": "#FCE8B2",
            "max_color": "#B7E1CD"})
    wb.close()
    buffer.seek(0)
    return buffer


def evaluation_cad_docx(db: Session, project: Project) -> BytesIO:
    synthese = service.synthese_evaluations(db, project.id)
    document = Document()
    section = document.sections[0]
    section.left_margin = section.right_margin = Cm(2)
    _en_tete_document(document, project, "RAPPORT D'ÉVALUATION",
                      "Appréciation selon les six critères du Comité d'aide au développement "
                      "de l'OCDE")

    titre = document.add_heading("1. Cadre d'évaluation", level=1)
    titre.runs[0].font.color.rgb = BLEU
    document.add_paragraph(
        "L'évaluation est conduite selon les six critères du CAD de l'OCDE, révisés en 2019. "
        "Chaque critère est apprécié sur une échelle de 1 à 6 et assorti d'une justification "
        "écrite : une note sans justification n'est ni contestable ni réutilisable.")
    table = _tableau(document, ["Critère", "Question évaluative", "Points d'examen"],
                     largeurs=[3, 6.5, 7])
    for critere in CRITERES_CAD:
        ligne = table.add_row()
        _texte_cellule(ligne.cells[0], critere["libelle"], gras=True, taille=8.5)
        _texte_cellule(ligne.cells[1], critere["question"], taille=8.5)
        _texte_cellule(ligne.cells[2], critere["points_examen"], taille=8.5)

    document.add_paragraph()
    table = _tableau(document, ["Note", "Appréciation"], largeurs=[3, 6])
    for entree in ECHELLE_NOTATION_CAD:
        ligne = table.add_row()
        _ombrer(ligne.cells[0], entree["couleur"].lstrip("#"))
        _texte_cellule(ligne.cells[0], str(entree["note"]), gras=True, blanc=True, centre=True)
        _texte_cellule(ligne.cells[1], entree["libelle"])

    if not synthese["evaluations"]:
        document.add_paragraph(
            "Aucun exercice évaluatif n'est encore enregistré pour ce projet.")
        _pied_de_page(document, f"{project.code} — Rapport d'évaluation — {APP_NAME}")
        return _sauver(document)

    titre = document.add_heading("2. Synthèse des exercices évaluatifs", level=1)
    titre.runs[0].font.color.rgb = BLEU
    document.add_paragraph(
        f"{synthese['total']} exercice(s) évaluatif(s) sont enregistrés, dont "
        f"{synthese['achevees']} achevé(s) ou validé(s). La note moyenne s'établit à "
        f"{synthese['note_moyenne'] if synthese['note_moyenne'] is not None else '—'} sur 6. "
        f"{synthese['nb_recommandations']} recommandation(s) ont été formulées, mises en œuvre "
        f"à {synthese['taux_mise_en_oeuvre'] if synthese['taux_mise_en_oeuvre'] is not None else '—'} % "
        f"en moyenne.")
    table = _tableau(document, ["Critère", "Note moyenne", "Appréciation", "Évaluations notées"],
                     largeurs=[5, 3, 4.5, 4])
    for libelle, valeurs in synthese["par_critere"].items():
        ligne = table.add_row()
        _texte_cellule(ligne.cells[0], libelle, gras=True)
        if valeurs["moyenne"] is None:
            _texte_cellule(ligne.cells[1], "—", centre=True)
            _texte_cellule(ligne.cells[2], "Non noté", centre=True)
        else:
            _ombrer(ligne.cells[1], valeurs["couleur"].lstrip("#"))
            _texte_cellule(ligne.cells[1], f"{valeurs['moyenne']} / 6", gras=True, blanc=True,
                           centre=True)
            _texte_cellule(ligne.cells[2], valeurs["libelle_note"], centre=True)
        _texte_cellule(ligne.cells[3], valeurs["nb_evaluations"], centre=True)

    for index, evaluation in enumerate(synthese["evaluations"], start=3):
        document.add_page_break()
        titre = document.add_heading(
            f"{index}. {evaluation['code'] or ''} — {evaluation['titre']}", level=1)
        titre.runs[0].font.color.rgb = BLEU
        table = _tableau(document, ["Rubrique", "Information"], largeurs=[5, 11])
        for libelle, valeur in [
            ("Type d'évaluation", evaluation["type"]),
            ("Période couverte", evaluation["periode"]),
            ("Statut", evaluation["statut"]),
            ("Évaluateur", evaluation["evaluateur"]),
            ("Degré d'indépendance", evaluation["independance"]),
            ("Période de réalisation", f"{_date_iso_fr(evaluation['date_debut'])} — "
                                       f"{_date_iso_fr(evaluation['date_fin'])}"),
            ("Budget", _nombre(evaluation["budget"])),
            ("Référence du rapport", evaluation["rapport"]),
            ("Note globale", f"{evaluation['note_globale']} / 6 — "
                             f"{evaluation['globale_libelle']}"
             if evaluation["note_globale"] is not None else "Non noté"),
        ]:
            ligne = table.add_row()
            _texte_cellule(ligne.cells[0], libelle, gras=True)
            _texte_cellule(ligne.cells[1], valeur or "—")

        for libelle, texte in [("Méthodologie", evaluation["methodologie"]),
                               ("Sources de données", evaluation["sources"]),
                               ("Échantillonnage", evaluation["echantillonnage"]),
                               ("Limites de l'évaluation", evaluation["limites"])]:
            if texte:
                sous = document.add_heading(libelle, level=2)
                sous.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
                document.add_paragraph(texte)

        sous = document.add_heading("Appréciation par critère", level=2)
        sous.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        table = _tableau(document, ["Critère", "Note", "Appréciation", "Justification"],
                         largeurs=[3, 1.8, 3, 8.5])
        for critere in evaluation["criteres"]:
            ligne = table.add_row()
            _texte_cellule(ligne.cells[0], critere["libelle"], gras=True, taille=8.5)
            if critere["note"] is None:
                _texte_cellule(ligne.cells[1], "—", centre=True, taille=8.5)
                _texte_cellule(ligne.cells[2], "Non noté", centre=True, taille=8.5)
            else:
                _ombrer(ligne.cells[1], critere["couleur"].lstrip("#"))
                _texte_cellule(ligne.cells[1], f"{critere['note']}/6", gras=True, blanc=True,
                               centre=True, taille=8.5)
                _texte_cellule(ligne.cells[2], critere["libelle_note"], centre=True, taille=8.5)
            _texte_cellule(ligne.cells[3], critere["justification"] or "—", taille=8.5)

        for libelle, texte in [("Constats principaux", evaluation["constats"]),
                               ("Leçons apprises", evaluation["lecons"]),
                               ("Appréciation générale", evaluation["appreciation_generale"])]:
            if texte:
                sous = document.add_heading(libelle, level=2)
                sous.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
                document.add_paragraph(texte)

        if evaluation["recommandations"]:
            sous = document.add_heading("Recommandations et réponse de la direction", level=2)
            sous.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
            table = _tableau(document, ["Code", "Critère", "Recommandation", "Priorité",
                                        "Réponse", "Responsable", "Échéance", "Mise en œuvre"],
                             largeurs=[1.2, 2, 5.5, 1.6, 2, 2.4, 1.8, 2])
            for reco in evaluation["recommandations"]:
                ligne = table.add_row()
                valeurs = [reco["code"], reco["critere"], reco["enonce"], reco["priorite"],
                           reco["reponse_management"], reco["responsable"],
                           _date_iso_fr(reco["echeance"]), f"{_nombre(reco['taux'])} %"]
                for col, valeur in enumerate(valeurs):
                    _texte_cellule(ligne.cells[col], valeur or "—",
                                   centre=col in (3, 4, 6, 7), taille=8)

    _pied_de_page(document, f"{project.code} — Rapport d'évaluation CAD-OCDE — {APP_NAME}")
    return _sauver(document)


# ---------------------------------------------------------------------------
# Évaluation d'impact
# ---------------------------------------------------------------------------
def evaluation_impact_docx(db: Session, project: Project) -> BytesIO:
    synthese = service.synthese_impact(db, project.id)
    document = Document()
    section = document.sections[0]
    section.left_margin = section.right_margin = Cm(2)
    _en_tete_document(document, project, "PROTOCOLE D'ÉVALUATION D'IMPACT",
                      "Devis expérimentaux et quasi-expérimentaux, puissance statistique "
                      "et résultats")

    titre = document.add_heading("1. Objet et principe", level=1)
    titre.runs[0].font.color.rgb = BLEU
    document.add_paragraph(
        "L'évaluation d'impact vise à établir un lien causal entre l'intervention et les "
        "changements observés. Elle se distingue du suivi, qui décrit ce qui s'est passé, et de "
        "l'évaluation par critères, qui porte un jugement de qualité. Son enjeu propre est la "
        "reconstitution du contrefactuel : ce qui serait advenu en l'absence de l'intervention.")
    document.add_paragraph(
        "La comparaison des bénéficiaires avant et après ne suffit pas : elle attribue à "
        "l'intervention tout ce qui a changé entre les deux dates, y compris les effets du "
        "contexte. Un groupe de comparaison crédible est donc indispensable, et la méthode "
        "retenue détermine sous quelle hypothèse l'effet estimé peut être interprété comme "
        "causal.")

    titre = document.add_heading("2. Méthodes disponibles et conditions d'emploi", level=1)
    titre.runs[0].font.color.rgb = BLEU
    table = _tableau(document, ["Méthode", "Approche", "Hypothèse d'identification",
                                "Conditions", "Forces", "Limites"],
                     largeurs=[3, 2.2, 4.5, 4, 3.5, 3.8])
    for methode in synthese["methodes"]:
        ligne = table.add_row()
        _texte_cellule(ligne.cells[0], methode["libelle"], gras=True, taille=7.5)
        _texte_cellule(ligne.cells[1], methode["approche"], centre=True, taille=7.5)
        _texte_cellule(ligne.cells[2], methode["hypothese"], taille=7.5)
        _texte_cellule(ligne.cells[3], methode["conditions"], taille=7.5)
        _texte_cellule(ligne.cells[4], methode["forces"], taille=7.5)
        _texte_cellule(ligne.cells[5], methode["limites"], taille=7.5)

    if not synthese["etudes"]:
        document.add_paragraph(
            "Aucune étude d'impact n'est enregistrée pour ce projet. Le devis doit être conçu "
            "avant la mesure de référence : une fois les activités engagées, la constitution "
            "d'un contrefactuel crédible devient nettement plus difficile.")
        _pied_de_page(document, f"{project.code} — Protocole d'évaluation d'impact — {APP_NAME}")
        return _sauver(document)

    for index, etude in enumerate(synthese["etudes"], start=3):
        document.add_page_break()
        titre = document.add_heading(f"{index}. {etude['code'] or ''} — {etude['titre']}", level=1)
        titre.runs[0].font.color.rgb = BLEU

        sous = document.add_heading(f"{index}.1 Question et devis", level=2)
        sous.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        table = _tableau(document, ["Rubrique", "Contenu"], largeurs=[4.5, 11.5])
        for libelle, valeur in [
            ("Question de recherche", etude["question"]),
            ("Hypothèse testée", etude["hypothese"]),
            ("Approche", etude["approche"]),
            ("Méthode d'identification", etude["methode"]),
            ("Hypothèse d'identification", etude["hypothese_identification"]),
            ("Règle d'affectation au traitement", etude["regle_affectation"]),
            ("Unité d'analyse", etude["unite_analyse"]),
            ("Indicateurs de résultat", ", ".join(etude["codes_indicateurs"]) or None),
            ("Variables de contrôle", etude["covariables"]),
        ]:
            ligne = table.add_row()
            _texte_cellule(ligne.cells[0], libelle, gras=True, taille=9)
            _texte_cellule(ligne.cells[1], valeur or "—", taille=9)

        sous = document.add_heading(f"{index}.2 Échantillon et puissance statistique", level=2)
        sous.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        table = _tableau(document, ["Paramètre", "Valeur", "Paramètre", "Valeur"],
                         largeurs=[4, 3.5, 4.5, 4])
        parametres = [
            ("Groupe de traitement", _nombre(etude["traitement"])),
            ("Groupe de contrôle", _nombre(etude["controle"])),
            ("Échantillon total", _nombre(etude["echantillon_total"])),
            ("Nombre de grappes", _nombre(etude["grappes"])),
            ("Corrélation intra-grappe", _nombre(etude["correlation_intra"])),
            ("Effet minimal détectable", _nombre(etude["effet_minimal_detectable"])),
            ("Écart-type de l'indicateur", _nombre(etude["ecart_type_resultat"])),
            ("Puissance visée", _nombre(etude["puissance"])),
            ("Seuil de signification", _nombre(etude["seuil"])),
            ("Attrition anticipée", _nombre(etude["attrition"])),
            ("Statut", etude["statut"]),
        ]
        for position in range(0, len(parametres), 2):
            ligne = table.add_row()
            _texte_cellule(ligne.cells[0], parametres[position][0], gras=True, taille=9)
            _texte_cellule(ligne.cells[1], parametres[position][1], centre=True, taille=9)
            if position + 1 < len(parametres):
                _texte_cellule(ligne.cells[2], parametres[position + 1][0], gras=True, taille=9)
                _texte_cellule(ligne.cells[3], parametres[position + 1][1], centre=True, taille=9)

        controle = etude.get("controle_puissance")
        if controle and controle.get("indisponible"):
            p = document.add_paragraph()
            run = p.add_run("Contrôle de puissance indisponible : " + controle["indisponible"])
            run.italic = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xEA, 0x86, 0x00)
        elif controle:
            p = document.add_paragraph()
            run = p.add_run(
                f"Contrôle de puissance : l'échantillon prévu de "
                f"{_nombre(controle['n_prevu'])} unités est comparé au minimum requis de "
                f"{_nombre(controle['n_requis'])} unités pour détecter l'effet minimal retenu "
                f"(effet de plan {controle['effet_de_plan']}). "
                + ("L'échantillon est suffisant." if controle["suffisant"] else
                   "L'échantillon est insuffisant : l'étude risque de conclure à l'absence "
                   "d'effet alors qu'un effet réel existe."))
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = (RGBColor(0x0F, 0x9D, 0x58) if controle["suffisant"]
                                  else RGBColor(0xD9, 0x30, 0x25))

        sous = document.add_heading(f"{index}.3 Calendrier, résultats et validité", level=2)
        sous.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        table = _tableau(document, ["Rubrique", "Contenu"], largeurs=[4.5, 11.5])
        for libelle, valeur in [
            ("Mesure de référence", _date_iso_fr(etude["date_baseline"])),
            ("Mesure intermédiaire", _date_iso_fr(etude["date_intermediaire"])),
            ("Mesure finale", _date_iso_fr(etude["date_finale"])),
            ("Effet estimé", f"{_nombre(etude['effet_estime'])} {etude['unite_effet'] or ''}"
             if etude["effet_estime"] is not None else None),
            ("Erreur type", _nombre(etude["erreur_type"])),
            ("Valeur p", _nombre(etude["p_value"])),
            ("Intervalle de confiance", etude["intervalle_confiance"]),
            ("Signification statistique",
             "Effet significatif au seuil retenu" if etude["significatif"] is True else
             "Effet non significatif" if etude["significatif"] is False else None),
            ("Tests de robustesse", etude["tests_robustesse"]),
            ("Menaces sur la validité", etude["menaces_validite"]),
            ("Conclusion", etude["conclusion"]),
            ("Avis éthique et consentement", etude["ethique"]),
            ("Dépôt des données", etude["depot_donnees"]),
        ]:
            ligne = table.add_row()
            _texte_cellule(ligne.cells[0], libelle, gras=True, taille=9)
            _texte_cellule(ligne.cells[1], valeur or "—", taille=9)

    document.add_page_break()
    titre = document.add_heading("Précautions d'interprétation", level=1)
    titre.runs[0].font.color.rgb = BLEU
    for texte in [
        "Un effet non significatif n'est pas la preuve d'une absence d'effet : il peut résulter "
        "d'un échantillon insuffisant. Le contrôle de puissance doit être présenté avec le "
        "résultat.",
        "L'effet estimé vaut pour la population et la période étudiées. Sa transposition à un "
        "autre contexte relève d'un jugement argumenté, non d'une déduction statistique.",
        "L'hypothèse d'identification de la méthode retenue doit être discutée explicitement : "
        "c'est elle, et non la technique d'estimation, qui fonde l'interprétation causale.",
        "Les tests de robustesse et l'analyse de l'attrition font partie intégrante du résultat. "
        "Un résultat présenté sans eux ne peut être considéré comme établi.",
        "Le protocole doit avoir reçu un avis éthique et recueillir le consentement éclairé des "
        "participants avant toute collecte.",
    ]:
        p = document.add_paragraph(style="List Bullet")
        run = p.add_run(texte)
        run.font.size = Pt(10)

    _pied_de_page(document, f"{project.code} — Protocole d'évaluation d'impact — {APP_NAME}")
    return _sauver(document)
